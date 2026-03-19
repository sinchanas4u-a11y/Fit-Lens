"""
Flask Backend API for Body Measurement System
Updated workflow: Upload -> Detect Reference -> YOLOv8 Masking -> MediaPipe -> Measurements
Includes Live Camera Mode via WebSockets
"""
from flask import Flask, request, jsonify
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import cv2
import numpy as np
import base64
import io
from PIL import Image
import traceback
import sys
import os
import hashlib
import tempfile
import time
import json
import datetime
import xml.etree.ElementTree as ET
from flask import Flask, request, jsonify, send_file

# Add parent directory to path to import modules
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from reference_detector import ReferenceDetector
from backend.measurement_engine import MeasurementEngine
from segmentation_model import SegmentationModel
try:
    from smpl.smpl_pipeline import (
        run_smpl_pipeline,
        run_multiview_smpl_pipeline,
    )
except Exception as smpl_import_err:
    print(f"SMPL pipeline not available: "
          f"{smpl_import_err}")
    def run_smpl_pipeline(*a, **kw):
        return {"success": False,
                "error": "SMPL unavailable"}
    def run_multiview_smpl_pipeline(*a, **kw):
        return {"success": False,
                "error": "SMPL unavailable"}
from smpl.smpl_estimator import SMPLEstimator
from processing.smplifyx_runner import run_smplifyx
from processing.smplifyx_reader import SMPLifyXReader

# LandmarkDetector will be imported and initialized on startup
landmark_detector = None

def get_landmark_detector():
    """Get the landmark detector instance (lazy load if not initialized)."""
    global landmark_detector
    if landmark_detector is None:
        try:
            print("Loading MediaPipe landmark detector...")
            from backend.landmark_detector import LandmarkDetector
            landmark_detector = LandmarkDetector()
            print("✓ LandmarkDetector loaded")
        except Exception as e:
            print(f"✗ Could not load landmark_detector: {e}")
            import traceback
            traceback.print_exc()
            return None
    return landmark_detector

# Directory for storing measurement images (metadata only now)
IMAGES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "measurement_images")
os.makedirs(IMAGES_DIR, exist_ok=True)

# Directory for storing measurement results (JSON)
RESULTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "measurement_results")
os.makedirs(RESULTS_DIR, exist_ok=True)

# Directory for storing generated 3D meshes
MESHES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_meshes")
os.makedirs(MESHES_DIR, exist_ok=True)

def save_mesh_as_obj(mesh_data, view_name='front', session_id=None):
    """Save mesh data (Plotly format) as a Wavefront OBJ file.

    Vertices in mesh_data are already in centimetres (cm).
    OBJ files conventionally use metres, so we divide by 100.
    The pipeline centres the mesh so Y=0 is the body midpoint;
    we shift by +height/2 so feet sit at Y=0 in the OBJ file.
    """
    if not mesh_data or 'x' not in mesh_data:
        return None

    subdir   = session_id if session_id else view_name
    view_dir = os.path.join(MESHES_DIR, subdir)
    os.makedirs(view_dir, exist_ok=True)
    obj_path = os.path.join(view_dir, "000.obj")

    try:
        xs = mesh_data['x']
        ys = mesh_data['y']
        zs = mesh_data['z']

        # Pipeline centres mesh so midpoint is Y=0.
        # Shift so feet (min Y) sit at Y=0 in the OBJ.
        y_min_cm = min(ys)
        y_shift  = -y_min_cm   # positive → lifts mesh up

        with open(obj_path, 'w') as f:
            f.write(f"# SMPL Mesh OBJ export - Session: {session_id or 'default'}\n")
            # Convert cm → metres for OBJ
            for x, y, z in zip(xs, ys, zs):
                f.write(f"v {x/100:.6f} {(y + y_shift)/100:.6f} {z/100:.6f}\n")
            # Faces (OBJ is 1-indexed)
            for i, j, k in zip(mesh_data['i'], mesh_data['j'], mesh_data['k']):
                f.write(f"f {i+1} {j+1} {k+1}\n")

        print(f"✓ Saved mesh to {obj_path}  "
              f"(y_shift={y_shift/100:.3f} m, "
              f"vertices={len(xs)})")
        return obj_path
    except Exception as e:
        print(f"Error saving OBJ: {e}")
        return None

# Global set to track hashes of results to prevent duplicates
saved_result_hashes = set()

def load_existing_result_hashes():
    """Load existing result hashes from results.jsonl on startup and remove duplicates."""
    results_path = os.path.join(RESULTS_DIR, "results.jsonl")
    if os.path.exists(results_path):
        unique_entries = []
        try:
            with open(results_path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.strip():
                        try:
                            entry = json.loads(line)
                            if 'data' in entry:
                                data_str = json.dumps(entry['data'], sort_keys=True)
                                digest = hashlib.md5(data_str.encode()).hexdigest()
                                if digest not in saved_result_hashes:
                                    saved_result_hashes.add(digest)
                                    unique_entries.append(line.strip())
                        except json.JSONDecodeError:
                            continue # Skip malformed lines
            
            # Rewrite file with only unique entries
            with open(results_path, 'w', encoding='utf-8') as f:
                for entry in unique_entries:
                    f.write(entry + '\n')
                    
            print(f"Loaded {len(saved_result_hashes)} unique result hashes. Removed duplicates if any.")
        except Exception as e:
            print(f"Error loading existing result hashes: {e}")

# Load existing hashes when the app starts
load_existing_result_hashes()

def save_body_measurements(results):
    """Save the final measurement results to a JSONL file."""
    try:
        results_path = os.path.join(RESULTS_DIR, "results.jsonl")
        # Add metadata to the result
        # We strip out large base64 strings if they exist to keep the file clean
        # IMPORTANT: Do this on a copy to avoid mutating the response being sent to the frontend
        import copy
        results_copy = copy.deepcopy(results)
        if 'results' in results_copy:
            for view in results_copy['results']:
                if isinstance(results_copy['results'][view], dict):
                    results_copy['results'][view].pop('original_image', None)
                    results_copy['results'][view].pop('visualization', None)
                    results_copy['results'][view].pop('mask', None)

        # Check for duplicates before saving
        data_str = json.dumps(results_copy, sort_keys=True)
        digest = hashlib.md5(data_str.encode()).hexdigest()
        
        if digest in saved_result_hashes:
            print(f"✓ Measurement results ignored (duplicate entry)")
            return
            
        saved_result_hashes.add(digest)

        store_data = {
            "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "data": results_copy
        }

        with open(results_path, "a", encoding='utf-8') as f:
            f.write(json.dumps(store_data) + "\n")
        print(f"✓ Measurement results saved to {results_path}")
    except Exception as e:
        print(f"Error saving measurement results: {e}")
        traceback.print_exc()

# Global set to track hashes of images we've already saved to prevent duplicates
saved_image_hashes = set()

def load_existing_hashes():
    """Load existing image hashes from metadata.jsonl on startup."""
    metadata_path = os.path.join(IMAGES_DIR, "metadata.jsonl")
    if os.path.exists(metadata_path):
        try:
            with open(metadata_path, 'r') as f:
                for line in f:
                    if line.strip():
                        try:
                            data = json.loads(line)
                            if 'hash' in data:
                                saved_image_hashes.add(data['hash'])
                        except json.JSONDecodeError:
                            continue
            print(f"Loaded {len(saved_image_hashes)} existing image hashes to prevent duplicates.")
        except Exception as e:
            print(f"Error loading existing hashes: {e}")

# Load existing hashes when the app starts
load_existing_hashes()

def save_measurement_image(image_data, view_name, source='upload'):
    """
    Save the measurement image to disk with metadata.
    Args:
        image_data: base64 string or numpy array
        view_name: 'front', 'side', 'right', 'left', etc.
        source: 'upload' or 'camera'
    """
    try:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        filename = f"{timestamp}_{source}_{view_name}.png"
        filepath = os.path.join(IMAGES_DIR, filename)
        
        # Determine image format and extract bytes for hashing
        img_bytes = None
        if isinstance(image_data, str):
            # Base64 string
            if ',' in image_data:
                image_data = image_data.split(',')[1]
            img_bytes = base64.b64decode(image_data)
        elif isinstance(image_data, np.ndarray):
            # Numpy array: encode back to bytes to get a consistent hash
            # If we don't encode, numpy arrays won't match base64 uploads correctly
            success, buffer = cv2.imencode('.png', image_data)
            if success:
                img_bytes = buffer.tobytes()
        else:
            print(f"Warning: Unsupported image format for saving: {type(image_data)}")
            return
            
        if img_bytes is None:
            print("Warning: Failed to extract image bytes for deduplication.")
            return

        # Compute MD5 hash and check for duplicates
        img_hash = hashlib.md5(img_bytes).hexdigest()
        if img_hash in saved_image_hashes:
            print(f"Duplicate image detected ({view_name}). Skipping save.")
            return
            
        # Add new hash to the tracked set
        saved_image_hashes.add(img_hash)

        # Save to disk - DISABLED as per user request to store results instead of images
        # if isinstance(image_data, str):
        #     with open(filepath, "wb") as f:
        #         f.write(img_bytes)
        # elif isinstance(image_data, np.ndarray):
        #     cv2.imwrite(filepath, image_data)
        
        print(f"Image deduplication check passed for {view_name} (File writing skipped).")

        # Save metadata info to keep track of seen images
        metadata = {
            "filename": filename,
            "timestamp": timestamp,
            "view": view_name,
            "source": source,
            "hash": img_hash
        }
        
        metadata_path = os.path.join(IMAGES_DIR, "metadata.jsonl")
        with open(metadata_path, "a") as f:
            f.write(json.dumps(metadata) + "\n")
            
        print(f"Saved measurement image: {filename}")
        
    except Exception as e:
        print(f"Error saving measurement image: {e}")
        traceback.print_exc()

app = Flask(__name__)
CORS(app)
socketio = SocketIO(app, cors_allowed_origins="*")

# Models will be initialized below with full error handling

# Live Session State
class LiveSession:
    def __init__(self):
        self.reset()

    def reset(self):
        self.captured_images = {} # Original base64
        self.processed_results = {} # Measurements + Visualizations
        self.current_view = 'front'  # front, right, back, left
        self.stability_start_time = None
        self.is_stable = False
        self.last_instruction = ""
        self.last_instruction_time = 0
        self.user_height_cm = 0
        self.scale_factor = 0

live_session = LiveSession()

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    ld = get_landmark_detector()
    return jsonify({
        'status': 'healthy',
        'models_loaded': {
            'yolov8_segmentation': segmentation_model.model is not None,
            'mediapipe_landmarks': ld is not None and hasattr(ld, 'pose') and ld.pose is not None,
            'reference_detector': True
        }
    })

# --- WebSocket Events for Live Camera ---

@socketio.on('connect')
def handle_connect():
    print('Client connected')
    live_session.reset()

@socketio.on('disconnect')
def handle_disconnect():
    print('Client disconnected')

@socketio.on('reset_session')
def handle_reset():
    live_session.reset()
    print('Session reset')

@socketio.on('process_frame')
def handle_frame(data):
    try:
        # Decode image
        image_data = data.get('image')
        view = data.get('view')
        user_height = data.get('user_height')
        height_unit = data.get('height_unit')

        if not image_data:
            return

        # Update user height if provided
        if user_height:
            live_session.user_height_cm = _normalize_height_to_cm(
                user_height,
                height_unit,
                fallback=live_session.user_height_cm
            )

        # Decode base64 to OpenCV image
        img = decode_image(image_data)
        if img is None:
            return

        # Process frame for alignment
        alignment, instruction, countdown = process_alignment(img, view)

        # Check if we should speak the instruction
        should_speak = False
        current_time = time.time()
        if instruction != live_session.last_instruction and (current_time - live_session.last_instruction_time > 3):
            should_speak = True
            live_session.last_instruction = instruction
            live_session.last_instruction_time = current_time

        # Auto-capture logic
        if countdown == 0:
            # Capture!
            print(f"Capturing {view} view!")
            live_session.captured_images[view] = image_data # Store base64
            
            # Save captured image
            save_measurement_image(image_data, view, source='camera')
            
            # Determine next view
            next_view = get_next_view(view)
            live_session.current_view = next_view
            live_session.stability_start_time = None # Reset timer
            
            # Generate "after capture" voice alert
            after_capture_alerts = {
                'front': 'Front view captured.',
                'right': 'Right side view captured.',
                'back': 'Back view captured.',
                'left': 'Left side view captured.'
            }
            
            # Simple "turn left" message for all except last
            before_msg = 'Turn towards your left.' if next_view != 'complete' else ''
            
            # Combine messages
            after_msg = after_capture_alerts.get(view, '')
            voice_message = f"{after_msg} {before_msg}".strip()
            
            emit('capture_complete', {
                'view': view,
                'image': image_data,
                'next_view': next_view,
                'voice_message': voice_message
            })
            
            # Note: Final processing is now triggered by 'finalize_session' from frontend


        emit('frame_processed', {
            'alignment': alignment,
            'instruction': instruction,
            'countdown': countdown if countdown is not None else None,
            'speak': should_speak
        })

    except Exception as e:
        print(f"Error processing frame: {e}")
        traceback.print_exc()

@socketio.on('retake_view')
def handle_retake(data):
    view = data.get('view')
    print(f"Retake requested for {view}")
    if view in live_session.captured_images:
        del live_session.captured_images[view]
    if view in live_session.processed_results:
        del live_session.processed_results[view]
    live_session.stability_start_time = None
    live_session.current_view = view

@socketio.on('process_selection')
def handle_process_selection(data):
    try:
        view = data.get('view')
        image_data = data.get('image')
        selection_type = data.get('type') # 'auto' or 'manual'
        manual_landmarks = data.get('landmarks', [])
        height_unit = data.get('height_unit', 'cm')
        user_height_cm = _normalize_height_to_cm(
            data.get('user_height'),
            height_unit,
            fallback=live_session.user_height_cm
        )

        # Keep session height in canonical centimeters for all later scale math.
        if user_height_cm and user_height_cm > 0:
            live_session.user_height_cm = user_height_cm

        # If client didn't pass image payload for this step, use captured copy.
        if not image_data and view in live_session.captured_images:
            image_data = live_session.captured_images.get(view)

        img = decode_image(image_data)
        if img is None:
            emit('selection_processed', {'error': 'Invalid image data'})
            return

        if selection_type == 'manual':
            # Reuse process_manual_view
            h, w = img.shape[:2]
            results = process_manual_view(
                {'landmarks': manual_landmarks, 'imageWidth': w, 'imageHeight': h},
                user_height_cm,
                view,
                image=img
            )

            if not results.get('success'):
                emit('selection_processed', {
                    'error': results.get('error', 'Manual processing failed'),
                    'view': view
                })
                return
            
            # Build upload-mode style visuals (mask + landmarks) for consistency.
            scale = live_session.scale_factor
            if view == 'front' or scale == 0:
                ld = get_landmark_detector()
                landmarks = ld.detect(img) if ld is not None else None
                landmarks = _ensure_pixel_landmarks(landmarks, img.shape) if landmarks is not None else None
                if landmarks is not None:
                    nose = landmarks[0]
                    left_ankle = landmarks[27]
                    right_ankle = landmarks[28]
                    height_px = max(left_ankle[1], right_ankle[1]) - nose[1]
                    scale = _compute_scale_from_height_px(user_height_cm, height_px, fallback=scale)
                    if view == 'front' and scale > 0:
                        live_session.scale_factor = scale

            engine_view = _normalize_engine_view(view)
            auto_visuals = process_single_view(img, scale, engine_view, user_height_cm=user_height_cm)
            visualization_b64 = auto_visuals.get('visualization') if auto_visuals.get('success') else results.get('visualization')
            mask_b64 = auto_visuals.get('mask') if auto_visuals.get('success') else results.get('mask')
            smpl_meta = auto_visuals.get('smpl') if auto_visuals.get('success') else None

            # Store in session using normalized payload (base64 strings)
            live_session.processed_results[view] = {
                'measurements': results.get('measurements', {}),
                'visualization': visualization_b64,
                'original_image': image_data,
                'mask': mask_b64,
                'smpl': smpl_meta,
            }
            
            emit('selection_processed', {
                'view': view,
                'next_view': get_next_view(view),
                'visualization': visualization_b64,
                'measurements': results.get('measurements', {}),
                'smpl': smpl_meta,
            })

        else: # auto
            # Process view with the same pipeline used by upload mode.
            ld = get_landmark_detector()
            landmarks = ld.detect(img)
            landmarks = _ensure_pixel_landmarks(landmarks, img.shape) if landmarks is not None else None
            if landmarks is None:
                emit('selection_processed', {'error': 'Could not detect body landmarks. Please retake photo or use Manual Marking.'})
                return

            # Defensive compatibility shim: some legacy paths may still reference
            # compute_measurements directly during auto processing.
            if 'compute_measurements' not in globals() or not callable(globals().get('compute_measurements')):
                def compute_measurements(image, landmarks, scale_factor, view_name):
                    engine_view = _normalize_engine_view(view_name)
                    raw = measurement_engine.calculate_measurements_with_confidence(
                        landmarks, scale_factor, engine_view
                    )
                    formatted = {}
                    for name, data in raw.items():
                        if isinstance(data, tuple) and len(data) == 3:
                            cm_value, confidence, source = data
                            px_value = cm_value / scale_factor if scale_factor and scale_factor > 0 else 0
                            formatted[name] = {
                                'value_cm': round(float(cm_value or 0), 2),
                                'value_px': round(float(px_value or 0), 2),
                                'confidence': round(float(confidence or 0), 3),
                                'source': source,
                                'label': name.replace('_', ' ').title(),
                            }
                    return formatted
                globals()['compute_measurements'] = compute_measurements
            
            # We'll calculate a temporary scale if it's the front view
            scale = live_session.scale_factor
            if view == 'front' or scale == 0:
                nose = landmarks[0]
                left_ankle = landmarks[27]
                right_ankle = landmarks[28]
                height_px = max(left_ankle[1], right_ankle[1]) - nose[1]
                scale = _compute_scale_from_height_px(user_height_cm, height_px, fallback=scale)
                if view == 'front' and scale > 0:
                    live_session.scale_factor = scale

            engine_view = _normalize_engine_view(view)
            auto_results = process_single_view(img, scale, engine_view, user_height_cm=user_height_cm)
            if not auto_results.get('success'):
                emit('selection_processed', {
                    'error': auto_results.get('error', 'Auto processing failed'),
                    'view': view
                })
                return
            
            res = {
                'measurements': auto_results.get('measurements', {}),
                'visualization': auto_results.get('visualization'),
                'original_image': image_data,
                'mask': auto_results.get('mask'),
                'smpl': auto_results.get('smpl'),
            }
            live_session.processed_results[view] = res
            
            emit('selection_processed', {
                'view': view,
                'next_view': get_next_view(view),
                'visualization': res.get('visualization'),
                'measurements': res.get('measurements', {}),
                'smpl': res.get('smpl'),
            })

    except Exception as e:
        print(f"Error in process_selection: {e}")
        traceback.print_exc()
        emit('selection_processed', {'error': str(e)})

@socketio.on('finalize_session')
def handle_finalize():
    print("Finalizing live session...")
    process_all_captured_images()

def get_next_view(current):
    # Order: Front → Right → Back → Left
    order = ['front', 'right', 'back', 'left']
    try:
        idx = order.index(current)
        if idx < len(order) - 1:
            return order[idx + 1]
    except ValueError:
        pass
    return 'complete'


def _normalize_engine_view(view_name):
    """Map live-capture view names to measurement engine view names."""
    if view_name in ('right', 'left', 'side'):
        return 'side'
    return 'front' if view_name == 'back' else view_name


def _estimate_height_from_landmarks_px(landmarks):
    """Estimate body height in pixels using nose-to-ankle span."""
    if landmarks is None or len(landmarks) < 29:
        return 0.0
    try:
        nose = landmarks[0]
        left_ankle = landmarks[27]
        right_ankle = landmarks[28]
        return float(max(left_ankle[1], right_ankle[1]) - nose[1])
    except Exception:
        return 0.0


def _landmarks_to_smpl_list(landmarks, image_shape):
    """Convert landmarks to SMPL input format (33 dict entries with normalized x/y)."""
    if landmarks is None or len(landmarks) < 33:
        return []

    h, w = image_shape[:2]
    smpl_landmarks = []
    for idx in range(33):
        lm = landmarks[idx]
        x = float(lm[0] or 0.0) if lm is not None and len(lm) > 0 else 0.0
        y = float(lm[1] or 0.0) if lm is not None and len(lm) > 1 else 0.0
        visibility = float(lm[2] or 1.0) if lm is not None and len(lm) > 2 else 1.0

        # Normalize only when x/y appear to be in pixels.
        if x > 1.5 or y > 1.5:
            x = x / float(w or 1) if w > 0 else 0.0
            y = y / float(h or 1) if h > 0 else 0.0

        smpl_landmarks.append({
            'x': float(max(0.0, min(1.0, x or 0.0))),
            'y': float(max(0.0, min(1.0, y or 0.0))),
            'visibility': float(max(0.0, min(1.0, visibility or 0.0))),
        })

    return smpl_landmarks


def run_gender_detection(image):
    """Best-effort gender detection. Returns male/female/neutral."""
    try:
        if face_verifier is None or not getattr(face_verifier, 'is_ready', False):
            return 'neutral'
        if not hasattr(face_verifier, 'app') or face_verifier.app is None:
            return 'neutral'

        rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        faces = face_verifier.app.get(rgb)
        if not faces:
            return 'neutral'

        face = max(
            faces,
            key=lambda f: float((f.bbox[2] - f.bbox[0]) * (f.bbox[3] - f.bbox[1]))
        )

        raw_gender = getattr(face, 'sex', None)
        if raw_gender is None:
            raw_gender = getattr(face, 'gender', None)

        if isinstance(raw_gender, str):
            g = raw_gender.strip().lower()
            if g in ('male', 'm', 'man'):
                return 'male'
            if g in ('female', 'f', 'woman'):
                return 'female'

        if isinstance(raw_gender, (int, float)):
            return 'male' if float(raw_gender or 0) >= 0.5 else 'female'
    except Exception:
        pass

    return 'neutral'


def _build_smpl_merged_measurements(mp_measurements, smpl_m, smpl_success):
    """Merge MediaPipe and SMPL measurements with graceful fallback behavior."""

    def _mp_entry(name):
        return mp_measurements.get(name, {}) if isinstance(mp_measurements, dict) else {}

    def _mp_cm(name, fallback=None):
        val = _mp_entry(name).get('value_cm')
        return val if val is not None else fallback

    def _mp_px(name, fallback=None):
        val = _mp_entry(name).get('value_px')
        return val if val is not None else fallback

    chest_width_cm = _mp_cm('chest_width', _mp_cm('chest_circumference'))
    waist_width_cm = _mp_cm('waist_width', _mp_cm('waist_circumference'))
    hip_width_cm = _mp_cm('hip_width')

    if smpl_success:
        chest_circ = smpl_m.get('chest_circumference', chest_width_cm)
        waist_circ = smpl_m.get('waist_circumference', waist_width_cm)
        hip_circ = smpl_m.get('hip_circumference', hip_width_cm)
        circ_source = 'SMPL 3D Model'
    else:
        chest_circ = round(chest_width_cm * 3.0, 2) if chest_width_cm is not None else None
        waist_circ = round(waist_width_cm * 3.0, 2) if waist_width_cm is not None else None
        hip_circ = round(hip_width_cm * 3.0, 2) if hip_width_cm is not None else None
        circ_source = 'Estimated'

    merged = {
        'full_height': {
            'value_cm': _mp_cm('full_height'),
            'value_px': _mp_px('full_height'),
            'source': 'MediaPipe',
            'label': 'Full Height',
        },
        'arm_length': {
            'value_cm': _mp_cm('arm_length'),
            'value_px': _mp_px('arm_length'),
            'source': 'MediaPipe',
            'label': 'Arm Length',
        },
        'leg_length': {
            'value_cm': _mp_cm('leg_length'),
            'value_px': _mp_px('leg_length'),
            'source': 'MediaPipe',
            'label': 'Leg Length',
        },
        'torso_length': {
            'value_cm': _mp_cm('torso_length'),
            'value_px': _mp_px('torso_length'),
            'source': 'MediaPipe',
            'label': 'Torso Length',
        },
        'chest_circumference': {
            'value_cm': chest_circ,
            'value_px': None,
            'source': circ_source,
            'label': 'Chest Circumference',
        },
        'waist_circumference': {
            'value_cm': waist_circ,
            'value_px': None,
            'source': circ_source,
            'label': 'Waist Circumference',
        },
        'hip_circumference': {
            'value_cm': hip_circ,
            'value_px': None,
            'source': circ_source,
            'label': 'Hip Circumference',
        },
        'shoulder_width': {
            'value_cm': smpl_m.get('shoulder_width', _mp_cm('shoulder_width')),
            'value_px': _mp_px('shoulder_width'),
            'source': 'SMPL + MediaPipe',
            'label': 'Shoulder Width',
        },
        'chest_width': {
            'value_cm': smpl_m.get('chest_width', chest_width_cm),
            'value_px': _mp_px('chest_width', _mp_px('chest_circumference')),
            'source': 'SMPL + MediaPipe',
            'label': 'Chest Width',
        },
        'waist_width': {
            'value_cm': smpl_m.get('waist_width', waist_width_cm),
            'value_px': _mp_px('waist_width', _mp_px('waist_circumference')),
            'source': 'SMPL + MediaPipe',
            'label': 'Waist Width',
        },
        'hip_width': {
            'value_cm': smpl_m.get('hip_width', hip_width_cm),
            'value_px': _mp_px('hip_width'),
            'source': 'SMPL + MediaPipe',
            'label': 'Hip Width',
        },
    }

    # Keep any additional MediaPipe values that are not part of the merged schema.
    for key, value in mp_measurements.items():
        if key not in merged:
            merged[key] = value

    cleaned = {}
    for key, value in merged.items():
        if not isinstance(value, dict):
            continue
        cm = value.get('value_cm')
        px = value.get('value_px')
        # Always keep circumference rows even when value is None —
        # the frontend will show "—" so the user knows the field exists.
        is_circumference = key in ('chest_circumference', 'waist_circumference', 'hip_circumference')
        if cm is None and px is None and not is_circumference:
            continue
        cleaned[key] = value

    return cleaned, circ_source


def _build_smpl_mesh_data(smpl_result, user_height_cm, gender='neutral'):
    """
    Build mesh data from SMPLify-X output.
    Uses output/meshes/front/000.obj
    """
    try:
        if not user_height_cm or float(
            user_height_cm
        ) <= 0:
            return None

        import glob
        import os
        import sys

        # Find project root
        project_root = os.path.dirname(
            os.path.dirname(
                os.path.abspath(__file__)
            )
        )

        # Always use front mesh — best pose
        mesh_candidates = [
            os.path.join(
                project_root,
                'output', 'meshes',
                'front', '000.obj'
            ),
            os.path.join(
                project_root,
                'output', 'meshes',
                'front', '000.ply'
            ),
        ]

        # Also search recursively
        found = glob.glob(
            os.path.join(
                project_root,
                'output', 'meshes',
                '**', '*.obj'
            ),
            recursive=True
        )

        mesh_path = None
        # Prefer front mesh
        for c in mesh_candidates:
            if os.path.exists(c):
                mesh_path = c
                break

        # Fallback to any found mesh
        if not mesh_path and found:
            # Prefer front in name
            front_meshes = [
                f for f in found
                if 'front' in f.lower()
            ]
            mesh_path = (
                front_meshes[0]
                if front_meshes
                else found[0]
            )

        if not mesh_path:
            print("No SMPLify-X mesh found")
            return None

        print(f"Loading mesh: {mesh_path}")

        # Add processing to path
        if project_root not in sys.path:
            sys.path.insert(0, project_root)

        from processing.smplifyx_reader import (
            SMPLifyXReader
        )

        reader = SMPLifyXReader(mesh_path)

        # Get measurements for beta fitting
        smplx_meas = reader.extract_measurements(
            float(user_height_cm)
        )

        # Find SMPL model path for beta fitting
        smpl_model_path = os.path.join(
            project_root, 'models', 'smpl'
        )

        # Export with beta fitting if models exist
        mesh_data = reader.export_for_plotly(
            user_height_cm = float(
                user_height_cm
            ),
            measurements   = smplx_meas,
            model_path     = smpl_model_path
            if os.path.exists(smpl_model_path)
            else None,
            gender         = gender or 'neutral'
        )

        if mesh_data:
            print(f"Mesh ready: "
                  f"{len(mesh_data['x'])} "
                  f"vertices")

        return mesh_data

    except Exception as e:
        print(f"_build_smpl_mesh_data "
              f"error: {e}")
        import traceback
        traceback.print_exc()
        return None


def _ensure_pixel_landmarks(landmarks, image_shape):
    """
    Ensure landmarks are in pixel coordinates.
    If landmarks are normalized (0..1), convert to pixels using image width/height.
    """
    if landmarks is None:
        return None

    if len(landmarks) == 0:
        return landmarks

    h, w = image_shape[:2]
    lm = np.array(landmarks, dtype=np.float32).copy()

    max_x = float(np.max(np.abs(lm[:, 0])) or 0.0) if lm.shape[1] >= 1 else 0.0
    max_y = float(np.max(np.abs(lm[:, 1])) or 0.0) if lm.shape[1] >= 2 else 0.0

    # Normalized landmarks are typically <= 1.0 (allow small tolerance).
    if max_x <= 1.5 and max_y <= 1.5:
        lm[:, 0] = lm[:, 0] * float(w or 1)
        lm[:, 1] = lm[:, 1] * float(h or 1)

    return lm


def _compute_scale_from_height_px(user_height_cm, height_px, fallback=0.0):
    """Compute cm-per-pixel scale using the correct formula."""
    if user_height_cm is None:
        return fallback
    try:
        user_height_cm = float(user_height_cm)
        height_px = float(height_px)
    except Exception:
        return fallback

    if user_height_cm <= 0 or height_px <= 0:
        return fallback

    # Correct conversion: cm_per_pixel = user_height_cm / pixel_height
    return user_height_cm / height_px


def _normalize_height_to_cm(user_height, height_unit='cm', fallback=0.0):
    """Normalize user height input to centimeters."""
    if user_height is None:
        return fallback
    try:
        height_cm = float(user_height)
    except Exception:
        return fallback

    unit = str(height_unit or 'cm').strip().lower()
    if unit == 'inches':
        height_cm *= 2.54
    elif unit == 'feet':
        height_cm *= 30.48

    return height_cm if height_cm > 0 else fallback


def compute_measurements(image, landmarks, scale_factor, view_name):
    """
    Compute measurements in the same shape expected by the live-camera frontend.
    Returns values with value_cm/value_px/source metadata.
    """
    if landmarks is None:
        return {}

    engine_view = _normalize_engine_view(view_name)
    raw = measurement_engine.calculate_measurements_with_confidence(
        landmarks, scale_factor, engine_view
    )

    formatted = {}
    for name, data in raw.items():
        if isinstance(data, tuple) and len(data) == 3:
            cm_value, confidence, source = data
            px_value = cm_value / scale_factor if scale_factor and scale_factor > 0 else 0
            formatted[name] = {
                'value_cm': round(float(cm_value or 0), 2),
                'value_px': round(float(px_value or 0), 2),
                'confidence': round(float(confidence or 0), 3),
                'source': source,
                'label': name.replace('_', ' ').title(),
            }

    return formatted


def visualize_measurements(image, landmarks, measurements):
    """Render a landmark visualization image for live-camera results."""
    if image is None:
        return None

    vis_image = image.copy()
    ld = get_landmark_detector()
    if ld is not None and landmarks is not None:
        vis_image = ld.draw_landmarks(vis_image, landmarks)

    # Overlay a compact summary so live view has immediate feedback.
    if isinstance(measurements, dict):
        y = 28
        for name, data in list(measurements.items())[:6]:
            value = data.get('value_cm') if isinstance(data, dict) else None
            if value is None:
                continue
            text = f"{name.replace('_', ' ')}: {value:.1f} cm"
            cv2.putText(
                vis_image,
                text,
                (12, y),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.55,
                (0, 255, 255),
                2,
                cv2.LINE_AA,
            )
            y += 24

    return vis_image


def process_alignment(image, view):
    """
    Check if user is aligned (centered + distance + full body visible)
    Returns: alignment_status ('red', 'green'), instruction, countdown
    
    Requirements for GREEN:
    1. Full body visible (all key landmarks detected with good confidence)
    2. User at approximately 1 meter distance
    3. User centered in frame
    """
    h, w, _ = image.shape
    ld = get_landmark_detector()
    if ld is None:
        return 'red', 'Landmark detector unavailable', None
    landmarks = ld.detect(image)
    landmarks = _ensure_pixel_landmarks(landmarks, image.shape) if landmarks is not None else None

    if landmarks is None:
        live_session.stability_start_time = None
        return 'red', 'No person detected. Please stand in front of camera.', None

    # 1. CHECK FULL BODY VISIBILITY
    # Verify all critical landmarks are visible and have good confidence
    # MediaPipe landmarks: 0=nose, 11,12=shoulders, 23,24=hips, 27,28=ankles
    critical_landmarks = [0, 11, 12, 23, 24, 27, 28]
    
    for idx in critical_landmarks:
        if idx >= len(landmarks):
            live_session.stability_start_time = None
            return 'red', 'Full body not visible. Step back.', None
        
        # Check if landmark is visible (confidence > 0.5)
        # MediaPipe returns [x, y, visibility] for each landmark
        if len(landmarks[idx]) >= 3 and landmarks[idx][2] < 0.5:
            live_session.stability_start_time = None
            return 'red', 'Full body not visible. Adjust position.', None
    
    # Check if feet are visible (ankles should be in frame)
    left_ankle = landmarks[27]
    right_ankle = landmarks[28]
    
    # Ankles should not be at the very edge of frame
    if left_ankle[1] > h * 0.95 or right_ankle[1] > h * 0.95:
        live_session.stability_start_time = None
        return 'red', 'Step back. Feet not fully visible.', None
    
    # Check if head is visible (nose should not be at top edge)
    nose = landmarks[0]
    if nose[1] < h * 0.05:
        live_session.stability_start_time = None
        return 'red', 'Move back. Head too close to edge.', None

    # 2. CHECK CENTERING (horizontal alignment)
    left_shoulder = landmarks[11]
    right_shoulder = landmarks[12]
    center_x = (left_shoulder[0] + right_shoulder[0]) / 2
    
    frame_center_x = w / 2
    offset_x = abs(center_x - frame_center_x)
    threshold_x = w * 0.08  # 8% tolerance

    if offset_x > threshold_x:
        live_session.stability_start_time = None
        direction = "left" if center_x < frame_center_x else "right"
        return 'red', f'Move {direction} to center yourself.', None

    # 3. CHECK DISTANCE (target: 1 meter)
    # Calculate person's height in pixels
    ankle_y = max(left_ankle[1], right_ankle[1])
    height_px = ankle_y - nose[1]
    
    # For 1 meter distance, person should occupy approximately 60-80% of frame height
    target_ratio_min = 0.60
    target_ratio_max = 0.80
    current_ratio = height_px / h
    
    if current_ratio < target_ratio_min:
        live_session.stability_start_time = None
        return 'red', 'Move closer. Stand at 1 meter distance.', None
    elif current_ratio > target_ratio_max:
        live_session.stability_start_time = None
        return 'red', 'Move back. Stand at 1 meter distance.', None

    # 4. ALL CONDITIONS MET - START/CONTINUE COUNTDOWN
    if live_session.stability_start_time is None:
        live_session.stability_start_time = time.time()
        return 'green', 'Perfect! Hold still...', 5
    
    # Calculate countdown (5 seconds)
    elapsed = time.time() - live_session.stability_start_time
    remaining = max(0, 5 - int(elapsed))
    

    return 'green', 'Hold still...', remaining

def process_all_captured_images():
    """
    Process all views, favoring manually marked/pre-processed results
    """
    print("Processing all captured images...")
    try:
        final_results = {}
        scale = live_session.scale_factor
        
        # If scale is still 0 (shouldn't happen if front was captured), try to calculate it
        if scale == 0 and 'front' in live_session.captured_images:
            front_img = decode_image(live_session.captured_images['front'])
            ld = get_landmark_detector()
            landmarks = ld.detect(front_img)
            landmarks = _ensure_pixel_landmarks(landmarks, front_img.shape) if landmarks is not None else None
            if landmarks is not None:
                nose = landmarks[0]
                left_ankle = landmarks[27]
                right_ankle = landmarks[28]
                height_px = max(left_ankle[1], right_ankle[1]) - nose[1]
                scale = _compute_scale_from_height_px(live_session.user_height_cm, height_px, fallback=scale)
                if scale > 0:
                    live_session.scale_factor = scale

        views = ['front', 'right', 'back', 'left']
        for v in views:
            if v in live_session.processed_results:
                # Use already processed result (manual or auto-selection)
                res = live_session.processed_results[v]

                visualization_b64 = res.get('visualization')
                mask_b64 = res.get('mask')
                measurements = res.get('measurements', {})
                smpl_meta = res.get('smpl')

                # If any output is missing, rebuild from the same upload-mode pipeline.
                if (not visualization_b64 or not mask_b64 or not measurements) and v in live_session.captured_images:
                    img = decode_image(live_session.captured_images[v])
                    if img is not None:
                        engine_view = _normalize_engine_view(v)
                        rebuilt = process_single_view(img, scale, engine_view, user_height_cm=live_session.user_height_cm)
                        if rebuilt.get('success'):
                            visualization_b64 = visualization_b64 or rebuilt.get('visualization')
                            mask_b64 = mask_b64 or rebuilt.get('mask')
                            measurements = measurements or rebuilt.get('measurements', {})
                            smpl_meta = smpl_meta or rebuilt.get('smpl')
                
                # Results are stored in normalized base64 format
                final_results[v] = {
                    'measurements': measurements,
                    'visualization': visualization_b64,
                    'original_image': res.get('original_image') or live_session.captured_images.get(v),
                    'mask': mask_b64,
                    'smpl': smpl_meta,
                }
            elif v in live_session.captured_images:
                # Fallback to auto-processing if not already processed
                img = decode_image(live_session.captured_images[v])
                if img is not None:
                    engine_view = _normalize_engine_view(v)
                    auto_results = process_single_view(img, scale, engine_view, user_height_cm=live_session.user_height_cm)
                    if auto_results.get('success'):
                        final_results[v] = {
                            'measurements': auto_results.get('measurements', {}),
                            'visualization': auto_results.get('visualization'),
                            'original_image': live_session.captured_images.get(v),
                            'mask': auto_results.get('mask'),
                            'smpl': auto_results.get('smpl'),
                        }

        # 3. Emit final complete payload
        payload = {
            'success': True,
            'results': final_results,
            'calibration': {
                'user_height_cm': live_session.user_height_cm,
                'scale_factor': scale
            }
        }
        
        save_body_measurements(payload)
        socketio.emit('processing_complete', payload)
        
    except Exception as e:
        print(f"Error in final processing: {e}")
        traceback.print_exc()
        socketio.emit('error', {'message': str(e)})

from face_verifier import FaceVerifier

# Initialize models
print("Initializing models...")
try:
    reference_detector = ReferenceDetector()
    measurement_engine = MeasurementEngine()
    segmentation_model = SegmentationModel(model_size='n')  # YOLOv8 nano for speed
    # Initialize landmark detector on startup
    print("Loading MediaPipe landmark detector...")
    landmark_detector = get_landmark_detector()
    if landmark_detector is None:
        print("⚠ Warning: LandmarkDetector failed to initialize - some features will be unavailable")
    else:
        print("✓ LandmarkDetector initialized successfully")
    print("✓ Core models initialized")
except Exception as e:
    print(f"Error initializing core models: {e}")
    raise

# Initialize face verifier for identity verification
try:
    face_verifier = FaceVerifier(model_name="buffalo_l", det_size=(640, 640))
    print("✓ Face verifier initialized")
except Exception as e:
    print(f"Warning: Could not initialize FaceVerifier: {e}")
    face_verifier = None

print("✓ Models initialized")

# --- Existing API Routes ---

@app.route('/mesh/<view>/000.obj')
@app.route('/mesh/<session_id>/<view>/000.obj')
def serve_mesh_obj(view, session_id=None):
    """Serve the 000.obj file for the specified session and view."""
    if session_id:
        obj_path = os.path.join(MESHES_DIR, session_id, "000.obj")
        if os.path.exists(obj_path):
            return send_file(obj_path, mimetype='text/plain')
            
    legacy_path = os.path.join(MESHES_DIR, view, "000.obj")
    if os.path.exists(legacy_path):
         return send_file(legacy_path, mimetype='text/plain')
         
    return jsonify({"error": "Mesh not found"}), 404

@app.route('/api/verify-identity', methods=['POST'])
def verify_identity():
    """
    Verify identity between two images using ArcFace cosine similarity.
    Thresholds:
    - similarity >= 0.65: Same person (verified)
    - 0.50 <= similarity < 0.65: Uncertain range (verified with warning)
    - similarity < 0.50: Different person (not verified)
    - No face detected: Not verified
    """
    try:
        if face_verifier is None:
            return jsonify({
                'success': False,
                'verified': False,
                'error': 'Face verification service unavailable'
            }), 503
        
        data = request.json
        print("\n" + "="*60)
        print("IDENTITY VERIFICATION REQUEST")
        print(f"Time: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        front_b64 = data.get('front_image', '')
        side_b64 = data.get('side_image', '')
        
        if not front_b64 or not side_b64:
            return jsonify({
                'success': False,
                'verified': False,
                'error': 'Both front and side images are required'
            }), 400
        
        print(f"Front image payload size: {len(front_b64) / 1024:.2f} KB")
        print(f"Side image payload size: {len(side_b64) / 1024:.2f} KB")
        
        # Decode base64 images
        front_img = decode_image(front_b64)
        side_img = decode_image(side_b64)
        
        if front_img is None or side_img is None:
            return jsonify({
                'success': False,
                'verified': False,
                'error': 'Invalid image data'
            }), 400
        
        # Run face verification
        result = face_verifier.verify_person(front_img, side_img)
        
        # Extract values from result
        verified = result.get('verified', False)
        similarity = result.get('similarity', 0.0)
        threshold = result.get('threshold', 0.65)
        no_face = result.get('no_face', False)
        has_warning = result.get('warning', False)
        error_msg = result.get('error', '')
        issues = result.get('issues', {'front': [], 'side': []})
        
        # Construct response
        response_data = {
            'success': True,
            'verified': verified,
            'similarity': float(similarity or 0),
            'threshold': float(threshold or 0.65),
            'issues': issues
        }
        
        # Build appropriate message based on verification result
        if verified:
            # PRIORITIZE similarity verification:
            # If similarity is high enough, we proceed even if there are "soft" quality issues
            
            response_data['message'] = 'Identity verified successfully'
            response_data['verified'] = True
            
            if has_warning:
                 response_data['warning'] = True
                 response_data['message'] += ' (uncertain range)'
                 
        else:
            # Verification Failed
            response_data['verified'] = False
            
            # 1. No Face Check (Blocking)
            if no_face:
                 response_data['message'] = 'Face verification failed - face not detected clearly'
            else:
                 # 2. Identity Mismatch vs Quality
                 # If similarity is VERY low, it's likely a mismatch.
                 # If similarity is borderline, maybe quality is to blame.
                 # User request: "If mismatch → show mismatch error first. Then optionally show supporting quality hints"
                 
                 main_error = "Identity verification failed"
                 if similarity < 0.40: # Very low similarity
                      main_error += " (Different people detected)"
                 elif similarity < threshold:
                      main_error += " (Mismatch detected)"
                      
                 # Collect hints
                 hints = []
                 if issues['front']:
                     hints.append(f"Front: {', '.join(issues['front'])}")
                 if issues['side']:
                     hints.append(f"Side: {', '.join(issues['side'])}")
                     
                 if hints:
                      # Append hints to main error
                      response_data['message'] = f"{main_error}. Note: {'; '.join(hints)}"
                 else:
                      response_data['message'] = f"{main_error}. Please ensure both photos are of the same person."
        
        print(f"\n📊 Verification Result:")
        print(f"   Verified: {verified}")
        print(f"   Similarity: {similarity:.4f}")
        print(f"   Message: {response_data['message']}")
        print("="*60)
        
        return jsonify(response_data), 200
        
    except Exception as e:
        print(f"✗ Server Error in verify_identity: {e}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'verified': False,
            'error': f'Server error: {str(e)}'
        }), 500

@app.route('/api/process', methods=['POST'])
def process_images():
    """
    Complete workflow:
    1. Upload photos
    2. Verify identity (Face Verification)
    3. Use user's height as reference
    4. YOLOv8-seg masking (detect only human body)
    5. MediaPipe pose (get landmarks)
    6. Compute pixel distances and convert to cm using height
    7. Return measurements as JSON
    """
    try:
        data = request.json
        
        print("\n" + "="*60)
        print("STEP 1: RECEIVING UPLOADED PHOTOS")
        print("="*60)
        
        # Decode images
        front_img = decode_image(data.get('front_image'))
        side_img = decode_image(data.get('side_image'))
        
        if front_img is None:
            return jsonify({'error': 'Front image is required', 'step': 1}), 400
        
        print(f"✓ Front image: {front_img.shape}")
        if side_img is not None:
            print(f"✓ Side image: {side_img.shape}")
            
            # --- SECURITY CHECK: FACE VERIFICATION ---
            print("\n" + "="*60)
            print("STEP 1.5: VERIFYING IDENTITY")
            print("="*60)

            if face_verifier is None:
                return jsonify({
                    'error': 'Face verification service unavailable',
                    'step': 1.5
                }), 503
            
            # Verify that both images belong to the same person
            verification_result = face_verifier.verify_person(front_img, side_img)
            
            # InsightFace returns similarity (higher = more similar)
            similarity = verification_result.get('similarity', 0.0)
            verified = verification_result.get('verified', False)
            issues = verification_result.get('issues', {'front': [], 'side': []})
            
            # Refined Logic: If Verified, proceed (ignore quality issues). If Not Verified, check issues.
            if not verified:
                # Verification failed. 
                # User request: Priority to Mismatch Error.
                
                if verification_result.get('no_face'):
                     return jsonify({'error': 'Face verification failed - face not detected clearly', 'step': 1.5}), 400

                main_error = "Identity verification failed"
                # If similarity is low, emphasize mismatch
                similarity = verification_result.get('similarity', 0.0)
                if similarity < 0.40:
                     main_error += " (Different people detected)"
                
                # Append hints if any
                hints = []
                if issues['front']:
                    hints.append(f"Front: {', '.join(issues['front'])}")
                if issues['side']:
                    hints.append(f"Side: {', '.join(issues['side'])}")
                
                full_error_msg = main_error
                if hints:
                     full_error_msg += f". Issues: {'; '.join(hints)}"
                else:
                     full_error_msg += ". Please ensure both photos are of the same person."

                return jsonify({'error': full_error_msg, 'step': 1.5}), 400
            
            # Log warning if in uncertain range
            if verification_result.get('warning'):
                print(f"⚠ Warning: Similarity in uncertain range (0.50-0.65): {similarity:.4f}")
            # ----------------------------------------

        # Save uploaded images
        if data.get('front_image'):
            save_measurement_image(data.get('front_image'), 'front', source='upload')
            
        if data.get('side_image'):
            save_measurement_image(data.get('side_image'), 'side', source='upload')
        
        # Get user's height
        # Get user's height safely
        try:
            # Priority 1: request.json.get('user_height')
            user_height_cm = float(data.get('user_height') or 0)
            
            # Priority 2: if still 0, check request.json.get('height')
            if user_height_cm <= 0:
                user_height_cm = float(data.get('height') or 0)
                
            # Priority 3: check form data if available
            if user_height_cm <= 0 and request.form.get('height'):
                user_height_cm = float(request.form.get('height') or 0)
                
            # Fallback
            if user_height_cm <= 0:
                user_height_cm = 170.0
                print("⚠ User height missing or invalid, using default 170.0 cm")
        except (TypeError, ValueError):
            user_height_cm = 170.0
            print("⚠ User height parse error, using default 170.0 cm")
        
        print(f"✓ Final user height for processing: {user_height_cm} cm")
        
        print("\n" + "="*60)
        print("STEP 2: CALCULATING SCALE USING USER HEIGHT")
        print("="*60)
        
        # First, we need to detect landmarks to measure height in pixels
        # We'll do a preliminary landmark detection on front image
        print("Detecting landmarks to measure height in pixels...")
        ld = get_landmark_detector()
        if ld is None:
            return jsonify({
                'error': 'Could not initialize landmark detector',
                'step': 2
            }), 503
        temp_landmarks = ld.detect(front_img)
        
        if temp_landmarks is None:
            return jsonify({
                'error': 'Could not detect person in image. Please ensure full body is visible.',
                'step': 2
            }), 400
        
        # Calculate height in pixels (from top of head to feet)
        # Use nose (top) to ankle (bottom) as height approximation
        nose = temp_landmarks[0]  # nose
        left_ankle = temp_landmarks[27]  # left ankle
        right_ankle = temp_landmarks[28]  # right ankle
        
        # Use the lower ankle
        ankle_y = max(left_ankle[1], right_ankle[1])
        height_px = ankle_y - nose[1]
        
        if height_px <= 0:
            return jsonify({
                'error': 'Could not measure height in image. Please ensure full body is visible.',
                'step': 2
            }), 400
        
        # Calculate scale factor
        scale_factor = user_height_cm / height_px
        
        print(f"✓ Height in image: {height_px:.2f} pixels")
        print(f"✓ User height: {user_height_cm} cm")
        print(f"✓ Scale factor: {scale_factor:.4f} cm/px")
        print(f"✓ Formula: {user_height_cm} cm ÷ {height_px:.2f} px = {scale_factor:.4f} cm/px")
        
        # Prepare for SMPL results (now integrated into process_single_view)
        smplx_status = 'unavailable'
        smplx_error = None
        smplx_meas = {}
        mesh_data = None

        # Process front view
        print("\n" + "="*60)
        print("PROCESSING FRONT VIEW")
        print("="*60)

        front_results = process_single_view(
            front_img, scale_factor, 'front', user_height_cm=user_height_cm
        )

        results = {
            'front': front_results
        }

        # Process side view if provided
        if side_img is not None:
            print("\n" + "="*60)
            print("PROCESSING SIDE VIEW")
            print("="*60)

            side_results = process_single_view(
                side_img, scale_factor, 'side', user_height_cm=user_height_cm
            )
            results['side'] = side_results

        # ── Multi-view mesh generation ────────────────────────────────────────
        # When both views are available, replace the front-only SMPL mesh with
        # a multi-view reconstruction that uses front widths + side depths.
        if side_img is not None and front_results.get('success') and results.get('side', {}).get('success'):
            print("\n" + "="*60)
            print("MULTI-VIEW SMPL RECONSTRUCTION")
            print("="*60)
            try:
                # Re-detect landmarks (already done inside process_single_view,
                # but we need them here for the multiview call)
                ld = get_landmark_detector()
                front_lm_raw = ld.detect(front_img) if ld else None
                side_lm_raw  = ld.detect(side_img)  if ld else None

                front_lm_raw = _ensure_pixel_landmarks(front_lm_raw, front_img.shape) if front_lm_raw is not None else None
                side_lm_raw  = _ensure_pixel_landmarks(side_lm_raw,  side_img.shape)  if side_lm_raw  is not None else None

                def _fmt_lm(lm_raw, img):
                    if lm_raw is None:
                        return []
                    h, w = img.shape[:2]
                    return [
                        {
                            'x': float(lm_raw[i][0]) / w,
                            'y': float(lm_raw[i][1]) / h,
                            'visibility': float(lm_raw[i][2]) if len(lm_raw[i]) > 2 else 1.0,
                        }
                        for i in range(len(lm_raw))
                    ]

                front_lm_fmt = _fmt_lm(front_lm_raw, front_img)
                side_lm_fmt  = _fmt_lm(side_lm_raw,  side_img)

                # Get masks from segmentation model
                front_mask_mv = segmentation_model.segment_person(front_img, conf_threshold=0.5)
                side_mask_mv  = segmentation_model.segment_person(side_img,  conf_threshold=0.5)

                detected_gender_mv = run_gender_detection(front_img)

                mv_res = run_multiview_smpl_pipeline(
                    front_landmarks_2d = front_lm_fmt,
                    front_image_width  = front_img.shape[1],
                    front_image_height = front_img.shape[0],
                    user_height_cm     = user_height_cm,
                    gender             = detected_gender_mv or 'neutral',
                    front_mask         = front_mask_mv,
                    side_mask          = side_mask_mv,
                    target_measurements = front_results.get('smpl_target_measurements') or {},
                )

                if mv_res.get('success') and mv_res.get('mesh_data'):
                    # Overwrite front mesh with the multiview result
                    front_results['mesh_data'] = mv_res['mesh_data']
                    front_results['measurements'] = mv_res.get('measurements', front_results.get('measurements', {}))
                    front_results['smpl'] = {
                        'success': True,
                        'fitted_to_user': True,
                        'status_text': mv_res['fit'].get('status_text', '✓ Multi-view reconstruction'),
                        'multiview': True,
                    }
                    print("✓ Multi-view mesh integrated into front results")
                else:
                    print(f"Multi-view pipeline returned no mesh: {mv_res.get('error')}")
            except Exception as mv_err:
                print(f"Multi-view reconstruction failed (keeping single-view mesh): {mv_err}")
                import traceback
                traceback.print_exc()

        # Integrate SMPL results from front view
        if front_results.get('success'):
            smpl_info = front_results.get('smpl', {})
            if smpl_info.get('success'):
                smplx_status = 'success'
                smplx_meas = front_results.get('measurements', {}) # Measurements include SMPL values
                mesh_data = front_results.get('mesh_data')
                print("✓ Internal SMPL posed mesh + measurements integrated from Front View")
            else:
                smplx_status = 'failed'
                smplx_error = smpl_info.get('error')
            
            # Additional metadata for backward compatibility or debugging
            front_results['smplx_status'] = smplx_status
            front_results['smplx_error'] = smplx_error
            front_results['smplx_measurements'] = smplx_meas
            if mesh_data:
                front_results['mesh_data'] = mesh_data
        
        # Prepare final response
        print("\n" + "="*60)
        print("STEP 6: RETURNING MEASUREMENTS AS JSON")
        print("="*60)

        # Prepare final response
        print("\n" + "="*60)
        print("STEP 6: RETURNING MEASUREMENTS AS JSON")
        print("="*60)

        # Generate a unique session ID for this request to ensure fresh mesh
        session_id = f"{int(time.time() * 1000)}"
        mesh_url = f"/mesh/{session_id}/front/000.obj"

        # Force use of front mesh data if available
        front_results = results.get('front', {})
        if isinstance(front_results, dict) and front_results.get('mesh_data'):
            mesh_data = front_results.get('mesh_data')
            smplx_status = 'success'
            
            # Inject session_id into mesh_data so frontend SMPLViewer finds it
            mesh_data['session_id'] = session_id
            mesh_metadata = mesh_data.setdefault('metadata', {})
            mesh_metadata['session_id'] = session_id
            mesh_metadata['mesh_url'] = mesh_url
            mesh_metadata['mesh_format'] = 'obj'
            
            # Export front mesh as OBJ with unique session ID
            try:
                obj_path = save_mesh_as_obj(mesh_data, 'front', session_id=session_id)
                mesh_metadata['mesh_file_exists'] = bool(obj_path and os.path.exists(obj_path))
                mesh_metadata['mesh_file_path'] = obj_path
            except Exception as e:
                print(f"Warning: Failed to export OBJ: {e}")
                mesh_metadata['mesh_file_exists'] = False
                mesh_metadata['mesh_file_path'] = None
        else:
            mesh_data = None
            smplx_status = 'unavailable'

        # Add original images to results for display
        if front_results.get('success'):
            front_results['original_image'] = data.get('front_image')
        if side_img is not None and results.get('side', {}).get('success'):
            results['side']['original_image'] = data.get('side_image')
        
        response = {
            'success': True,
            'mode': 'automatic',  # Set mode for frontend display
            'calibration': {
                'user_height_cm': user_height_cm,
                'height_in_image_px': float(height_px or 0),
                'scale_factor': float(scale_factor or 0),
                'formula': f'{user_height_cm} cm ÷ {height_px:.2f} px = {scale_factor:.4f} cm/px',
                'description': f'1 pixel = {scale_factor:.4f} cm'
            },
            'results': results,
            'mesh_data': mesh_data if mesh_data is not None else (front_results.get('mesh_data') if isinstance(front_results, dict) else None),
            'mesh_url': mesh_url if mesh_data is not None else None,
            'smplx_status': smplx_status,
            'smplx_error': smplx_error,
            'session_id': session_id
        }

        # Prepare final response metadata

        # Store results persistently
        save_body_measurements(response)
        
        print("✓ Processing complete!")
        print("="*60 + "\n")
        
        return jsonify(response)
        
    except Exception as e:
        print(f"\n✗ ERROR: {str(e)}")
        traceback.print_exc()
        return jsonify({
            'error': f'Processing failed: {str(e)}',
            'step': 'unknown'
        }), 500


def process_single_view(image, scale_factor, view_name, user_height_cm=None):
    """
    Process a single view through the complete pipeline:
    Step 3: YOLOv8 segmentation
    Step 4: MediaPipe landmarks
    Step 5: Compute measurements
    """
    try:
        print(f"\nSTEP 3: YOLOv8-SEG MASKING ({view_name.upper()} VIEW)")
        print("-" * 60)
        
        # Run YOLOv8 segmentation to detect only human body
        mask = segmentation_model.segment_person(image, conf_threshold=0.5)
        
        if mask is None:
            print("✗ No person detected in image")
            return {
                'error': 'No person detected',
                'step': 3
            }
        
        print("✓ Human body detected and masked")
        
        # Apply mask to get clean human outline
        masked_image = segmentation_model.apply_mask(image, mask, background_mode='dim')
        
        # Get bounding box of person
        bbox = segmentation_model.get_person_bbox(mask)
        if bbox:
            x, y, w, h = bbox
            print(f"✓ Person bounding box: {w}x{h} at ({x}, {y})")
        
        print(f"\nSTEP 4: MEDIAPIPE POSE LANDMARKS ({view_name.upper()} VIEW)")
        print("-" * 60)
        
        # Run MediaPipe on original image (works better than masked)
        ld = get_landmark_detector()
        if ld is None:
            raise Exception('Landmark detector unavailable')
        landmarks = ld.detect(image)
        landmarks = _ensure_pixel_landmarks(landmarks, image.shape) if landmarks is not None else None
        
        if landmarks is None:
            print("✗ No landmarks detected")
            return {
                'error': 'No landmarks detected',
                'step': 4,
                'mask_available': True
            }
        
        print(f"✓ Detected {len(landmarks)} body landmarks")

        # Step 4.1: Pose formatting for SMPL
        landmarks_formatted = []
        for i in range(len(landmarks)):
            landmarks_formatted.append({
                'x': landmarks[i][0] / image.shape[1],
                'y': landmarks[i][1] / image.shape[0],
                'visibility': landmarks[i][2] if len(landmarks[i]) > 2 else 1.0
            })

        # Step 4.2: Gender detection
        detected_gender = run_gender_detection(image)

        # Step 4.3: Effective height calculation
        try:
            effective_height_cm = float(user_height_cm) if user_height_cm and float(user_height_cm) > 0 else 0.0
        except (TypeError, ValueError):
            effective_height_cm = 0.0

        if effective_height_cm <= 0 and scale_factor:
            try:
                estimated_height_px = _estimate_height_from_landmarks_px(landmarks)
                if estimated_height_px > 0:
                    effective_height_cm = float(scale_factor or 0) * float(estimated_height_px or 0)
            except:
                pass

        # Step 4.5: Edge points extraction (MOVED UP)
        edge_reference_points = None
        print(f"\nSTEP 4.5: EXTRACTING EDGE REFERENCE POINTS ({view_name.upper()} VIEW)")
        print("-" * 60)
        
        if mask is not None:
            try:
                edge_reference_points = ld.extract_body_edge_keypoints(mask, landmarks)
                if edge_reference_points and edge_reference_points.get('is_valid'):
                    print("✓ Edge reference points extracted successfully")
            except Exception as e:
                print(f"Warning: Could not extract edge points: {e}")
                edge_reference_points = None

        # Step 5: Computing measurements (MOVED UP)
        print(f"\nSTEP 5: COMPUTING MEASUREMENTS ({view_name.upper()} VIEW)")
        print("-" * 60)
        
        if scale_factor is None or scale_factor <= 0:
            try:
                nose = landmarks[0]
                left_ankle = landmarks[27]
                right_ankle = landmarks[28]
                height_px = max(left_ankle[1], right_ankle[1]) - nose[1]
                scale_factor = user_height_cm / height_px if height_px > 0 else 1.0
            except:
                scale_factor = 1.0

        try:
            measurements = measurement_engine.calculate_measurements_with_confidence(
                landmarks, scale_factor, view_name, edge_reference_points=edge_reference_points
            )
        except Exception as e:
            print(f"⚠ Measurement error: {e}")
            measurements = {}

        # Prepare target measurements for SMPL personalization
        target_measurements = {}

        # Always include user height — the most reliable anchor
        if effective_height_cm > 0:
            target_measurements['height'] = float(effective_height_cm)

        if measurements:
            for name, m_data in measurements.items():
                if isinstance(m_data, tuple) and len(m_data) == 3:
                     val = m_data[0] # cm value

                     # VALIDATE measurements before passing to SMPL
                     # Realistic ranges:
                     # chest: 70-130 cm
                     # waist: 55-125 cm
                     # hip:   75-135 cm

                     if name == 'chest_circumference':
                         if not (70 <= val <= 130):
                             print(f"WARNING: chest {val:.2f} out of range, recalculating...")
                             chest_width_data = measurements.get('chest_width')
                             if chest_width_data and isinstance(chest_width_data, tuple):
                                 val = chest_width_data[0] * 3.2
                             else:
                                 val = 90.0
                         target_measurements[name] = val

                     elif name == 'waist_circumference':
                         if not (55 <= val <= 125):
                             print(f"WARNING: waist {val:.2f} out of range, capping...")
                             val = max(55, min(125, val))
                         target_measurements[name] = val

                     elif name == 'hip_circumference' or name == 'hip_width':
                         key = 'hip_circumference'
                         if val > 135:
                             print(f"WARNING: hip {val:.2f} too large, recalculating from width...")
                             hip_width_data = measurements.get('hip_width')
                             if hip_width_data and isinstance(hip_width_data, tuple):
                                 h_w = hip_width_data[0]
                                 if 33 <= h_w <= 55:
                                     val = h_w * 3.0
                                 else:
                                     waist_data = measurements.get('waist_circumference')
                                     w_c = waist_data[0] if waist_data else 85
                                     val = w_c + 15
                             else:
                                 val = 100.0
                         elif val < 75:
                             val = 85.0
                         target_measurements[key] = val

                     elif name == 'shoulder_width':
                         if 25.0 <= val <= 65.0:
                             target_measurements[name] = val

                     elif name in ('arm_length', 'leg_length'):
                         if val > 10.0:
                             target_measurements[name] = val

                     else:
                         target_measurements[name] = val
        
        # Step 5.5: SMPL integration
        print(f"\nSTEP 5.5: GENERATING PERSONALIZED 3D BODY MODEL ({view_name.upper()} VIEW)")
        print("-" * 60)
        
        smpl_success = False
        smpl_m = {}
        smpl_error = None
        smpl_mesh_data = None
        smpl_fit_info = {}
        
        if effective_height_cm > 0:
            print(f"Running personalized SMPL pipeline for {view_name} view...")
            smpl_res = run_smpl_pipeline(
                landmarks_2d   = landmarks_formatted,
                image_width    = image.shape[1],
                image_height   = image.shape[0],
                user_height_cm = effective_height_cm,
                gender         = detected_gender or 'neutral',
                view_type      = view_name,
                use_neutral_pose = True,
                target_measurements = target_measurements
            )
            
            if smpl_res.get('success'):
                smpl_success = True
                smpl_mesh_data = smpl_res.get('mesh_data')
                smpl_m = smpl_res.get('measurements', {})
                smpl_fit_info = smpl_res.get('fit', {})
                print(f"✓ SMPL mesh refined: {len(smpl_mesh_data['x']) if smpl_mesh_data else 0} vertices")
                
                # Export OBJ file immediately so the viewer can fetch it
                if smpl_mesh_data is not None and smpl_mesh_data.get('session_id') is None:
                    temp_session_id = f"{int(time.time() * 1000)}"
                    mesh_url = f"/mesh/{temp_session_id}/{view_name}/000.obj"
                    
                    smpl_mesh_data['session_id'] = temp_session_id
                    mesh_metadata = smpl_mesh_data.setdefault('metadata', {})
                    mesh_metadata['session_id'] = temp_session_id
                    mesh_metadata['mesh_url'] = mesh_url
                    mesh_metadata['mesh_format'] = 'obj'
                    
                    try:
                        obj_path = save_mesh_as_obj(smpl_mesh_data, view_name, session_id=temp_session_id)
                        mesh_metadata['mesh_file_exists'] = bool(obj_path and os.path.exists(obj_path))
                        mesh_metadata['mesh_file_path'] = obj_path
                    except Exception as e:
                        print(f"Warning: Failed to export OBJ in process_single_view: {e}")
                        mesh_metadata['mesh_file_exists'] = False
                        mesh_metadata['mesh_file_path'] = None

            else:
                smpl_error = smpl_res.get('error', 'SMPL fitting failed')
                print(f'SMPL pipeline failed: {smpl_error} — generating default neutral body model...')
                # Fallback: always provide a body model even when landmark fitting fails
                try:
                    fallback_est = SMPLEstimator(gender=detected_gender or 'neutral')
                    fb_betas = np.zeros(10, dtype=np.float64)
                    fb_verts = fallback_est.get_vertices(fb_betas)
                    fb_faces = np.asarray(fallback_est.faces, dtype=np.int32)
                    fb_y_span = (float(fb_verts[:, 1].max()) - float(fb_verts[:, 1].min())) * 100.0
                    if fb_y_span > 0 and effective_height_cm > 0:
                        fb_scale = effective_height_cm / fb_y_span
                        fb_cm = fb_verts * 100.0 * fb_scale
                        fb_cm[:, 1] -= (fb_cm[:, 1].max() + fb_cm[:, 1].min()) / 2.0
                        smpl_mesh_data = {
                            'x': fb_cm[:, 0].tolist(),
                            'y': fb_cm[:, 1].tolist(),
                            'z': fb_cm[:, 2].tolist(),
                            'i': fb_faces[:, 0].tolist(),
                            'j': fb_faces[:, 1].tolist(),
                            'k': fb_faces[:, 2].tolist(),
                            'metadata': {
                                'vertex_count': int(fb_cm.shape[0]),
                                'face_count': int(fb_faces.shape[0]),
                                'height_cm': float(effective_height_cm),
                                'gender': detected_gender or 'neutral',
                                'fitted_to_user': False,
                                'fit_status': 'default',
                                'status_text': 'Default body model (estimated from height)',
                                'pose_applied': False,
                                'landmarks_source': 'none',
                            }
                        }
                        smpl_success = True
                        smpl_error = None
                        smpl_fit_info = {
                            'fit_status': 'default',
                            'status_text': 'Default body model (estimated from height)',
                            'fitted_to_user': False,
                            'landmarks_source': 'none',
                            'landmarks_mode': 'unavailable',
                            'visible_landmark_count': 0,
                            'landmark_count': len(landmarks_formatted),
                            'pose_applied': False,
                        }
                        print(f'Fallback SMPL: generated ({len(fb_cm)} vertices)')
                except Exception as fb_err:
                    print(f'Fallback SMPL failed: {fb_err}')
        else:
            smpl_error = 'Insufficient data for SMPL (height unavailable)'

        print(f"\n✓ Measurement engine results processed")
        
        # Format MediaPipe measurements first, then merge with SMPL outputs.
        mp_measurements = {}
        print(f"\n📊 Raw measurements received: {len(measurements)} items")
        print(f"Measurements type: {type(measurements)}")
        print(f"Measurements keys: {list(measurements.keys()) if measurements else 'None'}")
        
        for name, data in measurements.items():
            print(f"\n  Processing measurement: {name}")
            print(f"    Data type: {type(data)}")
            print(f"    Data value: {data}")
            
            if isinstance(data, tuple) and len(data) == 3:
                try:
                    cm_value, confidence, source = data
                    # Guard scale_factor and cm_value
                    sf = float(scale_factor) if scale_factor else 0.0
                    cm = float(cm_value) if cm_value is not None else 0.0
                    
                    pixel_distance = cm / sf if sf > 0 else 0
                    
                    mp_measurements[name] = {
                        'value_cm': round(cm, 2),
                        'value_px': round(float(pixel_distance or 0), 2),
                        'confidence': round(float(confidence or 0), 3),
                        'source': source,
                        'source_raw': source,
                        'source_group': 'MediaPipe',
                        'formula': f"{pixel_distance:.2f} px × {sf:.4f} cm/px = {cm:.2f} cm"
                    }
                    
                    print(f"  ✓ {name}: {cm:.2f} cm (confidence: {float(confidence or 0):.2f})")
                except (TypeError, ValueError) as fe:
                    print(f"  ⚠ Measurement parse error for {name}: {fe}")
                    continue
            else:
                print(f"  ⚠️ Unexpected data format for {name}: {data}")
        
        formatted_measurements, smpl_source_label = _build_smpl_merged_measurements(
            mp_measurements,
            smpl_m,
            smpl_success,
        )

        print(f"\n✓ Formatted {len(formatted_measurements)} merged measurements")
        print(f"Formatted measurements keys: {list(formatted_measurements.keys())}")
        
        # Create visualization with landmarks and edge-based width overlays
        ld = get_landmark_detector()
        if ld is None:
            vis_image = masked_image.copy()
        else:
            vis_image = ld.draw_landmarks(masked_image.copy(), landmarks)

        # Draw outer edge points and true widths from YOLOv8 mask if available
        try:
            if edge_reference_points and edge_reference_points.get('is_valid'):
                # Colors (BGR)
                BLUE = (255, 0, 0)
                ORANGE = (0, 165, 255)
                GREEN = (0, 255, 0)

                def _draw_width_line(img, left_pt, right_pt, color, radius):
                    if not left_pt or not right_pt:
                        return
                    lx, ly = int(left_pt[0]), int(left_pt[1])
                    rx, ry = int(right_pt[0]), int(right_pt[1])
                    if (lx == 0 and ly == 0) or (rx == 0 and ry == 0):
                        return

                    # Draw using the exact same endpoints so dots and line are always aligned.
                    cv2.line(img, (lx, ly), (rx, ry), color, 2)
                    cv2.circle(img, (lx, ly), radius, color, -1)
                    cv2.circle(img, (rx, ry), radius, color, -1)

                # Shoulder width (blue)
                _draw_width_line(
                    vis_image,
                    edge_reference_points.get('shoulder_left'),
                    edge_reference_points.get('shoulder_right'),
                    BLUE,
                    8,
                )

                # Chest width (orange)
                _draw_width_line(
                    vis_image,
                    edge_reference_points.get('chest_left'),
                    edge_reference_points.get('chest_right'),
                    ORANGE,
                    6,
                )

                # Waist width (green)
                _draw_width_line(
                    vis_image,
                    edge_reference_points.get('waist_left'),
                    edge_reference_points.get('waist_right'),
                    GREEN,
                    6,
                )
        except Exception as draw_err:
            print(f"Warning: failed to draw edge-based width overlays: {draw_err}")
        vis_base64 = encode_image(vis_image)
        
        # Encode mask (it's already 2D, but encode_image expects 3D or handles grayscale differently)
        # Convert grayscale mask to 3-channel
        mask_3d = cv2.cvtColor(mask, cv2.COLOR_GRAY2BGR) if len(mask.shape) == 2 else mask
        mask_base64 = encode_image(mask_3d)
        
        # Add hybrid approach metadata
        source_summary = {
            'segmentation_edge': len([m for m in mp_measurements.values() if 'Edge' in str(m.get('source', ''))]),
            'mediapipe_landmarks': len([m for m in formatted_measurements.values() if m.get('source') == 'MediaPipe']),
            'smpl_model': len([m for m in formatted_measurements.values() if m.get('source') == 'SMPL 3D Model']),
            'smpl_mediapipe': len([m for m in formatted_measurements.values() if m.get('source') == 'SMPL + MediaPipe']),
            'estimated': len([m for m in formatted_measurements.values() if m.get('source') == 'Estimated']),
        }
        
        return {
            'success': True,
            'measurements': formatted_measurements,
            'landmark_count': len(landmarks),
            'visualization': vis_base64,
            'mask': mask_base64,
            'bbox': bbox if bbox else None,
            'smpl_target_measurements': target_measurements,
            'smpl': {
                'enabled': True,
                'success': smpl_success,
                'status': 'active' if smpl_success else 'estimated',
                'source': smpl_source_label,
                'gender': detected_gender or 'neutral',
                'error': smpl_error,
                'fit_status': smpl_fit_info.get('fit_status', 'fitted' if smpl_success else 'estimated'),
                'status_text': smpl_fit_info.get('status_text', '✓ Model fitted to your body' if smpl_success else 'Estimated from default body'),
                'fitted_to_user': bool(smpl_fit_info.get('fitted_to_user', smpl_success)),
                'landmarks_source': smpl_fit_info.get('landmarks_source', 'mediapipe'),
                'landmarks_mode': smpl_fit_info.get('landmarks_mode', 'real' if smpl_success else 'unavailable'),
                'landmark_count': int(smpl_fit_info.get('landmark_count', len(landmarks_formatted))),
                'visible_landmark_count': int(smpl_fit_info.get('visible_landmark_count', 0)),
                'pose_applied': bool(smpl_fit_info.get('pose_applied', False)),
            },
            'mesh_data': smpl_mesh_data,
            'hybrid_approach': {
                'enabled': True,
                'edge_points_available': edge_reference_points is not None and edge_reference_points.get('is_valid', False),
                'source_summary': source_summary
            }
        }
        
    except Exception as e:
        print(f"✗ Processing error: {e}")
        traceback.print_exc()
        return {
            'error': str(e),
            'step': 'processing'
        }


def decode_image(base64_str):
    """Decode base64 image to numpy array"""
    if not base64_str:
        return None
    
    try:
        # Remove data URL prefix if present
        if ',' in base64_str:
            base64_str = base64_str.split(',')[1]
        
        img_data = base64.b64decode(base64_str)
        img = Image.open(io.BytesIO(img_data))
        img_array = np.array(img)
        
        # Convert RGB to BGR for OpenCV
        if len(img_array.shape) == 3:
            if img_array.shape[2] == 3:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
            elif img_array.shape[2] == 4:
                img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2BGR)
        
        return img_array
    except Exception as e:
        print(f"Image decode error: {e}")
        return None


def encode_image(img_array):
    """Encode numpy array to base64"""
    try:
        # Convert BGR to RGB
        if len(img_array.shape) == 3 and img_array.shape[2] == 3:
            img_array = cv2.cvtColor(img_array, cv2.COLOR_BGR2RGB)
        
        img = Image.fromarray(img_array)
        buffer = io.BytesIO()
        img.save(buffer, format='PNG')
        img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        return f"data:image/png;base64,{img_base64}"
    except Exception as e:
        print(f"Image encode error: {e}")
        return None


def merge_manual_measurements(front_results, side_results):
    """
    Merge front view (width) and side view (depth) manual measurements into 
    a single consolidated measurements dictionary.
    
    Special handling:
    - leg_length: taken from front view
    - arm_length: taken from side view
    - Both front and side visualizations/masks are preserved and returned
    
    Args:
        front_results: Dictionary with front view measurements
        side_results: Dictionary with side view measurements
        
    Returns:
        Dictionary with merged measurements, both visualizations, and metadata
    """
    merged = {
        'success': True,
        'measurements': {},
        'front_visualization': None,
        'front_mask': None,
        'side_visualization': None,
        'side_mask': None,
        'visualization': None,  # Keep for backward compatibility
        'mask': None,  # Keep for backward compatibility
        'scale_factor': 0,
        'height_px': 0,
        'total_landmarks': 0
    }
    
    # Extract front measurements (widths + leg_length)
    if front_results and front_results.get('success'):
        front_measurements = front_results.get('measurements', {})
        
        # Add front measurements
        # leg_length MUST come from front view
        for name, data in front_measurements.items():
            if name == 'leg_length':
                # Prioritize front view for leg_length
                merged['measurements'][name] = data
                print(f"  ✓ Using leg_length from FRONT view: {data.get('value_cm')} cm")
            elif name != 'arm_length':
                # Add all other measurements except arm_length (which should come from side)
                merged['measurements'][name] = data
        
        # Store front visualization and metadata
        merged['front_visualization'] = front_results.get('visualization')
        merged['front_mask'] = front_results.get('mask')
        merged['visualization'] = front_results.get('visualization')  # Primary visualization
        merged['mask'] = front_results.get('mask')  # Primary mask
        merged['scale_factor'] = front_results.get('scale_factor', 0)
        merged['height_px'] = front_results.get('height_px', 0)
        merged['total_landmarks'] = front_results.get('total_landmarks', 0)
    
    # Extract side measurements (depths + arm_length)
    if side_results and side_results.get('success'):
        side_measurements = side_results.get('measurements', {})
        
        # Add side measurements
        # arm_length MUST come from side view
        for name, data in side_measurements.items():
            if name == 'arm_length':
                # Prioritize side view for arm_length
                merged['measurements'][name] = data
                print(f"  ✓ Using arm_length from SIDE view: {data.get('value_cm')} cm")
            elif name == 'leg_length':
                # Skip leg_length from side (already handled from front)
                print(f"  ⓘ Skipping leg_length from side view (using front view instead)")
                continue
            elif name not in merged['measurements']:
                # Add other measurements if not already present
                merged['measurements'][name] = data
            else:
                # If conflict, keep both with prefixes
                merged['measurements'][f'front_{name}'] = merged['measurements'][name]
                merged['measurements'][f'side_{name}'] = data
                del merged['measurements'][name]
        
        # Store side visualization and metadata
        merged['side_visualization'] = side_results.get('visualization')
        merged['side_mask'] = side_results.get('mask')
        
        # Update total landmarks count
        merged['total_landmarks'] += side_results.get('total_landmarks', 0)
    
    print(f"\n✓ Merged manual measurements: {len(merged['measurements'])} total measurements")
    print(f"  Measurements: {list(merged['measurements'].keys())}")
    if 'arm_length' in merged['measurements']:
        print(f"  ✓ arm_length: {merged['measurements']['arm_length'].get('value_cm')} cm (from SIDE view)")
    if 'leg_length' in merged['measurements']:
        print(f"  ✓ leg_length: {merged['measurements']['leg_length'].get('value_cm')} cm (from FRONT view)")
    
    return merged


@app.route('/api/process-manual', methods=['POST'])
def process_manual_landmarks():
    """
    Process manually marked landmarks and compute measurements.
    Uses the same pixel-to-scale conversion logic as automatic detection.
    """
    try:
        data = request.json
        
        print("\n" + "="*60)
        print("MANUAL LANDMARK PROCESSING")
        print("="*60)
        
        # Get user height for scaling safely
        try:
            user_height_cm = float(data.get('user_height') or 0)
        except (TypeError, ValueError):
            user_height_cm = 0
            
        if user_height_cm <= 0:
            return jsonify({'error': 'User height is required'}), 400
        
        print(f"✓ User height: {user_height_cm} cm")
        
        # Decode images if provided
        front_img = decode_image(data.get('front_image'))
        side_img = decode_image(data.get('side_image'))
        
        if front_img is not None:
            print(f"✓ Front image received: {front_img.shape}")
        
        # Process front view manual landmarks
        front_landmarks = data.get('front_landmarks', {})
        front_results = None
        side_results = None
        
        if front_landmarks:
            print("\n" + "="*60)
            print("PROCESSING FRONT VIEW MANUAL LANDMARKS")
            print("="*60)
            
            front_results = process_manual_view(
                front_landmarks,
                user_height_cm,
                'front',
                front_img
            )
        
        # Process side view if provided
        side_landmarks = data.get('side_landmarks', {})
        if side_landmarks:
            print("\n" + "="*60)
            print("PROCESSING SIDE VIEW MANUAL LANDMARKS")
            print("="*60)
            
            side_img_to_use = side_img if side_img is not None else front_img # Fallback? No, just None
            
            side_results = process_manual_view(
                side_landmarks,
                user_height_cm,
                'side',
                side_img
            )
        
        # MERGE FRONT AND SIDE MEASUREMENTS
        # Instead of returning separate front/side results,
        # merge them into a single consolidated result
        print("\n" + "="*60)
        print("MERGING MANUAL MEASUREMENTS")
        print("="*60)
        
        merged_result = merge_manual_measurements(front_results, side_results)
        
        # Prepare calibration data from merged result
        calibration_data = {
            'user_height_cm': user_height_cm,
            'method': 'manual_landmark_marking',
            'height_in_image_px': merged_result.get('height_px', 0),
            'scale_factor': merged_result.get('scale_factor', 0),
            'formula': f'{user_height_cm} cm / {merged_result.get("height_px", 1):.2f} px = {merged_result.get("scale_factor", 0):.4f} cm/px' if merged_result.get('height_px', 0) > 0 else 'N/A',
            'description': f'1 pixel = {merged_result.get("scale_factor", 0):.4f} cm' if merged_result.get('scale_factor', 0) > 0 else 'Manual mode'
        }

        response = {
            'success': True,
            'mode': 'manual',
            'calibration': calibration_data,
            'results': {
                'merged': merged_result  # Single consolidated result instead of separate front/side
            }
        }
        
        print("✓ Manual processing complete!")
        print("="*60 + "\n")
        
        return jsonify(response)
        
    except Exception as e:
        print(f"\n✗ ERROR: {str(e)}")
        traceback.print_exc()
        return jsonify({
            'error': f'Manual processing failed: {str(e)}'
        }), 500


def snap_point_to_edge(point, image, mask=None, search_radius=20, sample_count=8):
    """
    Snap a user-marked point to the nearest body contour edge.
    
    Uses Canny edge detection on the segmentation mask to find precise body boundaries,
    then searches for the closest edge point within a radius.
    
    Args:
        point: Tuple (x, y) of user-marked point
        image: Original image (numpy array)
        mask: Segmentation mask (optional, improves accuracy)
        search_radius: Pixel radius to search for edges
        sample_count: Number of radial samples to average
        
    Returns:
        Refined (x, y) point snapped to edge, or original point if no edge found
    """
    x, y = int(point[0]), int(point[1])
    h, w = image.shape[:2]
    
    # Validate point is within image bounds
    if x < 0 or x >= w or y < 0 or y >= h:
        return point
    
    try:
        # Use mask if available, otherwise use image
        if mask is not None and mask.size > 0:
            # Apply Canny edge detection on mask
            edges = cv2.Canny(mask, 50, 150)
        else:
            # Convert to grayscale and apply edge detection
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
            edges = cv2.Canny(gray, 50, 150)
        
        # Search in a circular region around the point
        y_min = max(0, y - search_radius)
        y_max = min(h, y + search_radius)
        x_min = max(0, x - search_radius)
        x_max = min(w, x + search_radius)
        
        # Extract region of interest
        roi = edges[y_min:y_max, x_min:x_max]
        
        if roi.size == 0:
            return point
        
        # Find edge pixels in ROI
        edge_points = np.where(roi > 0)
        
        if len(edge_points[0]) == 0:
            # No edges found, return original point
            return point
        
        # Convert to absolute coordinates
        edge_y = edge_points[0] + y_min
        edge_x = edge_points[1] + x_min
        
        # Find closest edge point
        distances = np.sqrt((edge_x - x)**2 + (edge_y - y)**2)
        min_idx = np.argmin(distances)
        
        snapped_x = edge_x[min_idx]
        snapped_y = edge_y[min_idx]
        
        # Average multiple nearby edge points for stability
        close_threshold = 3  # pixels
        close_mask = distances < close_threshold
        if np.sum(close_mask) >= 3:
            snapped_x = int(np.mean(edge_x[close_mask]))
            snapped_y = int(np.mean(edge_y[close_mask]))
        
        print(f"    Edge snap: ({x}, {y}) → ({snapped_x}, {snapped_y}), distance: {distances[min_idx]:.1f}px")
        return (float(snapped_x), float(snapped_y))
        
    except Exception as e:
        print(f"    ⚠ Edge snapping failed: {e}, using original point")
        return point


def refine_measurement_with_contours(p1, p2, image, mask=None, num_samples=5):
    """
    Refine a measurement by sampling along the line and averaging edge-snapped points.
    
    This improves accuracy by:
    1. Taking multiple samples along the measurement line
    2. Snapping each sample to nearest edge
    3. Averaging to reduce noise and improve robustness
    
    Args:
        p1: Start point (x, y)
        p2: End point (x, y)
        image: Original image
        mask: Segmentation mask (optional)
        num_samples: Number of samples to take along the line
        
    Returns:
        Refined (p1, p2) tuple
    """
    try:
        # Sample points along the line
        t_values = np.linspace(0, 1, num_samples)
        
        # For start point, sample nearby points
        start_samples = []
        for t in t_values[:2]:  # Sample near start
            sample_x = p1[0] + t * 0.1 * (p2[0] - p1[0])
            sample_y = p1[1] + t * 0.1 * (p2[1] - p1[1])
            snapped = snap_point_to_edge((sample_x, sample_y), image, mask, search_radius=15)
            start_samples.append(snapped)
        
        # Average start samples
        if start_samples:
            refined_p1 = (
                np.mean([s[0] for s in start_samples]),
                np.mean([s[1] for s in start_samples])
            )
        else:
            refined_p1 = snap_point_to_edge(p1, image, mask)
        
        # For end point, sample nearby points
        end_samples = []
        for t in t_values[-2:]:  # Sample near end
            sample_x = p1[0] + (0.9 + t * 0.1) * (p2[0] - p1[0])
            sample_y = p1[1] + (0.9 + t * 0.1) * (p2[1] - p1[1])
            snapped = snap_point_to_edge((sample_x, sample_y), image, mask, search_radius=15)
            end_samples.append(snapped)
        
        # Average end samples
        if end_samples:
            refined_p2 = (
                np.mean([s[0] for s in end_samples]),
                np.mean([s[1] for s in end_samples])
            )
        else:
            refined_p2 = snap_point_to_edge(p2, image, mask)
        
        return refined_p1, refined_p2
        
    except Exception as e:
        print(f"    ⚠ Contour refinement failed: {e}, using original points")
        return p1, p2


def process_manual_view(landmarks_data, user_height_cm, view_name, image=None):
    """
    Process manually marked landmarks for a single view.
    
    Args:
        landmarks_data: Dictionary with 'landmarks', 'imageWidth', 'imageHeight'
        user_height_cm: User's actual height in cm
        view_name: 'front' or 'side'
        image: Original image (numpy array) for visualization and automatic scale calibration
    
    Returns:
        Dictionary with measurements
    """
    try:
        landmarks = landmarks_data.get('landmarks', [])
        image_width = landmarks_data.get('imageWidth', 1)
        image_height = landmarks_data.get('imageHeight', 1)
        
        if not landmarks and not image:
             return {
                'success': False,
                'error': 'No landmarks provided',
                'measurements': {}
            }
        
        print(f"✓ Processing {len(landmarks)} manual landmarks")
        
        # 1. Determine Scale Factor
        scale_factor = 0
        height_px = 0
        
        # Method A: Try Automatic Detection first (Most Accurate)
        if image is not None:
            print("  Attempting automatic detection for scale calibration...")
            # Detect landmarks to measure height in pixels
            ld = get_landmark_detector()
            if ld is not None:
                auto_landmarks = ld.detect(image)
            else:
                auto_landmarks = None
            
            if auto_landmarks is not None:
                nose = auto_landmarks[0]
                left_ankle = auto_landmarks[27]
                right_ankle = auto_landmarks[28]
                ankle_y = max(left_ankle[1], right_ankle[1])
                height_px = ankle_y - nose[1]
                
                if height_px > 0:
                    scale_factor = user_height_cm / height_px
                    print(f"  ✓ Auto-calibration successful: {height_px:.1f} px height -> {scale_factor:.4f} cm/px")
                else:
                    print("  ⚠ Auto-detection found landmarks but height calculation failed.")
            else:
                print("  ⚠ Automatic detection failed (no person found).")
        
        # Method B: Estimate from manual landmarks if Auto failed
        if scale_factor <= 0:
            print("  Falling back to manual height estimation...")
            height_px = estimate_height_from_landmarks(landmarks, image_height)
            
            if height_px <= 0:
                 # Fallback: use image height as approximation
                height_px = image_height * 0.85 
                print(f"  ⚠ Could not estimate height from landmarks. Using 85% image height: {height_px:.1f} px")
            
            scale_factor = user_height_cm / height_px
            print(f"  ✓ Manual/Fallback calibration: {height_px:.1f} px height -> {scale_factor:.4f} cm/px")

        
        measurements = {}
        
        # 2. Generate segmentation mask for edge refinement
        mask = None
        if image is not None:
            try:
                print("  Generating segmentation mask for edge refinement...")
                mask = segmentation_model.segment_person(image, conf_threshold=0.3)
                if mask is not None:
                    print(f"  ✓ Mask generated: {mask.shape}")
                else:
                    print("  ⚠ Mask generation failed, will use image edges directly")
            except Exception as e:
                print(f"  ⚠ Could not generate mask: {e}")
                mask = None
        
        # 3. Process each manual landmark line with edge snapping and refinement
        # Create a visualization image
        vis_image = image.copy() if image is not None else np.zeros((image_height, image_width, 3), dtype=np.uint8)
        
        for landmark in landmarks:
            landmark_type = landmark.get('type', 'custom')
            landmark_label = landmark.get('label', 'Unknown')
            points = landmark.get('points', [])
            
            if len(points) == 2:
                # Get original user-marked points
                p1 = points[0]
                p2 = points[1]
                
                x1_orig, y1_orig = p1['x'], p1['y']
                x2_orig, y2_orig = p2['x'], p2['y']
                
                print(f"\n  Processing: {landmark_label}")
                print(f"    Original points: ({x1_orig:.1f}, {y1_orig:.1f}) → ({x2_orig:.1f}, {y2_orig:.1f})")
                
                # ACCURACY IMPROVEMENT: Snap points to body contour edges
                if image is not None:
                    # Refine with contour detection and multi-sample averaging
                    (x1, y1), (x2, y2) = refine_measurement_with_contours(
                        (x1_orig, y1_orig), 
                        (x2_orig, y2_orig),
                        image, 
                        mask,
                        num_samples=5
                    )
                    print(f"    Refined points: ({x1:.1f}, {y1:.1f}) → ({x2:.1f}, {y2:.1f})")
                else:
                    x1, y1, x2, y2 = x1_orig, y1_orig, x2_orig, y2_orig
                
                # Calculate pixel distance with refined points
                dx = x2 - x1
                dy = y2 - y1
                pixel_dist = np.sqrt(dx**2 + dy**2)
                
                # Calculate original distance for comparison
                dx_orig = x2_orig - x1_orig
                dy_orig = y2_orig - y1_orig
                pixel_dist_orig = np.sqrt(dx_orig**2 + dy_orig**2)
                
                improvement = abs(pixel_dist - pixel_dist_orig)
                if improvement > 0.5:
                    print(f"    ✓ Accuracy improvement: {improvement:.1f} px ({improvement*scale_factor:.2f} cm)")
                
                # Convert to cm using scale factor (UNCHANGED - using correct formula)
                cm_dist = pixel_dist * scale_factor
                
                measurements[landmark_type] = {
                    'value_cm': round(float(cm_dist or 0), 2),
                    'value_px': round(float(pixel_dist or 0), 2),
                    'confidence': 0.95,  # High confidence due to edge snapping
                    'source': 'Manual (Edge-Refined)',
                    'label': landmark_label,
                    'formula': f"{pixel_dist:.2f} px × {scale_factor:.4f} cm/px = {cm_dist:.2f} cm",
                    'refinement': {
                        'original_px': round(float(pixel_dist_orig or 0), 2),
                        'refined_px': round(float(pixel_dist or 0), 2),
                        'improvement_px': round(float(improvement or 0), 2),
                        'edge_snapped': True
                    }
                }

                # Draw on visualization with both original and refined points
                # Draw original points in gray (faded)
                pt1_orig = (int(x1_orig), int(y1_orig))
                pt2_orig = (int(x2_orig), int(y2_orig))
                cv2.line(vis_image, pt1_orig, pt2_orig, (128, 128, 128), 1) # Gray line (original)
                cv2.circle(vis_image, pt1_orig, 4, (128, 128, 128), 1) # Gray dots (original)
                cv2.circle(vis_image, pt2_orig, 4, (128, 128, 128), 1)
                
                # Draw refined points in bright colors
                pt1 = (int(x1), int(y1))
                pt2 = (int(x2), int(y2))
                cv2.line(vis_image, pt1, pt2, (0, 255, 255), 3) # Yellow line (refined)
                cv2.circle(vis_image, pt1, 6, (0, 255, 0), -1) # Green dots (refined)
                cv2.circle(vis_image, pt2, 6, (0, 255, 0), -1)
                
                # Draw label
                mid_x = int((x1 + x2) / 2)
                mid_y = int((y1 + y2) / 2)
                cv2.putText(vis_image, f"{cm_dist:.1f}cm", (mid_x, mid_y - 10), 
                           cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255, 255, 255), 2)
                
                print(f"  ✓ {landmark_label}: {cm_dist:.2f} cm ({pixel_dist:.1f} px) [Edge-Refined]")

        # 4. Generate outputs
        vis_base64 = encode_image(vis_image)
        
        mask_base64 = None
        if image is not None:
             # Try to generate a mask for aesthetics
             try:
                mask = segmentation_model.segment_person(image, conf_threshold=0.3)
                if mask is not None:
                    mask_base64 = encode_image(cv2.cvtColor(mask, cv2.COLOR_GRAY2BGR))
             except Exception as e:
                 print(f"  ⚠ Could not generate mask: {e}")

        return {
            'success': True,
            'measurements': measurements,
            'scale_factor': float(scale_factor or 0),
            'height_px': float(height_px or 0),
            'visualization': vis_base64,
            'mask': mask_base64,
            'total_landmarks': len(landmarks)
        }
        
    except Exception as e:
        print(f"✗ Manual view processing error: {e}")
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e),
            'measurements': {}
        }


def estimate_height_from_landmarks(landmarks, image_height):
    """
    Estimate body height in pixels from landmark positions.
    Looks for vertical measurements like shoulder-to-ankle or hip-to-ankle.
    """
    height_estimates = []
    
    for landmark in landmarks:
        points = landmark.get('points', [])
        if len(points) == 2:
            # Check if measurement is primarily vertical
            p1 = points[0]
            p2 = points[1]
            dy = abs(p2['y'] - p1['y'])
            dx = abs(p2['x'] - p1['x'])
            
            # If vertical component is much larger than horizontal
            if dy > dx * 2:
                height_estimates.append(dy)
    
    if height_estimates:
        # Use the largest vertical measurement as height estimate
        return max(height_estimates)
    
    return 0


# --- EXPORT ENDPOINTS ---

    return 0


# --- EXPORT HELPERS ---

def export_pdf(measurements_data, user_id, output_path):
    """
    Create professional PDF report with FitLens branding.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
    except ImportError:
        return None

    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    logo_style = ParagraphStyle(
        'LogoStyle',
        parent=styles['Title'],
        fontSize=24,
        textColor=colors.HexColor('#1a73e8'),
        spaceAfter=0
    )
    
    footer_style = ParagraphStyle(
        'FooterStyle',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey,
        alignment=1 # Center
    )
    
    elements = []
    
    # Header: FitLens Logo Text
    elements.append(Paragraph("FitLens", logo_style))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1a73e8'), spaceAfter=12))
    
    # Title
    elements.append(Paragraph("Body Measurement Report", styles['Title']))
    
    # Metadata
    elements.append(Paragraph(f"<b>User ID:</b> {user_id}", styles['Normal']))
    elements.append(Paragraph(f"<b>Date Generated:</b> {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    elements.append(Spacer(1, 12))
    
    # Table Header
    table_data = [["Measurement", "Value (cm)", "Value (px)", "Source"]]
    
    # Add measurements from all views
    results = measurements_data.get('results', {})
    
    def process_measurements(m_dict):
        for name, data in m_dict.items():
            if isinstance(data, dict):
                table_data.append([
                    name.replace('_', ' ').title(),
                    f"{data.get('value_cm', 0):.2f}",
                    f"{data.get('value_px', 0):.2f}",
                    data.get('source', 'N/A')
                ])

    for view_name, view_data in results.items():
        if isinstance(view_data, dict) and 'measurements' in view_data:
            process_measurements(view_data['measurements'])
            
    if 'merged' in results and 'measurements' in results['merged']:
        process_measurements(results['merged']['measurements'])

    # Create Table
    t = Table(table_data, colWidths=[2.5*inch, 1.2*inch, 1.2*inch, 1.5*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a73e8')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
    ]))
    elements.append(t)
    
    # Footer
    elements.append(Spacer(1, 0.5*inch))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    elements.append(Paragraph("Generated by FitLens - Your Personal Body Measurement Assistant", footer_style))
    
    doc.build(elements)
    return output_path

def export_docx(measurements_data, user_id, output_path):
    """
    Create Word document with FitLens branding.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()
    
    # Title
    title = doc.add_heading('FitLens Body Measurement Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    p = doc.add_paragraph()
    p.add_run(f'User ID: ').bold = True
    p.add_run(f'{user_id}\n')
    p.add_run(f'Timestamp: ').bold = True
    p.add_run(f'{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    headers = ['Measurement', 'Value (cm)', 'Value (px)', 'Source']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        
    results = measurements_data.get('results', {})
    
    def add_rows(m_dict):
        for name, data in m_dict.items():
            if isinstance(data, dict):
                row_cells = table.add_row().cells
                row_cells[0].text = name.replace('_', ' ').title()
                row_cells[1].text = f"{data.get('value_cm', 0):.2f}"
                row_cells[2].text = f"{data.get('value_px', 0):.2f}"
                row_cells[3].text = data.get('source', 'N/A')

    for view_name, view_data in results.items():
        if isinstance(view_data, dict) and 'measurements' in view_data:
            add_rows(view_data['measurements'])
            
    if 'merged' in results and 'measurements' in results['merged']:
        add_rows(results['merged']['measurements'])

    doc.save(output_path)
    return output_path


# --- EXPORT ENDPOINTS ---

@app.route('/api/download/pdf', methods=['POST'])
def download_pdf():
    """Generate and download a PDF report of measurements."""
    try:
        data = request.json
        results = data.get('results', {})
        user_id = data.get('user_id', 'Guest_User')
        
        # Create a temporary file
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf_path = temp_pdf.name
        temp_pdf.close()
        
        result_path = export_pdf(data, user_id, temp_pdf_path)
        
        if not result_path:
             return jsonify({'error': 'PDF generation libraries (reportlab) are not installed on the server.'}), 501
             
        return send_file(result_path, as_attachment=True, download_name=f"FitLens_Report_{user_id}.pdf")
    except Exception as e:
        print(f"PDF Export Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/docx', methods=['POST'])
def download_docx():
    """Generate and download a Word report of measurements."""
    try:
        data = request.json
        results = data.get('results', {})
        user_id = data.get('user_id', 'Guest_User')
        
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_docx_path = temp_docx.name
        temp_docx.close()
        
        result_path = export_docx(data, user_id, temp_docx_path)
        
        if not result_path:
            return jsonify({'error': 'DOCX generation libraries (python-docx) are not installed on the server.'}), 501
            
        return send_file(result_path, as_attachment=True, download_name=f"FitLens_Report_{user_id}.docx")
    except Exception as e:
        print(f"DOCX Export Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/xml', methods=['POST'])
def download_xml():
    """Generate and download an XML report of measurements."""
    try:
        data = request.json
        results = data.get('results', {})
        user_id = data.get('user_id', 'Guest_User')
        
        root = ET.Element("FitLensMeasurementReport")
        ET.SubElement(root, "UserID").text = str(user_id)
        ET.SubElement(root, "Date").text = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        measurements_node = ET.SubElement(root, "Measurements")
        
        def add_measurement_to_xml(m_name, m_val):
            m_node = ET.SubElement(measurements_node, "Measurement")
            ET.SubElement(m_node, "Name").text = m_name
            ET.SubElement(m_node, "ValueCM").text = str(m_val.get('value_cm', 0))
            ET.SubElement(m_node, "ValuePX").text = str(m_val.get('value_px', 0))
            ET.SubElement(m_node, "Confidence").text = str(m_val.get('confidence', 0))
            ET.SubElement(m_node, "Source").text = str(m_val.get('source', 'N/A'))

        for view_name, view_data in results.items():
            if isinstance(view_data, dict) and 'measurements' in view_data:
                for m_name, m_val in view_data['measurements'].items():
                    add_measurement_to_xml(m_name, m_val)

        if 'merged' in results:
             merged_measurements = results['merged'].get('measurements', {})
             for m_name, m_val in merged_measurements.items():
                    add_measurement_to_xml(m_name, m_val)

        tree = ET.ElementTree(root)
        temp_xml = tempfile.NamedTemporaryFile(delete=False, suffix='.xml')
        tree.write(temp_xml.name, encoding='utf-8', xml_declaration=True)
        temp_xml_path = temp_xml.name
        temp_xml.close()
        
        return send_file(temp_xml_path, as_attachment=True, download_name=f"FitLens_Report_{user_id}.xml")
    except Exception as e:
        print(f"XML Export Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    print("\n" + "="*60)
    print("BODY MEASUREMENT SYSTEM - BACKEND SERVER")
    print("="*60)
    print("\nWorkflow:")
    print("  1. Upload photos")
    print("  2. Detect reference object")
    print("  3. YOLOv8-seg masking (human body only)")
    print("  4. MediaPipe pose landmarks")
    print("  5. Compute measurements (px → cm)")
    print("  6. Return JSON results")
    print("\n" + "="*60 + "\n")
    
    # Use socketio.run instead of app.run
    socketio.run(app, host='0.0.0.0', port=5001, debug=True, use_reloader=False)
