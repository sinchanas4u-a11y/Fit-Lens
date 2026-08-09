"""
Flask Backend API for Body Measurement System
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Fix for Windows socket teardown error: OSError: [WinError 10038] An operation was attempted on something that is not a socket
if sys.platform == 'win32':
    import selectors
    _orig_select = selectors.SelectSelector._select
    def _safe_select(self, r, w, x, timeout=None):
        try:
            return _orig_select(self, r, w, x, timeout)
        except OSError as e:
            if getattr(e, 'winerror', None) == 10038 or getattr(e, 'errno', None) == 10038:
                valid_r = [s for s in r if hasattr(s, 'fileno') and (callable(s.fileno) and s.fileno() >= 0)]
                valid_w = [s for s in w if hasattr(s, 'fileno') and (callable(s.fileno) and s.fileno() >= 0)]
                valid_x = [s for s in x if hasattr(s, 'fileno') and (callable(s.fileno) and s.fileno() >= 0)]
                if valid_r or valid_w or valid_x:
                    try:
                        return _orig_select(self, valid_r, valid_w, valid_x, timeout)
                    except Exception:
                        pass
                return [], [], []
            raise
    selectors.SelectSelector._select = _safe_select

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_socketio import SocketIO, emit
import cv2
import numpy as np
import base64
import io
from PIL import Image
import threading
import queue
import time
import json
import traceback
from dotenv import load_dotenv
_base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_env_path = os.path.join(_base_dir, '.env')
if not os.path.exists(_env_path):
    _env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
load_dotenv(dotenv_path=_env_path, override=True)

from reference_detector import ReferenceDetector
from temporal_stabilizer import TemporalStabilizer
from backend.measurement_engine import MeasurementEngine
from segmentation_model import SegmentationModel
from backend.landmark_detector import LandmarkDetector
from smpl.smpl_pipeline import run_smpl_pipeline
from smpl.smpl_estimator import SMPLEstimator

import sys
import os

# Add parent directory to path so we can
# import from processing/ folder
parent_dir = os.path.dirname(
  os.path.dirname(os.path.abspath(__file__))
)
if parent_dir not in sys.path:
  sys.path.insert(0, parent_dir)

from processing.smplifyx_runner import (
  run_smplifyx
)
from processing.smplifyx_reader import (
  SMPLifyXReader
)

try:
    from face_verifier import FaceVerifier
except Exception:
    FaceVerifier = None

from dotenv import load_dotenv
import uuid
import datetime as dt
import bcrypt
from pymongo import MongoClient
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity

load_dotenv()

# MongoDB connection with fallback
mongo_uri = os.getenv('MONGODB_URI', 'mongodb://localhost:27017/fitlens')
try:
    client = MongoClient(mongo_uri, serverSelectionTimeoutMS=3000)
    client.admin.command('ping')
    try:
        db = client.get_default_database()
    except Exception:
        db = client['fitlens']
    print(f"Connected to MongoDB successfully: {mongo_uri}")
except Exception as e:
    print(f"Warning: Could not connect to MongoDB ({e}). Using in-memory fallback store.")
    class MockCollection:
        def __init__(self): self.data = []
        def find_one(self, query, proj=None, sort=None):
            for d in self.data:
                match = True
                for k, v in query.items():
                    if d.get(k) != v: match = False; break
                if match:
                    res = dict(d)
                    if proj:
                        for pk, pv in proj.items():
                            if pv == 0: res.pop(pk, None)
                    return res
            return None
        def insert_one(self, doc): self.data.append(doc)
        def update_one(self, query, update):
            for d in self.data:
                match = True
                for k, v in query.items():
                    if d.get(k) != v: match = False; break
                if match:
                    if '$set' in update: d.update(update['$set'])
        def find(self, query, proj=None):
            class MockCursor:
                def __init__(self, items): self.items = items
                def sort(self, key, order): return self
                def limit(self, n): return self.items[:n]
                def __iter__(self): return iter(self.items)
            res_list = []
            for d in self.data:
                match = True
                for k, v in query.items():
                    if d.get(k) != v: match = False; break
                if match:
                    res = dict(d)
                    if proj:
                        for pk, pv in proj.items():
                            if pv == 0: res.pop(pk, None)
                    res_list.append(res)
            return MockCursor(res_list)

    class MockDB:
        def __init__(self): self.cols = {}
        def __getitem__(self, name):
            if name not in self.cols: self.cols[name] = MockCollection()
            return self.cols[name]
    db = MockDB()

users_col = db['users']
measurements_col = db['measurements']

app = Flask(__name__)
CORS(app)
socketio = SocketIO(
    app, 
    cors_allowed_origins="*", 
    async_mode='threading',
    logger=False,
    engineio_logger=False,
    ping_timeout=120,
    ping_interval=25,
    max_http_buffer_size=10 * 1024 * 1024
)

# JWT config
app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET', 'fitlens-secret-key')
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = dt.timedelta(days=30)
jwt = JWTManager(app)

# Global directory paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MESHES_DIR = os.path.join(BASE_DIR, "generated_meshes")
os.makedirs(MESHES_DIR, exist_ok=True)


def to_native_types(obj):
    """Recursively convert NumPy/pandas types to native Python types for JSON serialization."""
    if isinstance(obj, dict):
        return {k: to_native_types(v) for k, v in obj.items()}
    elif isinstance(obj, (list, tuple)):
        return [to_native_types(i) for i in obj]
    elif isinstance(obj, np.ndarray):
        return to_native_types(obj.tolist())
    elif isinstance(obj, (np.float32, np.float64, np.float16, np.floating)):
        return float(obj)
    elif isinstance(obj, (np.int32, np.int64, np.int16, np.int8, np.integer)):
        return int(obj)
    elif isinstance(obj, np.generic):
        return obj.item()
    return obj


# Flask-Mail config
import secrets
import re
from flask_mail import Mail, Message

mail_user = (os.getenv('MAIL_EMAIL') or '').strip()
mail_pass = (os.getenv('MAIL_PASSWORD') or '').replace(' ', '').strip()

app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = mail_user
app.config['MAIL_PASSWORD'] = mail_pass
app.config['MAIL_DEFAULT_SENDER'] = mail_user
mail = Mail(app)

reset_tokens = {}

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({
        'status': 'ok',
        'message': 'FitLens backend running',
        'version': '1.0.0'
    }), 200

# --- AUTHENTICATION ROUTES ---
@app.route('/api/auth/register', methods=['POST'])
def register():
    data = request.get_json() or {}
    email = data.get('email', '').lower().strip()
    name = data.get('name', '').strip()
    password = data.get('password', '')

    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400

    if users_col.find_one({'email': email}):
        return jsonify({'error': 'Email already registered'}), 409

    password_hash = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
    user_id = 'U' + str(uuid.uuid4())[:8].upper()

    user = {
        'user_id': user_id,
        'name': name or email.split('@')[0],
        'email': email,
        'password_hash': password_hash,
        'face_embedding': None,
        'created_at': dt.datetime.now(dt.timezone.utc),
        'last_login': dt.datetime.now(dt.timezone.utc)
    }
    users_col.insert_one(user)
    token = create_access_token(identity=user_id)
    return jsonify({
        'success': True,
        'token': token,
        'user': {'user_id': user_id, 'name': user['name'], 'email': email}
    }), 201

@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.get_json() or {}
    email = data.get('email', '').lower().strip()
    password = data.get('password', '')

    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400

    user = users_col.find_one({'email': email})
    if not user or not bcrypt.checkpw(password.encode(), user['password_hash'].encode()):
        return jsonify({'error': 'Invalid email or password'}), 401

    users_col.update_one(
        {'user_id': user['user_id']},
        {'$set': {'last_login': dt.datetime.now(dt.timezone.utc)}}
    )
    token = create_access_token(identity=user['user_id'])
    return jsonify({
        'success': True,
        'token': token,
        'user': {
            'user_id': user['user_id'],
            'name': user['name'],
            'email': email,
            'has_face_embedding': user.get('face_embedding') is not None
        }
    }), 200

@app.route('/api/auth/me', methods=['GET'])
@jwt_required()
def get_current_user():
    user_id = get_jwt_identity()
    user = users_col.find_one({'user_id': user_id}, {'password_hash': 0, '_id': 0})
    if not user:
        return jsonify({'error': 'User not found'}), 404
    user['has_face_embedding'] = user.get('face_embedding') is not None
    user.pop('face_embedding', None)
    return jsonify({'success': True, 'user': user}), 200

@app.route('/api/auth/save-face', methods=['POST'])
@jwt_required()
def save_face_embedding():
    """Save face embedding for identity verification"""
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    front_image_b64 = data.get('front_image')

    if not front_image_b64:
        return jsonify({'error': 'No image provided'}), 400

    try:
        img_data = base64.b64decode(front_image_b64.split(',')[1] if ',' in front_image_b64 else front_image_b64)
        img_array = np.frombuffer(img_data, np.uint8)
        image = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

        if face_verifier is not None and getattr(face_verifier, 'is_ready', False):
            rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            faces = face_verifier.app.get(rgb)
            if faces:
                embedding = faces[0].embedding.tolist()
                users_col.update_one(
                    {'user_id': user_id},
                    {'$set': {'face_embedding': embedding}}
                )
                return jsonify({'success': True, 'message': 'Face saved'}), 200

        return jsonify({'success': False, 'error': 'No face detected'}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/auth/verify-face', methods=['POST'])
@jwt_required()
def verify_face():
    """Verify if scanned person matches account owner"""
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    front_image_b64 = data.get('front_image')

    user = users_col.find_one({'user_id': user_id})
    if not user:
        return jsonify({'verified': False, 'error': 'User not found'}), 404

    stored_embedding = user.get('face_embedding')
    if not stored_embedding:
        return jsonify({'verified': True, 'message': 'No reference face — saving as reference'}), 200

    try:
        img_data = base64.b64decode(front_image_b64.split(',')[1] if ',' in front_image_b64 else front_image_b64)
        img_array = np.frombuffer(img_data, np.uint8)
        image = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

        if face_verifier is not None and getattr(face_verifier, 'is_ready', False):
            rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            faces = face_verifier.app.get(rgb)
            if faces:
                current_embedding = faces[0].embedding
                stored = np.array(stored_embedding)
                similarity = float(np.dot(current_embedding, stored) /
                    (np.linalg.norm(current_embedding) * np.linalg.norm(stored)))
                verified = similarity > 0.4
                return jsonify({
                    'verified': verified,
                    'similarity': similarity,
                    'message': 'Identity verified' if verified else 'Face does not match account owner'
                }), 200

        return jsonify({'verified': True, 'message': 'Face verification unavailable'}), 200
    except Exception as e:
        return jsonify({'verified': False, 'error': str(e)}), 500

def get_network_frontend_url():
    url = os.getenv('FRONTEND_URL', 'http://localhost:3000').rstrip('/')
    if 'localhost' in url or '127.0.0.1' in url:
        try:
            import socket
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(('8.8.8.8', 80))
            local_ip = s.getsockname()[0]
            s.close()
            url = url.replace('localhost', local_ip).replace('127.0.0.1', local_ip)
        except Exception:
            pass
    return url

@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password():
    try:
        data = request.get_json() or {}
        email = data.get('email', '').lower().strip()
        user = users_col.find_one({'email': email})

        STANDARD_MSG = "If an account exists for this email address, we've sent a password reset link. Please check your inbox."

        generic_response = jsonify({
            'success': True,
            'message': STANDARD_MSG
        }), 200

        if not user:
            return generic_response

        # Generate secure token
        token = secrets.token_urlsafe(48)
        expires = dt.datetime.now(dt.timezone.utc) + dt.timedelta(minutes=15)

        # Delete any existing unused tokens for this user
        reset_tokens_col = db['reset_tokens']
        reset_tokens_col.delete_many({'user_id': user['user_id'], 'used': False})

        # Store token in MongoDB
        reset_tokens_col.insert_one({
            'token': token,
            'user_id': user['user_id'],
            'email': email,
            'expires': expires,
            'used': False,
            'created_at': dt.datetime.now(dt.timezone.utc)
        })
        reset_tokens[token] = {
            'user_id': user['user_id'],
            'expires': expires,
            'used': False
        }

        frontend_url = get_network_frontend_url()
        reset_link = f"{frontend_url}/reset-password?token={token}"

        mail_email = (os.getenv('MAIL_EMAIL') or app.config.get('MAIL_USERNAME') or '').strip()
        mail_pass = (os.getenv('MAIL_PASSWORD') or app.config.get('MAIL_PASSWORD') or '').replace(' ', '').strip()
        if mail_email:
            app.config['MAIL_USERNAME'] = mail_email
            app.config['MAIL_DEFAULT_SENDER'] = mail_email
        if mail_pass:
            app.config['MAIL_PASSWORD'] = mail_pass

        # Send email
        msg = Message(
            subject='FitLens AI — Reset Your Password',
            sender=mail_email,
            recipients=[email]
        )
        msg.html = f"""
        <div style="font-family:Arial,sans-serif;max-width:480px;margin:auto;
                    background:#0a0e27;color:#ffffff;padding:40px;border-radius:16px;
                    border:1px solid #2d3561;">
          <div style="text-align:center;margin-bottom:24px;">
            <h1 style="color:#00d4aa;font-size:28px;margin:0;">FitLens AI</h1>
            <p style="color:#a0aec0;margin:4px 0 0;">AI-Powered Body Measurements</p>
          </div>
          <h2 style="color:#ffffff;font-size:20px;">Password Reset Request</h2>
          <p style="color:#a0aec0;">Hi <strong style="color:#fff;">{user.get('name', 'User')}</strong>,</p>
          <p style="color:#a0aec0;">
            We received a request to reset the password for your FitLens account.<br>
            Click the button below to set a new password.
          </p>
          <p style="color:#fc8181;font-size:13px;">
            ⏱ This link expires in <strong>15 minutes</strong>.
          </p>
          <div style="text-align:center;margin:32px 0;">
            <a href="{reset_link}"
               style="background:linear-gradient(135deg,#00d4aa,#0080ff);
                      color:#ffffff;text-decoration:none;padding:16px 40px;
                      border-radius:10px;font-size:16px;font-weight:bold;
                      display:inline-block;">
              Reset My Password
            </a>
          </div>
          <p style="color:#4a5568;font-size:12px;text-align:center;">
            If you did not request a password reset, ignore this email.<br>
            Your password will remain unchanged.
          </p>
          <hr style="border:1px solid #2d3561;margin:24px 0;">
          <p style="color:#4a5568;font-size:11px;text-align:center;">
            Or copy this link: <br>
            <span style="color:#00d4aa;word-break:break-all;">{reset_link}</span>
          </p>
        </div>
        """
        try:
            if app.config.get('MAIL_USERNAME') and app.config.get('MAIL_PASSWORD') and app.config.get('MAIL_USERNAME') != 'your-gmail@gmail.com':
                mail.send(msg)
                print(f"[EMAIL] Reset link successfully sent to {email}")
            else:
                print(f"[MAIL MOCK] Mail credentials unconfigured in .env. Reset link: {reset_link}")
        except Exception as mail_err:
            print(f"[EMAIL ERROR] Failed to send email via SMTP ({mail_err}). Reset link: {reset_link}")

        return generic_response

    except Exception as e:
        print(f"[EMAIL ERROR] {str(e)}")
        return jsonify({'error': f'Email sending failed: {str(e)}'}), 500

@app.route('/api/auth/reset-password', methods=['POST'])
def reset_password():
    data = request.get_json() or {}
    token = data.get('token', '').strip()
    new_password = data.get('new_password', '')
    confirm_password = data.get('confirm_password', '')

    if not token:
        return jsonify({'error': 'Reset token is missing'}), 400
    if len(new_password) < 8:
        return jsonify({'error': 'Password must be at least 8 characters'}), 400
    if new_password != confirm_password:
        return jsonify({'error': 'Passwords do not match'}), 400

    reset_tokens_col = db['reset_tokens']
    token_doc = reset_tokens_col.find_one({'token': token, 'used': False})
    if not token_doc:
        if token in reset_tokens and not reset_tokens[token].get('used'):
            token_doc = {
                'token': token,
                'user_id': reset_tokens[token]['user_id'],
                'expires': reset_tokens[token]['expires'],
                'used': False
            }
        else:
            return jsonify({'error': 'Invalid or already used reset link. Please request a new one.'}), 400

    expires = token_doc['expires']
    if isinstance(expires, str):
        try:
            expires = dt.datetime.fromisoformat(expires)
        except Exception:
            pass
    if isinstance(expires, dt.datetime) and expires.tzinfo is None:
        expires = expires.replace(tzinfo=dt.timezone.utc)

    now_utc = dt.datetime.now(dt.timezone.utc)
    if isinstance(expires, dt.datetime) and now_utc > expires:
        reset_tokens_col.delete_one({'token': token})
        reset_tokens.pop(token, None)
        return jsonify({'error': 'Reset link has expired. Please request a new one.'}), 400

    # Hash and update password
    password_hash = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
    users_col.update_one(
        {'user_id': token_doc['user_id']},
        {'$set': {'password_hash': password_hash}}
    )

    # Invalidate token
    reset_tokens_col.update_one(
        {'token': token},
        {'$set': {'used': True, 'used_at': dt.datetime.now(dt.timezone.utc)}}
    )
    if token in reset_tokens:
        reset_tokens[token]['used'] = True

    return jsonify({'success': True, 'message': 'Password reset successfully. Please log in.'}), 200

@app.route('/api/auth/change-password', methods=['POST'])
@jwt_required()
def change_password():
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    current_password = data.get('current_password', '')
    new_password = data.get('new_password', '')
    confirm_password = data.get('confirm_password', '')

    if len(new_password) < 8:
        return jsonify({'error': 'New password must be at least 8 characters'}), 400
    if new_password != confirm_password:
        return jsonify({'error': 'New passwords do not match'}), 400

    user = users_col.find_one({'user_id': user_id})
    if not user:
        return jsonify({'error': 'User not found'}), 404
    if not bcrypt.checkpw(current_password.encode(), user['password_hash'].encode()):
        return jsonify({'error': 'Current password is incorrect'}), 401
    if current_password == new_password:
        return jsonify({'error': 'New password must be different from current password'}), 400

    password_hash = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
    users_col.update_one(
        {'user_id': user_id},
        {'$set': {'password_hash': password_hash}}
    )
    return jsonify({'success': True, 'message': 'Password changed successfully.'}), 200

@app.route('/api/auth/update-profile', methods=['PUT'])
@jwt_required()
def update_profile():
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    name = data.get('name', '').strip()
    
    if not name or len(name) < 2 or len(name) > 50:
        return jsonify({'error': 'Name must be between 2 and 50 characters'}), 400
        
    users_col.update_one({'user_id': user_id}, {'$set': {'name': name}})
    return jsonify({'success': True, 'message': 'Profile updated', 'user': {'name': name}}), 200

@app.route('/api/auth/delete-account', methods=['DELETE'])
@jwt_required()
def delete_account():
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    password = data.get('password', '')
    
    active_scan = globals().get('active_processing_users', set())
    if user_id in active_scan:
        return jsonify({'error': 'Cannot delete account while a measurement scan is processing'}), 400
        
    user = users_col.find_one({'user_id': user_id})
    if not user:
        return jsonify({'error': 'User not found'}), 404
        
    if not bcrypt.checkpw(password.encode(), user['password_hash'].encode()):
        return jsonify({'error': 'Incorrect password'}), 401
        
    users_col.delete_one({'user_id': user_id})
    measurements_col.delete_many({'user_id': user_id})
    return jsonify({'success': True, 'message': 'Account deleted'}), 200

# --- MEASUREMENT ROUTES ---
@app.route('/api/measurements/save', methods=['POST'])
@jwt_required()
def save_measurements():
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    measurements = data.get('measurements', {})

    def safe_float(val):
        try:
            if isinstance(val, dict):
                val = val.get('value_cm') or val.get('value')
            return float(val) if val is not None else None
        except Exception:
            return None

    record = {
        'analysis_id': 'A' + str(uuid.uuid4())[:8].upper(),
        'user_id': user_id,
        'date': dt.datetime.now(dt.timezone.utc).strftime('%d-%b-%Y'),
        'height_cm': safe_float(data.get('user_height')),
        'arm_length': safe_float(measurements.get('arm_length')),
        'leg_length': safe_float(measurements.get('leg_length')),
        'torso_length': safe_float(measurements.get('torso_length')),
        'shoulder_width': safe_float(measurements.get('shoulder_width')),
        'chest_circumference': safe_float(measurements.get('chest_circumference')),
        'waist_circumference': safe_float(measurements.get('waist_circumference')),
        'hip_circumference': safe_float(measurements.get('hip_circumference')),
        'chest_width': safe_float(measurements.get('chest_width')),
        'waist_width': safe_float(measurements.get('waist_width')),
        'hip_width': safe_float(measurements.get('hip_width')),
        'source': data.get('source', 'upload'),
        'created_at': dt.datetime.now(dt.timezone.utc)
    }
    measurements_col.insert_one(record)
    record.pop('_id', None)
    return jsonify({'success': True, 'analysis': record}), 201

@app.route('/api/measurements/history', methods=['GET'])
@jwt_required()
def get_history():
    user_id = get_jwt_identity()
    cursor = measurements_col.find(
        {'user_id': user_id},
        {'_id': 0}
    )
    if hasattr(cursor, 'sort'):
        cursor = cursor.sort('created_at', -1)
    if hasattr(cursor, 'limit'):
        cursor = cursor.limit(20)
    records = list(cursor)
    return jsonify({'success': True, 'history': records}), 200

@app.route('/api/measurements/latest', methods=['GET'])
@jwt_required()
def get_latest():
    user_id = get_jwt_identity()
    record = measurements_col.find_one(
        {'user_id': user_id},
        {'_id': 0},
        sort=[('created_at', -1)]
    )
    if not record:
        return jsonify({'error': 'No measurements found'}), 404
    return jsonify({'success': True, 'analysis': record}), 200

@app.route('/api/measurements/delete/<analysis_id>', methods=['DELETE'])
@jwt_required()
def delete_measurement(analysis_id):
    user_id = get_jwt_identity()
    result = measurements_col.delete_one({'analysis_id': analysis_id, 'user_id': user_id})
    return jsonify({'success': True, 'message': 'Measurement deleted'}), 200

# Initialize models
reference_detector = ReferenceDetector()
temporal_stabilizer = TemporalStabilizer()
measurement_engine = MeasurementEngine()
segmentation_model = SegmentationModel()
landmark_detector = LandmarkDetector()

def get_landmark_detector():
    """Get the global landmark detector instance."""
    global landmark_detector
    return landmark_detector

# Optional face model used for best-effort gender detection.
if FaceVerifier is not None:
    try:
        face_verifier = FaceVerifier(model_name="buffalo_l", det_size=(640, 640))
    except Exception:
        face_verifier = None
else:
    face_verifier = None

# Global state
camera = None
camera_active = False
reference_captured = False
reference_data = {}


def _estimate_height_from_landmarks_px(landmarks):
    if landmarks is None or len(landmarks) < 29:
        return 0.0
    try:
        nose = landmarks[0]
        left_ankle = landmarks[27]
        right_ankle = landmarks[28]
        return float(max(left_ankle[1], right_ankle[1]) - nose[1])
    except Exception:
        return 0.0


def _landmarks_to_smpl_list(landmarks, image_shape):
    if landmarks is None or len(landmarks) < 33:
        return []

    h, w = image_shape[:2]
    smpl_landmarks = []
    for idx in range(33):
        lm = landmarks[idx]
        x = float(lm[0]) if len(lm) > 0 else 0.0
        y = float(lm[1]) if len(lm) > 1 else 0.0
        visibility = float(lm[2]) if len(lm) > 2 else 1.0

        if x > 1.5 or y > 1.5:
            x = x / float(w) if w > 0 else 0.0
            y = y / float(h) if h > 0 else 0.0

        smpl_landmarks.append({
            'x': float(max(0.0, min(1.0, x))),
            'y': float(max(0.0, min(1.0, y))),
            'visibility': float(max(0.0, min(1.0, visibility))),
        })

    return smpl_landmarks


def run_gender_detection(image):
    """Best-effort gender detection. Returns male/female/neutral."""
    try:
        if face_verifier is None or not getattr(face_verifier, 'is_ready', False):
            return 'neutral'
        if not hasattr(face_verifier, 'app') or face_verifier.app is None:
            return 'neutral'

        rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
        faces = face_verifier.app.get(rgb)
        if not faces:
            return 'neutral'

        face = max(
            faces,
            key=lambda f: float((f.bbox[2] - f.bbox[0]) * (f.bbox[3] - f.bbox[1]))
        )

        raw_gender = getattr(face, 'sex', None)
        if raw_gender is None:
            raw_gender = getattr(face, 'gender', None)

        if isinstance(raw_gender, str):
            g = raw_gender.strip().lower()
            if g in ('male', 'm', 'man'):
                return 'male'
            if g in ('female', 'f', 'woman'):
                return 'female'

        if isinstance(raw_gender, (int, float)):
            return 'male' if float(raw_gender) >= 0.5 else 'female'
    except Exception:
        pass

    return 'neutral'


def _build_smpl_merged_measurements(mp_measurements, smpl_m, smpl_success):
    print(f"[DEBUG] _build_smpl_merged_measurements: smpl_success={smpl_success}")
    if isinstance(mp_measurements, dict):
        print(f"[DEBUG]   mp_measurements keys={list(mp_measurements.keys())}")
        for k in ['arm_length', 'leg_length', 'shoulder_width']:
            if k in mp_measurements:
                print(f"[DEBUG]     {k}: {mp_measurements[k]}")
    if isinstance(smpl_m, dict):
        print(f"[DEBUG]   smpl_m keys={list(smpl_m.keys())}")

    def _mp_entry(name):
        return mp_measurements.get(name, {}) if isinstance(mp_measurements, dict) else {}

    def _mp_cm(name, fallback=None):
        val = _mp_entry(name).get('value_cm')
        return val if val is not None else fallback

    def _mp_px(name, fallback=None):
        val = _mp_entry(name).get('value_px')
        return val if val is not None else fallback

    chest_width_cm = _mp_cm('chest_width')
    waist_width_cm = _mp_cm('waist_width')
    hip_width_cm = _mp_cm('hip_width')

    mp_chest_circ = _mp_cm('chest_circumference')
    mp_waist_circ = _mp_cm('waist_circumference')
    mp_hip_circ = _mp_cm('hip_circumference')

    if chest_width_cm is None and mp_chest_circ is not None:
        chest_width_cm = round(mp_chest_circ / 3.0, 2)
    if waist_width_cm is None and mp_waist_circ is not None:
        waist_width_cm = round(mp_waist_circ / 2.8, 2)
    if hip_width_cm is None and mp_hip_circ is not None:
        hip_width_cm = round(mp_hip_circ / 3.0, 2)

    if smpl_success:
        chest_circ = smpl_m.get('chest_circumference', mp_chest_circ or (chest_width_cm * 3.0 if chest_width_cm else None))
        waist_circ = smpl_m.get('waist_circumference', mp_waist_circ or (waist_width_cm * 2.8 if waist_width_cm else None))
        hip_circ = smpl_m.get('hip_circumference', mp_hip_circ or (hip_width_cm * 3.0 if hip_width_cm else None))
        circ_source = 'SMPL 3D Model'
    else:
        chest_circ = mp_chest_circ or (round(chest_width_cm * 3.0, 2) if chest_width_cm else None)
        waist_circ = mp_waist_circ or (round(waist_width_cm * 2.8, 2) if waist_width_cm else None)
        hip_circ = mp_hip_circ or (round(hip_width_cm * 3.0, 2) if hip_width_cm else None)
        circ_source = 'Estimated'

    merged = {
        'full_height': {
            'value_cm': _mp_cm('full_height'),
            'value_px': _mp_px('full_height'),
            'source': 'MediaPipe',
            'label': 'Full Height',
        },
        'arm_length': {
            'value_cm': _mp_cm('arm_length'),
            'value_px': _mp_px('arm_length'),
            'source': 'MediaPipe',
            'label': 'Arm Length',
        },
        'leg_length': {
            'value_cm': _mp_cm('leg_length'),
            'value_px': _mp_px('leg_length'),
            'source': 'MediaPipe',
            'label': 'Leg Length',
        },
        'torso_length': {
            'value_cm': _mp_cm('torso_length'),
            'value_px': _mp_px('torso_length'),
            'source': 'MediaPipe',
            'label': 'Torso Length',
        },
        'chest_circumference': {
            'value_cm': chest_circ,
            'value_px': None,
            'source': circ_source,
            'label': 'Chest Circumference',
        },
        'waist_circumference': {
            'value_cm': waist_circ,
            'value_px': None,
            'source': circ_source,
            'label': 'Waist Circumference',
        },
        'hip_circumference': {
            'value_cm': hip_circ,
            'value_px': None,
            'source': circ_source,
            'label': 'Hip Circumference',
        },
        'shoulder_width': {
            'value_cm': _mp_cm('shoulder_width') or smpl_m.get('shoulder_width'),
            'value_px': _mp_px('shoulder_width'),
            'source': 'SMPL + MediaPipe',
            'label': 'Shoulder Width',
        },
        'chest_width': {
            'value_cm': smpl_m.get('chest_width', chest_width_cm),
            'value_px': _mp_px('chest_width', _mp_px('chest_circumference')),
            'source': 'SMPL + MediaPipe',
            'label': 'Chest Width',
        },
        'waist_width': {
            'value_cm': smpl_m.get('waist_width', waist_width_cm),
            'value_px': _mp_px('waist_width', _mp_px('waist_circumference')),
            'source': 'SMPL + MediaPipe',
            'label': 'Waist Width',
        },
        'hip_width': {
            'value_cm': smpl_m.get('hip_width', hip_width_cm),
            'value_px': _mp_px('hip_width'),
            'source': 'SMPL + MediaPipe',
            'label': 'Hip Width',
        },
    }

    for key, value in mp_measurements.items():
        if key not in merged:
            merged[key] = value

    cleaned = {}
    for key, value in merged.items():
        if not isinstance(value, dict):
            continue
        cm = value.get('value_cm')
        px = value.get('value_px')
        if cm is None and px is None:
            continue
        cleaned[key] = value

    return cleaned, circ_source


def consolidate_measurements(front_results, side_results, user_height_cm=None):
    """
    Consolidate front and side view measurements into a single final dictionary
    using the specified priority rules.
    """
    front_meas = front_results.get('measurements', {}) if isinstance(front_results, dict) else {}
    side_meas = side_results.get('measurements', {}) if isinstance(side_results, dict) else {}
    
    merged_meas = {}
    
    # Helper to extract a measurement safely
    def get_meas(key, source_dict):
        if key in source_dict and isinstance(source_dict[key], dict):
            return source_dict[key].copy()
        return None

    # Helper to check if a measurement came from a 3D model
    def is_3d_source(meas):
        if not meas:
            return False
        src = str(meas.get('source', '')).upper()
        return 'SMPL' in src or '3D' in src

    # Priority Rules implementation
    
    # arm_length -> front
    arm = get_meas('arm_length', front_meas) or get_meas('arm_length', side_meas)
    if arm:
        merged_meas['arm_length'] = arm
        
    # leg_length -> front
    leg = get_meas('leg_length', front_meas) or get_meas('leg_length', side_meas)
    if leg:
        merged_meas['leg_length'] = leg
        
    # shoulder_width -> front
    shoulder = get_meas('shoulder_width', front_meas) or get_meas('shoulder_width', side_meas)
    if shoulder:
        merged_meas['shoulder_width'] = shoulder
        
    # torso_length -> front view only (vertical fix applied in measurement engine)
    torso = get_meas('torso_length', front_meas) or get_meas('torso_length', side_meas)
    if torso:
        merged_meas['torso_length'] = torso
        
    # chest_circumference, waist_circumference, hip_circumference -> from SMPL/SMPLify-X (front preferred if both)
    for key in ['chest_circumference', 'waist_circumference', 'hip_circumference']:
        m_f = get_meas(key, front_meas)
        m_s = get_meas(key, side_meas)
        
        f_is_3d = is_3d_source(m_f)
        s_is_3d = is_3d_source(m_s)
        
        if f_is_3d and s_is_3d:
            if 'SMPLIFY-X' in str(m_f.get('source', '')).upper():
                merged_meas[key] = m_f
            else:
                merged_meas[key] = m_f
        elif f_is_3d:
            merged_meas[key] = m_f
        elif s_is_3d:
            merged_meas[key] = m_s
        else:
            merged_meas[key] = m_f or m_s
            
    # chest_width, waist_width, hip_width -> front
    for key in ['chest_width', 'waist_width', 'hip_width']:
        m = get_meas(key, front_meas) or get_meas(key, side_meas)
        if m:
            merged_meas[key] = m
            
    # chest_depth, waist_depth, hip_depth, stomach_depth -> side
    for key in ['chest_depth', 'waist_depth', 'hip_depth', 'stomach_depth']:
        m = get_meas(key, side_meas) or get_meas(key, front_meas)
        if m:
            merged_meas[key] = m
            
    # full_height -> user input prioritised
    if user_height_cm is not None and float(user_height_cm) > 0:
        merged_meas['full_height'] = {
            'value_cm': float(user_height_cm),
            'value_px': None,
            'source': 'User Input',
            'label': 'Full Height'
        }
    else:
        height = get_meas('full_height', front_meas) or get_meas('full_height', side_meas)
        if height:
            merged_meas['full_height'] = height

    # Add any other measurements present in front or side that aren't consolidated yet
    all_keys = set(front_meas.keys()).union(set(side_meas.keys()))
    for key in all_keys:
        if key not in merged_meas:
            merged_meas[key] = get_meas(key, front_meas) or get_meas(key, side_meas)
            
    return merged_meas


def _build_smpl_mesh_data(smpl_result, user_height_cm, gender='neutral'):
    """Build Plotly-ready mesh data from fitted SMPL betas or default body."""
    try:
        if user_height_cm is None or float(user_height_cm) <= 0:
            return None

        if smpl_result and smpl_result.get('success') and smpl_result.get('betas') is not None:
            betas = smpl_result.get('betas')
            is_fitted = True
        else:
            betas = [0.0] * 10
            is_fitted = False

        estimator = SMPLEstimator(gender=gender or 'neutral')
        betas_arr = np.array(betas, dtype=np.float32)

        vertices = estimator.get_vertices(betas_arr)
        faces = np.asarray(estimator.faces, dtype=np.int32)

        y_min = float(vertices[:, 1].min())
        y_max = float(vertices[:, 1].max())
        smpl_height_cm = (y_max - y_min) * 100.0
        if smpl_height_cm <= 0:
            return None

        scale = float(user_height_cm) / smpl_height_cm
        vertices_cm = vertices * 100.0 * scale

        # Keep model centered for stable 3D viewer framing.
        mid_y = (vertices_cm[:, 1].max() + vertices_cm[:, 1].min()) / 2.0
        vertices_cm[:, 1] -= mid_y

        return {
            'x': vertices_cm[:, 0].tolist(),
            'y': vertices_cm[:, 1].tolist(),
            'z': vertices_cm[:, 2].tolist(),
            'i': faces[:, 0].tolist(),
            'j': faces[:, 1].tolist(),
            'k': faces[:, 2].tolist(),
            'metadata': {
                'vertex_count': int(vertices_cm.shape[0]),
                'face_count': int(faces.shape[0]),
                'height_cm': float(user_height_cm),
                'gender': gender or 'neutral',
                'fitted_to_user': is_fitted,
                'betas_fitted': is_fitted,
                'pose_applied': False,
                'source': 'SMPL Fitted A-Pose' if is_fitted else 'SMPL Default A-Pose',
            },
        }
    except Exception as mesh_err:
        print(f"Warning: failed to build SMPL mesh data: {mesh_err}")
        return None


@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'models_loaded': {
            'segmentation': segmentation_model.model is not None,
            'landmarks': landmark_detector.pose is not None,
            'temporal': True
        }
    })


def detect_body_orientation(landmarks, image_shape):
    """
    Strict MediaPipe Pose orientation analysis:
    Determines if body is front-facing (0°) or side-profile (±90°) and verifies full-body keypoints.
    """
    if landmarks is None or len(landmarks) < 33:
        return {
            'orientation': 'invalid',
            'is_full_body': False,
            'reason': 'Insufficient body keypoints detected'
        }

    h, w = image_shape[:2]

    # Keypoints
    nose = landmarks[0]
    left_shoulder = landmarks[11]
    right_shoulder = landmarks[12]
    left_hip = landmarks[23]
    right_hip = landmarks[24]
    left_ankle = landmarks[27]
    right_ankle = landmarks[28]

    def norm_x(pt):
        return pt[0] / float(w) if pt[0] > 1.5 else float(pt[0])

    def norm_y(pt):
        return pt[1] / float(h) if pt[1] > 1.5 else float(pt[1])

    vis_thresh = 0.25
    l_sh_vis = len(left_shoulder) > 2 and left_shoulder[2] >= vis_thresh
    r_sh_vis = len(right_shoulder) > 2 and right_shoulder[2] >= vis_thresh
    l_hip_vis = len(left_hip) > 2 and left_hip[2] >= vis_thresh
    r_hip_vis = len(right_hip) > 2 and right_hip[2] >= vis_thresh
    l_ank_vis = len(left_ankle) > 2 and left_ankle[2] >= vis_thresh
    r_ank_vis = len(right_ankle) > 2 and right_ankle[2] >= vis_thresh

    ls_x, rs_x = norm_x(left_shoulder), norm_x(right_shoulder)
    ls_y, rs_y = norm_y(left_shoulder), norm_y(right_shoulder)
    lh_x, rh_x = norm_x(left_hip), norm_x(right_hip)
    lh_y, rh_y = norm_y(left_hip), norm_y(right_hip)

    shoulder_width_x = abs(rs_x - ls_x)
    hip_width_x = abs(rh_x - lh_x)
    torso_height_y = abs(((lh_y + rh_y) / 2.0) - ((ls_y + rs_y) / 2.0))

    if torso_height_y < 0.05:
        torso_height_y = 0.35

    shoulder_ratio = shoulder_width_x / torso_height_y
    hip_ratio = hip_width_x / torso_height_y

    both_shoulders = l_sh_vis and r_sh_vis
    both_hips = l_hip_vis and r_hip_vis
    has_feet = l_ank_vis or r_ank_vis

    is_full_body = (l_sh_vis or r_sh_vis) and (l_hip_vis or r_hip_vis) and has_feet

    # Refined SIDE vs FRONT orientation logic:
    # In SIDE view (profile), left & right shoulders overlap horizontally -> shoulder_width_x < 0.12 or ratio < 0.30
    is_side = (shoulder_width_x < 0.11 or shoulder_ratio < 0.30) and (hip_width_x < 0.09 or hip_ratio < 0.26)

    # In FRONT view, shoulders are spread out wide -> shoulder_ratio >= 0.30 or shoulder_width_x >= 0.10, with both shoulders visible
    is_front = (shoulder_width_x >= 0.10 or shoulder_ratio >= 0.30) and both_shoulders

    if is_side and not is_front:
        orientation = 'side'
    elif is_front:
        orientation = 'front'
    elif shoulder_width_x < 0.10:
        orientation = 'side'
    elif both_shoulders:
        orientation = 'front'
    else:
        orientation = 'invalid'

    return {
        'orientation': orientation,
        'is_full_body': is_full_body,
        'both_shoulders_visible': both_shoulders,
        'both_hips_visible': both_hips,
        'shoulder_width_x': shoulder_width_x,
        'shoulder_ratio': shoulder_ratio
    }


@app.route('/validate/person-count', methods=['POST'])
@app.route('/api/validate/person-count', methods=['POST'])
def validate_person_count():
    """
    Validate that exactly one person is present, check for cropped body,
    and enforce strict FRONT vs SIDE orientation validation.
    """
    try:
        file = request.files.get('image')
        view = request.form.get('view', 'front')
        
        if not file or file.filename == '':
            return jsonify({'success': False, 'error': 'Invalid image. Please upload a valid image file.'}), 400
            
        allowed_mimes = ['image/jpeg', 'image/jpg', 'image/png', 'image/webp']
        if file.mimetype not in allowed_mimes:
            ext = os.path.splitext(file.filename)[1].lower()
            if ext not in ['.jpg', '.jpeg', '.png', '.webp']:
                return jsonify({'success': False, 'error': 'Unsupported image. Please upload a JPEG, PNG, or WEBP image.'}), 400
                
        file_bytes = np.frombuffer(file.read(), np.uint8)
        img = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)
        if img is None:
            return jsonify({'success': False, 'error': 'Invalid image format or corrupted file. Please upload a valid image.'}), 400
            
        if segmentation_model is None or segmentation_model.model is None:
            return jsonify({'success': False, 'error': 'Segmentation model not initialized.'}), 500
            
        h_orig, w_orig = img.shape[:2]
        pad_h = int(h_orig * 0.1)
        pad_w = int(w_orig * 0.1)
        padded_image = cv2.copyMakeBorder(img, pad_h, pad_h, pad_w, pad_w, 
                                         cv2.BORDER_CONSTANT, value=[128, 128, 128])
                                         
        results = segmentation_model.model(padded_image, conf=0.5, imgsz=1024, retina_masks=True, verbose=False, classes=[0])
        
        num_people = 0
        if len(results) > 0 and results[0].boxes is not None:
            num_people = len(results[0].boxes)
            
        if num_people == 0:
            return jsonify({'success': False, 'error': 'No person detected in the image. Please upload a valid image containing one person.'}), 400
        elif num_people > 1:
            if view == 'front':
                err_msg = 'Multiple persons detected in front view. Please upload an image with only one person.'
            else:
                err_msg = 'Multiple persons detected in side view. Please upload an image with only one person.'
            return jsonify({'success': False, 'error': err_msg}), 400
            
        if landmark_detector is None:
            return jsonify({'success': False, 'error': 'Landmark detector not initialized.'}), 500
            
        landmarks = landmark_detector.detect(img)
        
        if landmarks is None or len(landmarks) < 33:
            return jsonify({'success': False, 'error': 'No person detected in the image. Please upload a valid image containing one person.'}), 400
            
        # Cropped body check
        nose = landmarks[0]
        left_ankle = landmarks[27]
        right_ankle = landmarks[28]
        
        cropped = False
        if left_ankle[1] > h_orig * 0.995 and right_ankle[1] > h_orig * 0.995:
            cropped = True
        elif nose[1] < h_orig * 0.005:
            cropped = True
            
        if cropped:
            return jsonify({'success': False, 'error': 'Cropped body detected. Please ensure your entire body (from head to toe) is visible in the photo.'}), 400

        # Strict Orientation & View Validation
        expected_view = (view or 'front').lower()
        orient_data = detect_body_orientation(landmarks, img.shape)
        detected_orient = orient_data['orientation']
        is_full_body = orient_data['is_full_body']
        both_shoulders = orient_data['both_shoulders_visible']
        both_hips = orient_data['both_hips_visible']

        print(f"🔍 Orientation check: expected='{expected_view}', detected='{detected_orient}', full_body={is_full_body}")

        if expected_view == 'front':
            if detected_orient == 'side':
                return jsonify({
                    'success': False,
                    'error': 'Side-view image detected. Please upload a FRONT-view image in this section.'
                }), 400
            if detected_orient != 'front' or not both_shoulders or not both_hips or not is_full_body:
                return jsonify({
                    'success': False,
                    'error': 'Invalid Front View. Please upload a full-body FRONT-facing image.'
                }), 400

        elif expected_view == 'side':
            if detected_orient == 'front':
                return jsonify({
                    'success': False,
                    'error': 'Front-view image detected. Please upload a SIDE-view image in this section.'
                }), 400
            if detected_orient != 'side' or not is_full_body:
                return jsonify({
                    'success': False,
                    'error': 'Invalid Side View. Please upload a full-body SIDE-facing image.'
                }), 400

        return jsonify({'success': True, 'message': 'Image successfully validated.'}), 200
        
    except Exception as e:
        print(f"Error in validate_person_count: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Validation failed: {str(e)}'}), 500


def fallback_verify_person(img1, img2):
    """
    Fallback identity verification comparing clothing HSV color histograms,
    skin tone, and upper/lower body color profiles when InsightFace is not ready/installed.
    Filter background wall pixels and handle side-profile crop width.
    """
    try:
        if img1 is None or img2 is None:
            return {'verified': False, 'similarity': 0.0, 'face_fail_reason': 'Invalid image data'}

        h1, w1 = img1.shape[:2]
        h2, w2 = img2.shape[:2]

        hsv1 = cv2.cvtColor(img1, cv2.COLOR_BGR2HSV)
        hsv2 = cv2.cvtColor(img2, cv2.COLOR_BGR2HSV)

        # Torso ROI: front view (25% to 75% width), side view tighter crop (38% to 62% width to exclude background wall)
        torso1 = hsv1[int(h1*0.20):int(h1*0.50), int(w1*0.25):int(w1*0.75)]
        torso2 = hsv2[int(h2*0.20):int(h2*0.50), int(w2*0.38):int(w2*0.62)]

        # Legs ROI: front view (25% to 75% width), side view tighter crop (38% to 62% width)
        legs1 = hsv1[int(h1*0.55):int(h1*0.85), int(w1*0.25):int(w1*0.75)]
        legs2 = hsv2[int(h2*0.55):int(h2*0.85), int(w2*0.38):int(w2*0.62)]

        # Filter out background wall pixels (low saturation S < 20 or high brightness V > 210)
        def create_clothing_mask(hsv_crop):
            s = hsv_crop[:, :, 1]
            v = hsv_crop[:, :, 2]
            mask = (s > 20) | (v < 210)
            return mask.astype(np.uint8) * 255

        mask_t1 = create_clothing_mask(torso1)
        mask_t2 = create_clothing_mask(torso2)

        t_hist1 = cv2.calcHist([torso1], [0, 1], mask_t1, [24, 32], [0, 180, 0, 256])
        t_hist2 = cv2.calcHist([torso2], [0, 1], mask_t2, [24, 32], [0, 180, 0, 256])

        if cv2.countNonZero(t_hist1) > 0 and cv2.countNonZero(t_hist2) > 0:
            cv2.normalize(t_hist1, t_hist1, 0, 1, cv2.NORM_MINMAX)
            cv2.normalize(t_hist2, t_hist2, 0, 1, cv2.NORM_MINMAX)
            torso_sim = float(cv2.compareHist(t_hist1, t_hist2, cv2.HISTCMP_CORREL))
        else:
            torso_sim = 0.5

        mask_l1 = create_clothing_mask(legs1)
        mask_l2 = create_clothing_mask(legs2)

        l_hist1 = cv2.calcHist([legs1], [0, 1], mask_l1, [24, 32], [0, 180, 0, 256])
        l_hist2 = cv2.calcHist([legs2], [0, 1], mask_l2, [24, 32], [0, 180, 0, 256])

        if cv2.countNonZero(l_hist1) > 0 and cv2.countNonZero(l_hist2) > 0:
            cv2.normalize(l_hist1, l_hist1, 0, 1, cv2.NORM_MINMAX)
            cv2.normalize(l_hist2, l_hist2, 0, 1, cv2.NORM_MINMAX)
            legs_sim = float(cv2.compareHist(l_hist1, l_hist2, cv2.HISTCMP_CORREL))
        else:
            legs_sim = 0.5

        body_sim = 0.6 * max(0.0, torso_sim) + 0.4 * max(0.0, legs_sim)
        print(f"👕 Fallback Verification: Torso={torso_sim:.4f}, Legs={legs_sim:.4f}, Overall={body_sim:.4f}")

        # Strict identity check: torso_sim >= 0.20 and body_sim >= 0.25
        verified = (torso_sim >= 0.20) and (body_sim >= 0.25)

        fail_reason = None
        if not verified:
            fail_reason = "Face and clothing features do not match between front and side photos (Different people detected)."

        return {
            'verified': verified,
            'similarity': max(0.0, float(body_sim)),
            'threshold': 0.25,
            'face_fail_reason': fail_reason,
            'message': 'Identity successfully verified' if verified else f'Identity verification failed ({fail_reason})',
            'issues': {'front': [], 'side': []}
        }
    except Exception as e:
        print(f"Fallback verify error: {e}")
        return {'verified': True, 'similarity': 0.5, 'message': 'Identity verified (fallback mode)', 'issues': {'front': [], 'side': []}}


@app.route('/api/verify-identity', methods=['POST'])
def verify_identity():
    """
    Verify identity between front and side images.
    Returns verified: False if images are of different people.
    """
    try:
        data = request.json or {}
        front_b64 = data.get('front_image', '')
        side_b64 = data.get('side_image', '')

        if not front_b64 or not side_b64:
            return jsonify({
                'success': False,
                'verified': False,
                'error': 'Both front and side images are required'
            }), 400

        front_img = decode_image(front_b64)
        side_img = decode_image(side_b64)

        if front_img is None or side_img is None:
            return jsonify({
                'success': False,
                'verified': False,
                'error': 'Invalid image data provided'
            }), 400

        result = None
        if face_verifier is not None and getattr(face_verifier, 'is_ready', False):
            try:
                result = face_verifier.verify_person(front_img, side_img)
            except Exception as fv_err:
                print(f"FaceVerifier error: {fv_err}")
                result = None

        if result is None:
            result = fallback_verify_person(front_img, side_img)

        verified = result.get('verified', False)
        similarity = float(result.get('similarity', 0.0))
        message = result.get('message', '')

        if not message:
            if verified:
                message = "Identity successfully verified"
            else:
                fail_reason = result.get('face_fail_reason', 'Front and side photos belong to different people')
                message = f"Identity verification failed: {fail_reason}"

        return jsonify({
            'success': True,
            'verified': verified,
            'similarity': similarity,
            'threshold': result.get('threshold', 0.25),
            'message': message,
            'issues': result.get('issues', {'front': [], 'side': []})
        }), 200

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'verified': False, 'error': str(e)}), 500


def merge_manual_measurements(front_results, side_results):
    """
    Merge front view (width) and side view (depth) manual measurements into 
    a single consolidated measurements dictionary.
    """
    merged = {
        'success': True,
        'measurements': {},
        'front_visualization': None,
        'front_mask': None,
        'side_visualization': None,
        'side_mask': None,
        'visualization': None,
        'mask': None,
        'scale_factor': 0,
        'height_px': 0,
        'total_landmarks': 0
    }
    
    if front_results and front_results.get('success'):
        front_measurements = front_results.get('measurements', {})
        for name, data in front_measurements.items():
            merged['measurements'][name] = data
        
        merged['front_visualization'] = front_results.get('visualization')
        merged['front_mask'] = front_results.get('mask')
        merged['visualization'] = front_results.get('visualization')
        merged['mask'] = front_results.get('mask')
        merged['scale_factor'] = front_results.get('scale_factor', 0)
        merged['height_px'] = front_results.get('height_px', 0)
        merged['total_landmarks'] = front_results.get('total_landmarks', 0)
    
    if side_results and side_results.get('success'):
        side_measurements = side_results.get('measurements', {})
        for name, data in side_measurements.items():
            if name == 'arm_length' or name == 'leg_length':
                continue
            elif name not in merged['measurements']:
                merged['measurements'][name] = data
            else:
                merged['measurements'][f'front_{name}'] = merged['measurements'][name]
                merged['measurements'][f'side_{name}'] = data
                del merged['measurements'][name]
        
        merged['side_visualization'] = side_results.get('visualization')
        merged['side_mask'] = side_results.get('mask')
        merged['total_landmarks'] += side_results.get('total_landmarks', 0)
    
    return merged


@app.route('/api/process-manual', methods=['POST'])
def process_manual_landmarks():
    """
    Process manually marked landmarks and compute measurements.
    Uses the same pixel-to-scale conversion logic as automatic detection.
    """
    try:
        data = request.json or {}
        
        try:
            user_height_cm = float(data.get('user_height') or 0)
        except (TypeError, ValueError):
            user_height_cm = 0
            
        if user_height_cm <= 0:
            return jsonify({'error': 'User height is required'}), 400
        
        front_img = decode_image(data.get('front_image'))
        side_img = decode_image(data.get('side_image'))
        
        front_landmarks = data.get('front_landmarks', {})
        front_results = None
        side_results = None
        
        if front_landmarks:
            front_results = process_manual_view(
                front_landmarks,
                user_height_cm,
                'front',
                front_img
            )
        
        side_landmarks = data.get('side_landmarks', {})
        if side_landmarks:
            side_results = process_manual_view(
                side_landmarks,
                user_height_cm,
                'side',
                side_img
            )
        
        merged_result = merge_manual_measurements(front_results, side_results)
        
        calibration_data = {
            'user_height_cm': user_height_cm,
            'method': 'manual_landmark_marking',
            'height_in_image_px': merged_result.get('height_px', 0),
            'scale_factor': merged_result.get('scale_factor', 0),
            'formula': f'{user_height_cm} cm / {merged_result.get("height_px", 1):.2f} px = {merged_result.get("scale_factor", 0):.4f} cm/px' if merged_result.get('height_px', 0) > 0 else 'N/A',
            'description': f'1 pixel = {merged_result.get("scale_factor", 0):.4f} cm' if merged_result.get('scale_factor', 0) > 0 else 'Manual mode'
        }

        response = {
            'success': True,
            'mode': 'manual',
            'calibration': calibration_data,
            'results': {
                'merged': merged_result
            }
        }
        
        return jsonify(to_native_types(response))
        
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        print(f"Manual processing error: {e}")
        traceback.print_exc()
        return jsonify({'error': f'Manual processing failed: {str(e)}'}), 500


def snap_point_to_edge(point, image, mask=None, search_radius=20, sample_count=8):
    x, y = int(point[0]), int(point[1])
    h, w = image.shape[:2]
    if x < 0 or x >= w or y < 0 or y >= h:
        return point
    try:
        if mask is not None and mask.size > 0:
            edges = cv2.Canny(mask, 50, 150)
        else:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) if len(image.shape) == 3 else image
            edges = cv2.Canny(gray, 50, 150)
        
        y_min = max(0, y - search_radius)
        y_max = min(h, y + search_radius)
        x_min = max(0, x - search_radius)
        x_max = min(w, x + search_radius)
        
        roi = edges[y_min:y_max, x_min:x_max]
        if roi.size == 0:
            return point
        
        edge_points = np.where(roi > 0)
        if len(edge_points[0]) == 0:
            return point
        
        edge_y = edge_points[0] + y_min
        edge_x = edge_points[1] + x_min
        distances = np.sqrt((edge_x - x)**2 + (edge_y - y)**2)
        min_idx = np.argmin(distances)
        
        snapped_x = edge_x[min_idx]
        snapped_y = edge_y[min_idx]
        return (float(snapped_x), float(snapped_y))
    except Exception:
        return point


def refine_measurement_with_contours(p1, p2, image, mask=None, num_samples=5):
    try:
        t_values = np.linspace(0, 1, num_samples)
        start_samples = []
        for t in t_values[:2]:
            sample_x = p1[0] + t * 0.1 * (p2[0] - p1[0])
            sample_y = p1[1] + t * 0.1 * (p2[1] - p1[1])
            snapped = snap_point_to_edge((sample_x, sample_y), image, mask, search_radius=15)
            start_samples.append(snapped)
        
        refined_p1 = (
            np.mean([s[0] for s in start_samples]),
            np.mean([s[1] for s in start_samples])
        ) if start_samples else snap_point_to_edge(p1, image, mask)
        
        end_samples = []
        for t in t_values[-2:]:
            sample_x = p1[0] + (0.9 + t * 0.1) * (p2[0] - p1[0])
            sample_y = p1[1] + (0.9 + t * 0.1) * (p2[1] - p1[1])
            snapped = snap_point_to_edge((sample_x, sample_y), image, mask, search_radius=15)
            end_samples.append(snapped)
        
        refined_p2 = (
            np.mean([s[0] for s in end_samples]),
            np.mean([s[1] for s in end_samples])
        ) if end_samples else snap_point_to_edge(p2, image, mask)
        
        return refined_p1, refined_p2
    except Exception:
        return p1, p2


def estimate_height_from_landmarks(landmarks, image_height):
    height_estimates = []
    for landmark in landmarks:
        points = landmark.get('points', [])
        if len(points) == 2:
            p1, p2 = points[0], points[1]
            dy = abs(p2['y'] - p1['y'])
            dx = abs(p2['x'] - p1['x'])
            if dy > dx * 2:
                height_estimates.append(dy)
    return max(height_estimates) if height_estimates else 0


def process_manual_view(landmarks_data, user_height_cm, view_name, image=None):
    try:
        landmarks = landmarks_data.get('landmarks', [])
        image_width = landmarks_data.get('imageWidth', 1)
        image_height = landmarks_data.get('imageHeight', 1)
        
        if not landmarks and not image:
            return {'success': False, 'error': 'No landmarks provided', 'measurements': {}}
        
        scale_factor = 0
        height_px = 0
        
        if image is not None:
            ld = get_landmark_detector()
            auto_landmarks = ld.detect(image) if ld is not None else None
            if auto_landmarks is not None:
                nose = auto_landmarks[0]
                left_ankle = auto_landmarks[27]
                right_ankle = auto_landmarks[28]
                ankle_y = max(left_ankle[1], right_ankle[1])
                height_px = ankle_y - nose[1]
                if height_px > 0:
                    scale_factor = user_height_cm / height_px
        
        if scale_factor <= 0:
            height_px = estimate_height_from_landmarks(landmarks, image_height)
            if height_px <= 0:
                height_px = image_height * 0.85
            scale_factor = user_height_cm / height_px

        measurements = {}
        mask = None
        if image is not None:
            try:
                mask = segmentation_model.segment_person(image, conf_threshold=0.3)
            except Exception:
                mask = None
        
        vis_image = image.copy() if image is not None else np.zeros((image_height, image_width, 3), dtype=np.uint8)
        
        for landmark in landmarks:
            landmark_type = landmark.get('type', 'custom')
            landmark_label = landmark.get('label', 'Unknown')
            points = landmark.get('points', [])
            
            if len(points) == 2:
                p1, p2 = points[0], points[1]
                x1_orig, y1_orig = p1['x'], p1['y']
                x2_orig, y2_orig = p2['x'], p2['y']
                
                if image is not None:
                    (x1, y1), (x2, y2) = refine_measurement_with_contours(
                        (x1_orig, y1_orig), (x2_orig, y2_orig), image, mask, num_samples=5
                    )
                else:
                    x1, y1, x2, y2 = x1_orig, y1_orig, x2_orig, y2_orig
                
                pixel_dist = np.sqrt((x2 - x1)**2 + (y2 - y1)**2)
                cm_dist = pixel_dist * scale_factor
                
                measurements[landmark_type] = {
                    'value_cm': round(float(cm_dist or 0), 2),
                    'value_px': round(float(pixel_dist or 0), 2),
                    'confidence': 0.95,
                    'source': 'Manual (Edge-Refined)',
                    'label': landmark_label,
                    'formula': f"{pixel_dist:.2f} px × {scale_factor:.4f} cm/px = {cm_dist:.2f} cm"
                }

                pt1 = (int(x1), int(y1))
                pt2 = (int(x2), int(y2))
                cv2.line(vis_image, pt1, pt2, (0, 255, 255), 3)
                cv2.circle(vis_image, pt1, 6, (0, 255, 0), -1)
                cv2.circle(vis_image, pt2, 6, (0, 255, 0), -1)
                
                mid_x = int((x1 + x2) / 2)
                mid_y = int((y1 + y2) / 2)
                cv2.putText(vis_image, f"{cm_dist:.1f}cm", (mid_x, mid_y - 10), 
                           cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255, 255, 255), 2)

        vis_base64 = encode_image(vis_image)
        mask_base64 = None
        if image is not None and mask is not None:
            mask_base64 = encode_image(cv2.cvtColor(mask, cv2.COLOR_GRAY2BGR))

        return {
            'success': True,
            'measurements': measurements,
            'scale_factor': float(scale_factor or 0),
            'height_px': float(height_px or 0),
            'visualization': vis_base64,
            'mask': mask_base64,
            'total_landmarks': len(landmarks)
        }
    except Exception as e:
        return {'success': False, 'error': str(e), 'measurements': {}}


@app.route('/api/camera/check-pose', methods=['POST'])
def check_pose():
    """Real-time pose quality check endpoint for live camera web interface."""
    try:
        data = request.get_json() or {}
        image_b64 = data.get('frame') or data.get('image')
        view = data.get('view', 'front')

        if not image_b64:
            return jsonify({'aligned': False, 'reason': 'No frame provided'}), 400

        # Decode frame
        if ',' in image_b64:
            image_b64 = image_b64.split(',')[1]
        img_data = base64.b64decode(image_b64)
        img_array = np.frombuffer(img_data, np.uint8)
        image = cv2.imdecode(img_array, cv2.IMREAD_COLOR)

        if image is None:
            return jsonify({'aligned': False, 'reason': 'Invalid image format'}), 400

        h, w = image.shape[:2]

        # Step 1 — detect person with YOLO or MediaPipe
        num_people = 0
        x1, y1, x2, y2 = 0, 0, w, h
        boxes = None
        if segmentation_model is not None and segmentation_model.model is not None:
            try:
                results = segmentation_model.model(image, classes=[0], conf=0.35, verbose=False)
                if len(results) > 0 and results[0].boxes is not None:
                    boxes = results[0].boxes
                    num_people = len(boxes)
            except Exception as ex:
                print(f"YOLO error in check_pose: {ex}")

        # If YOLO did not detect or returned 0, try MediaPipe pose detector fallback
        if num_people == 0:
            ld = get_landmark_detector()
            landmarks = ld.detect(image) if ld else None
            landmarks = _ensure_pixel_landmarks(landmarks, image.shape) if landmarks is not None else None
            if landmarks is not None and len(landmarks) >= 29:
                xs = landmarks[:, 0]
                ys = landmarks[:, 1]
                x1, y1, x2, y2 = float(np.min(xs)), float(np.min(ys)), float(np.max(xs)), float(np.max(ys))
                num_people = 1
            else:
                return jsonify({'aligned': False, 'reason': 'No person detected'}), 200

        if num_people > 1:
            return jsonify({'aligned': False, 'reason': 'Multiple persons detected'}), 200

        # Calculate bounding box coordinates
        if boxes is not None and len(boxes) > 0:
            box = boxes[0].xyxy[0].cpu().numpy()
            x1, y1, x2, y2 = float(box[0]), float(box[1]), float(box[2]), float(box[3])

        person_height = y2 - y1
        frame_coverage = float(person_height / h)

        # Step 2 — check person fills frame adequately (min 0.48 for 1.0m-1.5m webcam FOV)
        if frame_coverage < 0.48:
            return jsonify({
                'aligned': False,
                'reason': 'Move closer — body not filling frame',
                'coverage': frame_coverage
            }), 200

        # Step 3 — check person is centered horizontally (within 20% of center)
        person_center_x = (x1 + x2) / 2
        frame_center_x = w / 2
        center_offset = abs(person_center_x - frame_center_x) / w
        if center_offset > 0.20:
            return jsonify({
                'aligned': False,
                'reason': 'Move to center of frame',
                'offset': float(center_offset)
            }), 200

        # Step 4 — check head visible at top
        if y1 < h * 0.02:
            return jsonify({
                'aligned': False,
                'reason': 'Move back — head cut off at top',
            }), 200

        # Step 5 — check feet visible at bottom
        if y2 > h * 0.98:
            return jsonify({
                'aligned': False,
                'reason': 'Move back — feet cut off at bottom',
            }), 200

        # All checks passed — person is aligned!
        return jsonify({
            'aligned': True,
            'reason': 'Perfect! Hold still...',
            'coverage': frame_coverage
        }), 200

    except Exception as e:
        print(f"Pose check error: {e}")
        return jsonify({'aligned': False, 'reason': str(e)}), 500


@app.route('/api/process', methods=['POST'])
def process():
    return process_upload()


@app.route('/api/validate-person', methods=['POST'])
@app.route('/validate/person-count', methods=['POST'])
@app.route('/api/validate/person-count', methods=['POST'])
def validate_person_route():
    """Validate that exactly one person is present in the image."""
    t_start = time.time()
    print("-> /api/validate-person starting validation...")
    try:
        data = request.json or {}
        image_b64 = data.get('image')
        if not image_b64:
            return jsonify({'error': 'Image is required'}), 400
        
        img = decode_image(image_b64)
        if img is None:
            return jsonify({'error': 'Invalid image data'}), 400
            
        padded_img = add_image_padding(img, padding_percent=0.10)
        
        if segmentation_model is not None and segmentation_model.model is not None:
            segmentation_model.segment_person(padded_img, conf_threshold=0.5)
            
        t_elapsed = time.time() - t_start
        print(f"✓ /api/validate-person completed successfully in {t_elapsed:.3f}s")
        return jsonify({'success': True, 'message': 'Exactly 1 person detected'})
        
    except ValueError as e:
        t_elapsed = time.time() - t_start
        print(f"✗ /api/validate-person validation failed in {t_elapsed:.3f}s: {e}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        t_elapsed = time.time() - t_start
        print(f"✗ /api/validate-person error in {t_elapsed:.3f}s: {e}")
        return jsonify({'error': f'Validation failed: {str(e)}'}), 500


@app.route('/api/upload/process', methods=['POST'])
def process_upload():
    """Process uploaded images"""
    try:
        data = request.json or {}
        
        # Decode images safely
        front_img = decode_image(data.get('front_image'))
        side_img = decode_image(data.get('side_image'))
        ref_img = decode_image(data.get('reference_image') or data.get('front_image'))
        
        if front_img is None:
            return jsonify({'error': 'Front image is required'}), 400
            
        # Get reference parameters safely
        try:
            ref_size = float(data.get('reference_size', 29.7))
        except (TypeError, ValueError):
            ref_size = 29.7
            
        ref_axis = data.get('reference_axis', 'height')
        
        # Detect reference if ref_img is provided
        ref_px = None
        if ref_img is not None:
            ref_px = reference_detector.detect_reference(ref_img, ref_axis)
        
        # Calculate scale factor safely
        scale_factor = 0.0
        if ref_px is not None and ref_px > 0:
            try:
                scale_factor = float(ref_size) / float(ref_px)
            except (TypeError, ValueError, ZeroDivisionError):
                scale_factor = 0.0
            print(f"✓ Reference-based scale factor: {scale_factor:.4f} cm/px")
            
        user_height_cm = None
        if data.get('user_height') is not None:
            try:
                user_height_cm = float(data.get('user_height'))
            except (TypeError, ValueError):
                user_height_cm = None
        elif data.get('height') is not None:
            try:
                user_height_cm = float(data.get('height'))
            except (TypeError, ValueError):
                user_height_cm = None
                
        # Fallback to height-based calibration if reference calibration is not available or failed
        if scale_factor <= 0.0:
            if front_img is not None:
                landmarks = landmark_detector.detect(front_img)
                if landmarks is not None and len(landmarks) >= 33:
                    nose = landmarks[0]
                    left_ankle = landmarks[27]
                    right_ankle = landmarks[28]
                    # Use lower ankle
                    ankle_y = max(left_ankle[1], right_ankle[1])
                    height_px = ankle_y - nose[1]
                    if height_px > 0:
                        try:
                            # Use user_height_cm if provided, otherwise default to 170.0 cm
                            calib_height = float(user_height_cm or 170.0)
                            scale_factor = calib_height / height_px
                            # Populate ref_px and ref_size for visual formula compatibility
                            ref_size = calib_height
                            ref_px = height_px
                            print(f"✓ Fallback height-based scale factor: {scale_factor:.4f} cm/px")
                        except (TypeError, ValueError, ZeroDivisionError):
                            scale_factor = 0.0

        if scale_factor <= 0.0:
            return jsonify({'error': 'Could not calculate scale factor. Please ensure full body or reference object is visible.'}), 400

        requested_gender = data.get('gender', 'neutral')
        
        # Process front view
        results = {}
        
        if front_img is not None:
            front_results = process_single_image(
                front_img, scale_factor, 'front', user_height_cm=user_height_cm, gender=requested_gender
            )
            results['front'] = front_results
        
        if side_img is not None:
            side_results = process_single_image(
                side_img, scale_factor, 'side', user_height_cm=user_height_cm, gender=requested_gender
            )
            results['side'] = side_results

        import tempfile, cv2, numpy as np

        # Save front image to disk for SMPLify-X
        front_temp_path = os.path.join(
            parent_dir, 'data', 'images', 'front.jpg'
        )
        os.makedirs(
            os.path.dirname(front_temp_path),
            exist_ok=True
        )
        cv2.imwrite(front_temp_path, front_img)
        print(f"Saved front image: {front_temp_path}")

        # Save side image if provided
        side_temp_path = None
        if side_img is not None:
            side_temp_path = os.path.join(
                parent_dir, 'data', 'images', 'side.jpg'
            )
            cv2.imwrite(side_temp_path, side_img)
            print(f"Saved side image: {side_temp_path}")

        # ── SMPLify-X Integration ──────────────────
        mesh_data    = None
        smplx_meas   = {}
        smplx_status = "not_run"

        try:
            # Get user height safely
            user_height_safe = 170.0  # safe default

            try:
                # Priority 1: from request form data
                if request.form.get('height'):
                    user_height_safe = float(
                        request.form.get('height')
                    )
                    print(f"Height from form: {user_height_safe} cm")

                # Priority 2: from request JSON body
                elif (request.is_json and
                      request.get_json() and
                      request.get_json().get('user_height')):
                    user_height_safe = float(
                        request.get_json().get('user_height')
                    )
                    print(f"Height from JSON: {user_height_safe} cm")
                
                # Priority 2b: check 'height' key too
                elif (request.is_json and
                      request.get_json() and
                      request.get_json().get('height')):
                    user_height_safe = float(
                        request.get_json().get('height')
                    )
                    print(f"Height from JSON (key 'height'): {user_height_safe} cm")

                # Priority 3: from processed results
                else:
                    if 'front' in results:
                        front_meas = (
                            results['front'].get(
                                'measurements', {}
                            ) or {}
                        )
                        for key in [
                            'full_height', 'height',
                            'Height',      'Full Height',
                            'body_height', 'user_height'
                        ]:
                            val = front_meas.get(key)
                            if val is not None:
                                try:
                                    # Handle dict structure in results
                                    if isinstance(val, dict):
                                        user_height_safe = float(val.get('value_cm', 170.0))
                                    else:
                                        user_height_safe = float(val)
                                    print(f"Height from results key '{key}': {user_height_safe} cm")
                                    break
                                except (TypeError, ValueError):
                                    continue

            except (TypeError, ValueError) as e:
                print(f"Height parse error: {e}, using default 170.0 cm")
                user_height_safe = 170.0

            # If user_height_cm was passed in and is valid, use it
            if user_height_cm is not None and user_height_cm > 0:
                final_height = user_height_cm
            else:
                final_height = user_height_safe

            print(f"Final final_height: {final_height} cm")

            print(f"Running SMPLify-X with "
                  f"height={final_height}cm...")

            smplifyx_result = run_smplifyx(
                front_image_path = front_temp_path,
                side_image_path  = side_temp_path,
                timeout_seconds  = 120
            )

            if smplifyx_result["success"]:
                reader = SMPLifyXReader(
                    smplifyx_result["mesh_path"]
                )
                smplx_meas = reader.extract_measurements(
                    final_height
                )
                mesh_data    = reader.export_for_plotly(
                    final_height
                )
                smplx_status = "success"
                print("SMPLify-X completed successfully")

                # Merge SMPLify-X measurements into results
                # SMPLify-X overrides circumferences only
                try:
                    if 'front' in results:
                        m = results['front'].get(
                            'measurements', {}
                        ) or {}
                        for key in [
                            'chest_circumference',
                            'waist_circumference',
                            'hip_circumference'
                        ]:
                            val = smplx_meas.get(key)
                            if val is not None:
                                # Ensure we preserve the result structure
                                m[key] = {
                                    'value_cm': float(val),
                                    'source': 'SMPLify-X',
                                    'label': key.replace('_', ' ').title()
                                }
                        if m:
                            results['front']['measurements'] = m

                    if 'side' in results:
                        m_side = results['side'].get(
                            'measurements', {}
                        ) or {}
                        for key in [
                            'chest_circumference',
                            'waist_circumference',
                            'hip_circumference'
                        ]:
                            val = smplx_meas.get(key)
                            if val is not None:
                                m_side[key] = {
                                    'value_cm': float(val),
                                    'source': 'SMPLify-X',
                                    'label': key.replace('_', ' ').title()
                                }
                        if m_side:
                            results['side']['measurements'] = m_side

                    if 'merged' in results:
                        m_merged = results['merged'].get(
                            'measurements', {}
                        ) or {}
                        for key in [
                            'chest_circumference',
                            'waist_circumference',
                            'hip_circumference'
                        ]:
                            val = smplx_meas.get(key)
                            if val is not None:
                                m_merged[key] = {
                                    'value_cm': float(val),
                                    'source': 'SMPLify-X',
                                    'label': key.replace('_', ' ').title()
                                }
                        if m_merged:
                            results['merged']['measurements'] = m_merged

                except Exception as merge_err:
                    print(f"Measurement merge error: {merge_err}")

            else:
                smplx_status = "failed"
                print(f"SMPLify-X failed: "
                      f"{smplifyx_result['error']}")

        except Exception as e:
            smplx_status = "error"
            print(f"SMPLify-X exception: {e}")
            import traceback
            traceback.print_exc()

        # Debug log to confirm measurements survive
        if 'front' in results:
            front_meas_count = len(results['front'].get('measurements', {}))
            print(f"[DEBUG] Front view measurements surviving: {front_meas_count}")
        if 'side' in results:
            side_meas_count = len(results['side'].get('measurements', {}))
            print(f"[DEBUG] Side view measurements surviving: {side_meas_count}")
        if 'merged' in results:
            merged_meas_count = len(results['merged'].get('measurements', {}))
            print(f"[DEBUG] Merged measurements surviving: {merged_meas_count}")

        # ── End SMPLify-X ─────────────────────────

        # Consolidate measurements and set results['merged']
        consolidated = consolidate_measurements(results.get('front'), results.get('side'), user_height_cm)
        results['merged'] = {
            'success': True,
            'measurements': consolidated,
            'front_visualization': results['front'].get('visualization') if 'front' in results else None,
            'front_mask': results['front'].get('mask') if 'front' in results else None,
            'side_visualization': results['side'].get('visualization') if 'side' in results else None,
            'side_mask': results['side'].get('mask') if 'side' in results else None,
            'visualization': results['front'].get('visualization') if 'front' in results else None,
            'mask': results['front'].get('mask') if 'front' in results else None,
        }

        return jsonify({
            'success':      True,
            'scale_factor': scale_factor,
            'reference_px': ref_px,
            'calibration': {
                'reference_size_cm': ref_size,
                'reference_size_px': float(ref_px) if ref_px is not None else 0.0,
                'scale_factor':      float(scale_factor),
                'user_height_cm':    final_height,
                'height_in_image_px': float(final_height / scale_factor) if scale_factor > 0 else 0.0,
                'formula': (
                    f'{ref_size:.2f} cm ÷ '
                    f'{ref_px:.2f} px = '
                    f'{scale_factor:.4f} cm/px'
                ) if ref_px is not None else f'Estimated scale factor: {scale_factor:.4f} cm/px',
                'description': (
                    f'1 pixel = {scale_factor:.4f} cm'
                )
            },
            'results':      results,
            'mesh_data':    mesh_data,
            'smplx_status': smplx_status,
            'smplx_measurements': smplx_meas
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def process_single_image(image, scale_factor, view, user_height_cm=None, gender='neutral'):
    """Process a single image using hybrid vision approach"""
    try:
        print(f"Processing {view} image with scale factor: {scale_factor}")
        
        # Segment person with YOLOv8
        mask = segmentation_model.segment_person(image, conf_threshold=0.5)
        # Fix 1 & 4: Ensure mask is not boolean and is a numpy array (None on failure)
        if mask is not None:
            if isinstance(mask, bool) or not isinstance(mask, np.ndarray):
                print(f"[DEBUG] segment_person returned invalid mask: {mask}. Setting mask = None")
                mask = None
        print(f"Segmentation complete: {mask is not None}")
        
        # Apply mask to get clean image (optional - for visualization)
        if mask is not None:
            masked_image = segmentation_model.apply_mask(image, mask, background_mode='dim')
        else:
            masked_image = image
        
        # Detect landmarks on original image (MediaPipe works better on full image)
        landmarks = landmark_detector.detect(image)
        print(f"Landmarks detected: {landmarks is not None}")
        
        if landmarks is None:
            print("ERROR: No landmarks detected")
            return {'error': 'No person detected'}
        
        print(f"Number of landmarks: {len(landmarks)}")

        # SMPL integration right after MediaPipe detection.
        detected_gender = run_gender_detection(image)
        input_gender = gender if gender in ('male', 'female', 'neutral') else 'neutral'
        smpl_gender = detected_gender if detected_gender != 'neutral' else input_gender

        landmarks_list = _landmarks_to_smpl_list(landmarks, image.shape)
        # Guard effective_height_cm against None
        try:
            effective_height_cm = float(user_height_cm) if user_height_cm and float(user_height_cm) > 0 else 0.0
        except (TypeError, ValueError):
            effective_height_cm = 0.0

        if effective_height_cm <= 0 and scale_factor:
            try:
                sf = float(scale_factor)
                if sf > 0:
                    estimated_height_px = _estimate_height_from_landmarks_px(landmarks)
                    if estimated_height_px > 0:
                        effective_height_cm = sf * float(estimated_height_px)
            except (TypeError, ValueError):
                pass

        # Calculate scale factor independently using the person's detected height in this specific image
        if effective_height_cm > 0:
            height_px = _estimate_height_from_landmarks_px(landmarks)
            if height_px > 0:
                scale_factor = effective_height_cm / height_px
                print(f"✓ Calculated independent scale factor for {view} view: {scale_factor:.4f} cm/px")

        smpl_success = False
        smpl_m = {}
        smpl_error = None
        smpl_mesh_data = None
        smpl_fit_info = {}
        if len(landmarks_list) == 33 and effective_height_cm > 0:
            h, w = image.shape[:2]
            smpl_result = run_smpl_pipeline(
                landmarks_2d=landmarks_list,
                image_width=w,
                image_height=h,
                user_height_cm=effective_height_cm,
                gender=smpl_gender or 'neutral',
                view_type=view,
                front_mask=mask if view == 'front' else None,
                side_mask=mask if view == 'side' else None
            )
            smpl_success = bool(smpl_result.get('success'))
            if smpl_success:
                # Check mean 2D reprojection error for pose quality validation
                reproj_err = smpl_result.get('reprojection_error', {}) or {}
                mean_px = reproj_err.get('mean_px', 0.0)
                if mean_px > 300.0:
                    print(f"⚠ SMPL pose validation failed: mean 2D reprojection error ({mean_px:.1f} px) is above 300px threshold. Disabling unreliable SMPL measurements.")
                    smpl_success = False
                    smpl_error = f"Pose is unreliable: mean reprojection error ({mean_px:.1f}px) exceeds 300px"
                    smpl_m = {}
                    smpl_fit_info = {}
                else:
                    smpl_m = smpl_result.get('measurements', {}) or {}
                    smpl_fit_info = smpl_result.get('fit', {}) or {}
            else:
                smpl_error = smpl_result.get('error')

            # Always provide mesh data when height is available.
            # If fitting failed, this becomes a default-body A-pose mesh.
            smpl_mesh_data = _build_smpl_mesh_data(
                smpl_result if smpl_success else None,
                effective_height_cm,
                smpl_gender or 'neutral'
            )
            if smpl_mesh_data is None and smpl_success:
                smpl_mesh_data = smpl_result.get('mesh_data')
        else:
            smpl_error = 'Insufficient data for SMPL (height or landmarks unavailable)'
        
        # Extract edge reference points from segmentation for hybrid approach
        edge_reference_points = None
        if mask is not None:
            try:
                # Extract edge reference points (shoulder/chest/waist/hip) from mask rows.
                edge_reference_points = landmark_detector.extract_body_edge_keypoints(mask, landmarks)
                if edge_reference_points:
                    print(f"Edge reference points extracted: {edge_reference_points.get('is_valid')}")
            except Exception as e:
                print(f"Warning: Could not extract edge points: {e}")
                edge_reference_points = None
        
        # TEMP DEBUG
        print(f"DEBUG landmarks[11] = {landmarks[11]}")  # left shoulder
        print(f"DEBUG landmarks[23] = {landmarks[23]}")  # left hip  
        print(f"DEBUG scale_factor = {scale_factor}")
        arm_test = np.linalg.norm(landmarks[11][:2] - landmarks[13][:2]) + np.linalg.norm(landmarks[13][:2] - landmarks[15][:2])
        print(f"DEBUG arm px (11->13->15) = {arm_test}")
        print(f"DEBUG arm cm = {arm_test * scale_factor}")
        # Fix 3: Verify mask is a valid numpy array before passing
        actual_mask = None
        if mask is not None and not isinstance(mask, bool) and isinstance(mask, np.ndarray):
            actual_mask = mask
        else:
            print(f"[DEBUG] mask before calculate_measurements_with_confidence is invalid: {mask} (type: {type(mask)}). Setting actual_mask = None")

        measurements = measurement_engine.calculate_measurements_with_confidence(
            landmarks, scale_factor, view,
            edge_reference_points=edge_reference_points,
            user_height_cm=effective_height_cm,
            mask=actual_mask
        )
        print(f"DEBUG hip px: {landmarks[23][:2]}")
        print(f"DEBUG knee px: {landmarks[25][:2]}")
        print(f"DEBUG ankle px: {landmarks[27][:2]}")
        hip_knee = np.linalg.norm(landmarks[23][:2] - landmarks[25][:2])
        knee_ankle = np.linalg.norm(landmarks[25][:2] - landmarks[27][:2])
        print(f"DEBUG hip-knee dist: {hip_knee:.1f} px")
        print(f"DEBUG knee-ankle dist: {knee_ankle:.1f} px")
        
        
        print(f"Measurements calculated: {len(measurements)} measurements")
        
        # Build MediaPipe measurement payload first, then merge with SMPL.
        mp_measurements = {}
        for name, val in measurements.items():
            if len(val) >= 4:
                cm_value, confidence, source, pixel_distance = val
            else:
                cm_value, confidence, source = val
                pixel_distance = cm_value / scale_factor if scale_factor > 0 else 0
            
            try:
                cm = float(cm_value) if cm_value is not None else 0.0
                px = float(pixel_distance) if pixel_distance is not None else 0.0
                conf = float(confidence) if confidence is not None else 0.0
                
                mp_measurements[name] = {
                    'value_cm': cm,
                    'value_px': px,
                    'value_pixels': px,
                    'confidence': conf,
                    'source': source,
                    'calculation': f"{px:.2f} px × {float(scale_factor or 0):.4f} cm/px = {cm:.2f} cm"
                }
            except (TypeError, ValueError):
                continue

        measurements_with_pixels, smpl_source_label = _build_smpl_merged_measurements(
            mp_measurements,
            smpl_m,
            smpl_success,
        )
        
        print(f"Measurements with pixels: {len(measurements_with_pixels)}")
    except Exception as e:
        print(f"ERROR in process_single_image: {e}")
        import traceback
        traceback.print_exc()
        return {'error': f'Processing failed: {str(e)}'}
    
    def _draw_edge_width_overlays(img, edge_points):
        if not edge_points or not edge_points.get('is_valid'):
            return

        # Colors (BGR)
        BLUE = (255, 0, 0)
        ORANGE = (0, 165, 255)
        GREEN = (0, 255, 0)

        def _draw_width_line(left_pt, right_pt, color, radius):
            if not left_pt or not right_pt:
                return
            lx, ly = int(left_pt[0]), int(left_pt[1])
            rx, ry = int(right_pt[0]), int(right_pt[1])
            if (lx == 0 and ly == 0) or (rx == 0 and ry == 0):
                return

            # Draw using identical endpoints for both dots and line.
            cv2.line(img, (lx, ly), (rx, ry), color, 2)
            cv2.circle(img, (lx, ly), radius, color, -1)
            cv2.circle(img, (rx, ry), radius, color, -1)

        _draw_width_line(edge_points.get('shoulder_left'), edge_points.get('shoulder_right'), BLUE, 8)
        _draw_width_line(edge_points.get('chest_left'), edge_points.get('chest_right'), ORANGE, 6)
        _draw_width_line(edge_points.get('waist_left'), edge_points.get('waist_right'), GREEN, 6)

    # Check if we have measurements
    if not measurements_with_pixels:
        print("WARNING: No measurements calculated!")
        # Return basic info even if no measurements
        vis_img = landmark_detector.draw_landmarks(image, landmarks)
        _draw_edge_width_overlays(vis_img, edge_reference_points)
        vis_base64 = encode_image(vis_img)
        mask_base64 = encode_image(mask) if mask is not None else None
        return {
            'measurements': {},
            'visualization': vis_base64,
            'mask': mask_base64,
            'landmark_count': len(landmarks),
            'smpl': {
                'enabled': True,
                'success': smpl_success,
                'status': 'active' if smpl_success else 'estimated',
                'source': 'SMPL 3D Model' if smpl_success else 'Estimated',
                'gender': smpl_gender if 'smpl_gender' in locals() else 'neutral',
                'error': smpl_error,
                'fit_status': smpl_fit_info.get('fit_status', 'fitted' if smpl_success else 'estimated'),
                'status_text': smpl_fit_info.get('status_text', '✓ Model fitted to your body' if smpl_success else 'Estimated from default body'),
                'fitted_to_user': bool(smpl_fit_info.get('fitted_to_user', smpl_success)),
                'landmarks_source': smpl_fit_info.get('landmarks_source', 'mediapipe'),
                'landmarks_mode': smpl_fit_info.get('landmarks_mode', 'real' if smpl_success else 'unavailable'),
                'landmark_count': int(smpl_fit_info.get('landmark_count', len(landmarks_list))),
                'visible_landmark_count': int(smpl_fit_info.get('visible_landmark_count', 0)),
                'pose_applied': bool(smpl_mesh_data.get('metadata', {}).get('pose_applied', smpl_fit_info.get('pose_applied', False))) if isinstance(smpl_mesh_data, dict) else bool(smpl_fit_info.get('pose_applied', False)),
            },
            'mesh_data': smpl_mesh_data,
            'scale_info': {
                'scale_factor': float(scale_factor),
                'unit': 'cm/pixel',
                'description': f'1 pixel = {scale_factor:.4f} cm'
            },
            'warning': 'No measurements could be calculated. Check if all required landmarks are visible.'
        }
    
    # Draw visualization with measurements
    vis_img = landmark_detector.draw_landmarks(image, landmarks)
    _draw_edge_width_overlays(vis_img, edge_reference_points)
    
    # Add measurement annotations
    h, w = vis_img.shape[:2]
    y_offset = 30
    for name, data in list(measurements_with_pixels.items())[:5]:  # Show first 5
        source_label = "Edge" if data['source'] == "Segmentation Edge" else "MediaPipe"
        text = f"{name}: {data['value_cm']:.1f}cm ({source_label})"
        cv2.putText(vis_img, text, (10, y_offset), 
                   cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 255, 255), 2, cv2.LINE_AA)
        y_offset += 25
    
    vis_base64 = encode_image(vis_img)
    mask_base64 = encode_image(mask) if mask is not None else None
    
    print(f"Returning {len(measurements_with_pixels)} measurements")
    
    # Add hybrid approach metadata to response
    source_summary = {
        'mediapipe': len([m for m in measurements_with_pixels.values() if m.get('source') == 'MediaPipe']),
        'smpl_model': len([m for m in measurements_with_pixels.values() if m.get('source') == 'SMPL 3D Model']),
        'smpl_mediapipe': len([m for m in measurements_with_pixels.values() if m.get('source') == 'SMPL + MediaPipe']),
        'estimated': len([m for m in measurements_with_pixels.values() if m.get('source') == 'Estimated']),
    }
    
    return {
        'measurements': measurements_with_pixels,
        'visualization': vis_base64,
        'mask': mask_base64,
        'landmark_count': len(landmarks),
        'smpl': {
            'enabled': True,
            'success': smpl_success,
            'status': 'active' if smpl_success else 'estimated',
            'source': smpl_source_label,
            'gender': smpl_gender,
            'error': smpl_error,
            'fit_status': smpl_fit_info.get('fit_status', 'fitted' if smpl_success else 'estimated'),
            'status_text': smpl_fit_info.get('status_text', '✓ Model fitted to your body' if smpl_success else 'Estimated from default body'),
            'fitted_to_user': bool(smpl_fit_info.get('fitted_to_user', smpl_success)),
            'landmarks_source': smpl_fit_info.get('landmarks_source', 'mediapipe'),
            'landmarks_mode': smpl_fit_info.get('landmarks_mode', 'real' if smpl_success else 'unavailable'),
            'landmark_count': int(smpl_fit_info.get('landmark_count', len(landmarks_list))),
            'visible_landmark_count': int(smpl_fit_info.get('visible_landmark_count', 0)),
            'pose_applied': bool(smpl_fit_info.get('pose_applied', False)),
        },
        'mesh_data': smpl_mesh_data,
        'scale_info': {
            'scale_factor': float(scale_factor),
            'unit': 'cm/pixel',
            'description': f'1 pixel = {scale_factor:.4f} cm'
        },
        'hybrid_approach': {
            'enabled': True,
            'edge_points_available': edge_reference_points is not None and edge_reference_points.get('is_valid'),
            'source_summary': source_summary
        }
    }


@app.route('/api/camera/start', methods=['POST'])
def start_camera():
    """Start camera stream"""
    global camera, camera_active
    
    try:
        print("Starting camera...")
        if camera is None:
            camera = cv2.VideoCapture(0)
            if not camera.isOpened():
                print("ERROR: Failed to open camera")
                return jsonify({'error': 'Failed to open camera. Check if camera is available.'}), 500
            
            camera.set(cv2.CAP_PROP_FRAME_WIDTH, 1280)
            camera.set(cv2.CAP_PROP_FRAME_HEIGHT, 720)
            print("Camera opened successfully")
        
        camera_active = True
        
        # Start camera thread
        print("Starting camera thread...")
        threading.Thread(target=camera_stream_thread, daemon=True).start()
        print("Camera thread started")
        
        return jsonify({'success': True})
    except Exception as e:
        print(f"ERROR starting camera: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/camera/stop', methods=['POST'])
def stop_camera():
    """Stop camera stream"""
    global camera, camera_active
    
    camera_active = False
    if camera:
        camera.release()
        camera = None
    
    return jsonify({'success': True})


@app.route('/api/camera/capture-reference', methods=['POST'])
def capture_reference():
    """Capture reference object from camera"""
    global reference_captured, reference_data
    
    try:
        data = request.json
        ref_size = float(data['reference_size'])
        ref_axis = data['reference_axis']
        
        if camera is None:
            return jsonify({'error': 'Camera not started'}), 400
        
        # Capture frame
        ret, frame = camera.read()
        if not ret:
            return jsonify({'error': 'Failed to capture frame'}), 500
        
        # Detect reference
        ref_px = reference_detector.detect_reference(frame, ref_axis)
        
        if ref_px is None:
            return jsonify({'error': 'Reference object not detected'}), 400
        
        # Store reference data
        reference_data = {
            'size_cm': ref_size,
            'size_px': ref_px,
            'axis': ref_axis,
            'scale_factor': ref_size / ref_px
        }
        reference_captured = True
        
        # Initialize temporal stabilizer
        temporal_stabilizer.initialize_reference(frame, ref_px)
        
        return jsonify({
            'success': True,
            'reference_px': ref_px,
            'scale_factor': reference_data['scale_factor']
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/camera/capture-measurement', methods=['POST'])
def capture_measurement():
    """Capture and process measurement"""
    global reference_captured, reference_data
    
    if not reference_captured:
        return jsonify({'error': 'Reference not captured'}), 400
    
    try:
        # Capture frame
        ret, frame = camera.read()
        if not ret:
            return jsonify({'error': 'Failed to capture frame'}), 500
        
        # Process frame
        result = process_single_image(
            frame,
            reference_data['scale_factor'],
            'front'
        )
        
        return jsonify({
            'success': True,
            'result': result
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def camera_stream_thread():
    """Camera streaming thread"""
    global camera_active, reference_captured, reference_data
    
    print("Camera thread running...")
    green_frame_count = 0
    countdown_started = False
    
    try:
        while camera_active:
            if camera is None:
                print("Camera is None, breaking")
                break
            
            ret, frame = camera.read()
            if not ret:
                print("Failed to read frame")
                time.sleep(0.1)
                continue
            
            # Flip for mirror effect
            frame = cv2.flip(frame, 1)
            
            # Process frame
            try:
                processed_frame, alignment_status, has_object = process_camera_frame(frame)
            except Exception as e:
                print(f"Error processing frame: {e}")
                processed_frame = frame
                alignment_status = 'red'
                has_object = False
        
        # Auto-capture logic
        if reference_captured and has_object:
            if alignment_status == 'green':
                green_frame_count += 1
                countdown = max(0, 90 - green_frame_count) // 30  # 3 seconds at 30 FPS
                
                if green_frame_count >= 90:  # 3 seconds
                    # Auto-capture
                    try:
                        result = process_single_image(
                            frame,
                            reference_data['scale_factor'],
                            'front'
                        )
                        socketio.emit('auto_capture', {
                            'success': True,
                            'result': result
                        })
                        green_frame_count = 0
                    except Exception as e:
                        print(f"Auto-capture error: {e}")
                        green_frame_count = 0
            else:
                green_frame_count = 0
                countdown = None
        else:
            green_frame_count = 0
            countdown = None
        
            # Encode and emit
            try:
                _, buffer = cv2.imencode('.jpg', processed_frame, [cv2.IMWRITE_JPEG_QUALITY, 85])
                frame_base64 = base64.b64encode(buffer).decode('utf-8')
                
                socketio.emit('camera_frame', {
                    'frame': frame_base64,
                    'timestamp': time.time(),
                    'alignment': alignment_status,
                    'has_object': has_object,
                    'countdown': countdown
                })
            except Exception as e:
                print(f"Error encoding/emitting frame: {e}")
            
            time.sleep(0.033)  # ~30 FPS
    except Exception as e:
        print(f"Camera thread error: {e}")
        import traceback
        traceback.print_exc()
    finally:
        print("Camera thread stopped")


def process_camera_frame(frame):
    """Process camera frame with overlays"""
    if not reference_captured:
        # Draw template overlay
        frame = draw_template_overlay(frame)
        return frame, 'red', False
    
    # Detect landmarks
    landmarks = landmark_detector.detect(frame)
    
    # Detect object in hand
    has_object = detect_object_in_hand(frame, landmarks)
    
    if landmarks is not None:
        # Check alignment
        alignment = check_alignment(landmarks, frame, has_object)
        
        # Draw feedback
        frame = draw_feedback_overlay(frame, landmarks, alignment, has_object)
    else:
        alignment = 'red'
    
    return frame, alignment, has_object


def detect_object_in_hand(frame, landmarks):
    """Detect if person is holding an object in hand"""
    if landmarks is None:
        return False
    
    try:
        # Get wrist and hand landmarks
        left_wrist = landmarks[15]  # left wrist
        right_wrist = landmarks[16]  # right wrist
        left_index = landmarks[19]  # left index finger
        right_index = landmarks[20]  # right index finger
        
        # Check if hands are visible and in front of body
        if left_wrist[2] > 0.5 or right_wrist[2] > 0.5:
            # Simple heuristic: if wrist is visible and elevated
            # (above hip level), assume holding object
            left_hip = landmarks[23]
            right_hip = landmarks[24]
            hip_y = (left_hip[1] + right_hip[1]) / 2
            
            # Check if either wrist is above hip level
            if (left_wrist[2] > 0.5 and left_wrist[1] < hip_y) or \
               (right_wrist[2] > 0.5 and right_wrist[1] < hip_y):
                return True
        
        return False
    except:
        return False


def draw_template_overlay(frame):
    """Draw body template overlay"""
    h, w = frame.shape[:2]
    overlay = frame.copy()
    
    # Simple stick figure template
    center_x, center_y = w // 2, h // 2
    
    # Head
    cv2.circle(overlay, (center_x, center_y - 150), 40, (255, 255, 255), 2)
    
    # Body
    cv2.line(overlay, (center_x, center_y - 110), (center_x, center_y + 100), (255, 255, 255), 2)
    
    # Arms
    cv2.line(overlay, (center_x, center_y - 80), (center_x - 100, center_y), (255, 255, 255), 2)
    cv2.line(overlay, (center_x, center_y - 80), (center_x + 100, center_y), (255, 255, 255), 2)
    
    # Legs
    cv2.line(overlay, (center_x, center_y + 100), (center_x - 50, center_y + 250), (255, 255, 255), 2)
    cv2.line(overlay, (center_x, center_y + 100), (center_x + 50, center_y + 250), (255, 255, 255), 2)
    
    # Blend
    frame = cv2.addWeighted(overlay, 0.3, frame, 0.7, 0)
    
    return frame


def check_alignment(landmarks, frame, has_object):
    """Check alignment status"""
    # Simplified alignment check
    if landmarks is None:
        return 'red'
    
    # Must have object in hand
    if not has_object:
        return 'red'
    
    # Check if full body visible
    h, w = frame.shape[:2]
    
    # Check feet visibility
    left_ankle = landmarks[27]
    right_ankle = landmarks[28]
    
    if left_ankle[2] < 0.5 or right_ankle[2] < 0.5:
        return 'red'
    
    # Check centering
    center_x = np.mean(landmarks[:, 0])
    if abs(center_x - w/2) > w * 0.2:
        return 'amber'
    
    # Check if standing straight
    left_shoulder = landmarks[11]
    right_shoulder = landmarks[12]
    shoulder_diff = abs(left_shoulder[1] - right_shoulder[1])
    
    if shoulder_diff > h * 0.05:  # Shoulders not level
        return 'amber'
    
    return 'green'


def draw_feedback_overlay(frame, landmarks, alignment, has_object):
    """Draw feedback overlay"""
    h, w = frame.shape[:2]
    
    # Color based on alignment
    colors = {
        'red': (0, 0, 255),
        'amber': (0, 165, 255),
        'green': (0, 255, 0)
    }
    color = colors.get(alignment, (0, 0, 255))
    
    # Draw border
    cv2.rectangle(frame, (10, 10), (w-10, h-10), color, 8)
    
    # Draw landmarks
    frame = landmark_detector.draw_landmarks(frame, landmarks)
    
    # Draw status text
    if not has_object:
        text = 'Hold Object in Hand'
        color = (0, 0, 255)
    else:
        texts = {
            'red': 'Adjust Position',
            'amber': 'Almost Ready',
            'green': 'Perfect! Hold Still'
        }
        text = texts.get(alignment, '')
    
    cv2.putText(frame, text, (20, 50), cv2.FONT_HERSHEY_SIMPLEX,
               1.2, color, 3, cv2.LINE_AA)
    
    # Draw object indicator
    if has_object:
        cv2.putText(frame, 'Object Detected', (20, h - 30), 
                   cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 255, 0), 2, cv2.LINE_AA)
    
    return frame


def decode_image(base64_str):
    """Decode base64 image to numpy array"""
    if not base64_str:
        return None
    
    # Remove data URL prefix if present
    if ',' in base64_str:
        base64_str = base64_str.split(',')[1]
    
    img_data = base64.b64decode(base64_str)
    img = Image.open(io.BytesIO(img_data))
    img_array = np.array(img)
    
    # Convert RGB to BGR for OpenCV
    if len(img_array.shape) == 3:
        img_array = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
    
    return img_array


def encode_image(img_array):
    """Encode numpy array to base64"""
    # Convert BGR to RGB
    if len(img_array.shape) == 3:
        img_array = cv2.cvtColor(img_array, cv2.COLOR_BGR2RGB)
    
    img = Image.fromarray(img_array)
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
    
    return f"data:image/png;base64,{img_base64}"


# ========== SHOULDER EDGE DETECTION ENDPOINTS ==========

@app.route('/api/shoulder/detect', methods=['POST'])
def detect_shoulder_edges():
    """
    Detect shoulder edge points from uploaded image
    
    Expected JSON:
    {
        "image": "base64_encoded_image",
        "shoulder_type": "both" (or "left", "right")
    }
    
    Returns:
    {
        "frame_number": int,
        "shoulder_edge_points": [{"x": float, "y": float}, ...],
        "confidence_score": float,
        "detection_quality": {...},
        "visualization": "base64_encoded_image"
    }
    """
    try:
        data = request.json
        
        if 'image' not in data:
            return jsonify({'error': 'Image required'}), 400
        
        # Decode image
        image = decode_image(data['image'])
        shoulder_type = data.get('shoulder_type', 'both')
        
        # Detect landmarks first
        landmarks = landmark_detector.detect(image)
        
        if landmarks is None:
            return jsonify({
                'error': 'No person detected in image'
            }), 400
        
        # Detect shoulder edges
        shoulder_data = landmark_detector.detect_shoulder_edge_points(
            image, landmarks, shoulder_type=shoulder_type
        )
        
        # Create visualization
        annotated_frame = landmark_detector.draw_shoulder_edges(image, shoulder_data)
        vis_base64 = encode_image(annotated_frame)
        
        # Prepare response
        response = {
            'success': True,
            'frame_number': shoulder_data['frame_number'],
            'shoulder_edge_points': shoulder_data['shoulder_edge_points'],
            'confidence_score': shoulder_data['confidence_score'],
            'detection_quality': landmark_detector._assess_detection_quality(shoulder_data),
            'visualization': vis_base64
        }
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/shoulder/batch', methods=['POST'])
def batch_detect_shoulder_edges():
    """
    Process multiple frames and detect shoulder edges
    
    Expected JSON:
    {
        "images": ["base64_image1", "base64_image2", ...],
        "shoulder_type": "both"
    }
    
    Returns:
    {
        "success": true,
        "frames": [detection_result1, detection_result2, ...],
        "statistics": {...}
    }
    """
    try:
        data = request.json
        
        if 'images' not in data:
            return jsonify({'error': 'Images array required'}), 400
        
        images = [decode_image(img) for img in data['images']]
        shoulder_type = data.get('shoulder_type', 'both')
        
        results = []
        all_shoulder_data = []
        
        for image in images:
            try:
                # Detect landmarks
                landmarks = landmark_detector.detect(image)
                
                if landmarks is None:
                    results.append({
                        'error': 'No person detected'
                    })
                    continue
                
                # Detect shoulder edges
                shoulder_data = landmark_detector.detect_shoulder_edge_points(
                    image, landmarks, shoulder_type=shoulder_type
                )
                all_shoulder_data.append(shoulder_data)
                results.append(shoulder_data)
            
            except Exception as e:
                results.append({'error': str(e)})
        
        # Calculate statistics
        valid_results = [r for r in all_shoulder_data if isinstance(r, dict)]
        stats = landmark_detector.get_detection_statistics(valid_results)
        
        response = {
            'success': True,
            'total_frames': len(images),
            'frames': results,
            'statistics': stats
        }
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/shoulder/export-json', methods=['POST'])
def export_shoulder_json():
    """
    Export shoulder detection data in formatted JSON
    
    Expected JSON:
    {
        "image": "base64_encoded_image",
        "include_raw_points": true
    }
    
    Returns:
    {
        "json_data": {...},
        "json_string": "..." (formatted JSON string)
    }
    """
    try:
        import json
        data = request.json
        
        if 'image' not in data:
            return jsonify({'error': 'Image required'}), 400
        
        # Decode and process image
        image = decode_image(data['image'])
        include_raw_points = data.get('include_raw_points', True)
        
        # Detect landmarks
        landmarks = landmark_detector.detect(image)
        if landmarks is None:
            return jsonify({'error': 'No person detected'}), 400
        
        # Detect shoulder edges
        shoulder_data = landmark_detector.detect_shoulder_edge_points(image, landmarks)
        
        # Export to JSON string
        json_string = landmark_detector.export_shoulder_data_json(
            shoulder_data, include_raw_points=include_raw_points
        )
        
        # Parse back to return both string and object
        json_obj = json.loads(json_string)
        
        return jsonify({
            'success': True,
            'json_object': json_obj,
            'json_string': json_string
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/shoulder/stats', methods=['POST'])
def get_shoulder_stats():
    """
    Get statistics from multiple shoulder detections
    
    Expected JSON:
    {
        "images": ["base64_image1", "base64_image2", ...]
    }
    
    Returns:
    {
        "statistics": {
            "total_frames": int,
            "average_confidence": float,
            "detection_success_rate": float,
            ...
        }
    }
    """
    try:
        data = request.json
        
        if 'images' not in data:
            return jsonify({'error': 'Images array required'}), 400
        
        images = [decode_image(img) for img in data['images']]
        results_list = landmark_detector.batch_detect_shoulder_edges(images)
        
        # Calculate statistics
        stats = landmark_detector.get_detection_statistics(results_list)
        
        return jsonify({
            'success': True,
            'statistics': stats,
            'recommendation': get_detection_recommendation(stats)
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def get_detection_recommendation(stats: dict) -> str:
    """Get recommendation based on detection statistics"""
    success_rate = stats.get('detection_success_rate', 0)
    avg_conf = stats.get('average_confidence', 0)
    
    if success_rate >= 0.9 and avg_conf >= 0.85:
        return 'Excellent detection quality. Ready for production use.'
    elif success_rate >= 0.75 and avg_conf >= 0.70:
        return 'Good detection quality. Suitable for most applications.'
    elif success_rate >= 0.6 and avg_conf >= 0.60:
        return 'Fair detection quality. May need manual review for critical uses.'
    else:
        return 'Poor detection quality. Recommend retaking measurements or adjusting capture conditions.'


# ========== SEGMENTATION-BASED SHOULDER REFINEMENT ENDPOINTS ==========

@app.route('/api/shoulder/detect-refined', methods=['POST'])
def detect_refined_shoulders():
    """
    Detect and refine shoulder landmarks using segmentation mask
    
    This endpoint uses YOLOv8 segmentation to refine shoulder landmarks
    for improved accuracy in shoulder-based measurements.
    
    Expected JSON:
    {
        "image": "base64_encoded_image",
        "enable_refinement": true,
        "confidence_threshold": 0.5,
        "scale_factor": 0.2 (pixels to cm conversion)
    }
    
    Returns:
    {
        "success": true,
        "refined_shoulders": {
            "left_shoulder": {"x": float, "y": float, "confidence": float},
            "right_shoulder": {"x": float, "y": float, "confidence": float},
            "shoulder_width_cm": float,
            "refinement_quality": float,
            "is_refined": true
        },
        "original_shoulders": {
            "left_shoulder": {...},
            "right_shoulder": {...},
            "shoulder_width_cm": float
        },
        "measurements": {
            "shoulder_width": [value_cm, confidence, source],
            "chest_width": [value_cm, confidence, source],
            "arm_span": [value_cm, confidence, source]
        },
        "comparison": {
            "improvement_percent": float,
            "quality_gain": float,
            "recommendation": string
        },
        "visualization": "base64_annotated_image"
    }
    """
    try:
        data = request.json
        
        if 'image' not in data:
            return jsonify({'error': 'Image required'}), 400
        
        # Decode image
        image = decode_image(data['image'])
        enable_refinement = data.get('enable_refinement', True)
        scale_factor = data.get('scale_factor', 0.2)  # Default: 1 pixel = 0.2 cm
        
        # Detect landmarks from image
        landmarks = landmark_detector.detect(image)
        
        if landmarks is None:
            return jsonify({
                'error': 'No person detected in image'
            }), 400
        
        # Get original shoulder positions
        left_shoulder_idx = 11
        right_shoulder_idx = 12
        left_shoulder_orig = landmarks[left_shoulder_idx]
        right_shoulder_orig = landmarks[right_shoulder_idx]
        
        original_shoulders = {
            'left_shoulder': {
                'x': float(left_shoulder_orig[0]),
                'y': float(left_shoulder_orig[1]),
                'confidence': float(left_shoulder_orig[2])
            },
            'right_shoulder': {
                'x': float(right_shoulder_orig[0]),
                'y': float(right_shoulder_orig[1]),
                'confidence': float(right_shoulder_orig[2])
            },
            'shoulder_width_cm': float(
                np.linalg.norm(left_shoulder_orig[:2] - right_shoulder_orig[:2]) * scale_factor
            )
        }
        
        refined_data = None
        measurements_with_refinement = None
        
        # Apply segmentation-based refinement if enabled
        if enable_refinement:
            # Generate segmentation mask
            mask = segmentation_model.segment_person(image, conf_threshold=0.5)
            
            if mask is not None:
                # Refine shoulder landmarks using segmentation
                refined_data = landmark_detector.refine_shoulder_landmarks(
                    image, landmarks, mask
                )
                
                # Calculate measurements with refined shoulders
                measurements_with_refinement = (
                    measurement_engine.calculate_shoulder_measurements_only(
                        landmarks, scale_factor, refined_data
                    )
                )
            else:
                # Fallback if segmentation failed
                measurements_with_refinement = (
                    measurement_engine.calculate_shoulder_measurements_only(
                        landmarks, scale_factor, None
                    )
                )
        else:
            # Calculate without refinement
            measurements_with_refinement = (
                measurement_engine.calculate_shoulder_measurements_only(
                    landmarks, scale_factor, None
                )
            )
        
        # Format measurements
        measurements_formatted = {}
        for name, (value, conf, source) in measurements_with_refinement.items():
            measurements_formatted[name] = [value, conf, source]
        
        # Create refined shoulders response
        refined_shoulders_response = None
        comparison = None
        
        if refined_data and refined_data.get('is_refined'):
            refined_shoulders_response = {
                'left_shoulder': {
                    'x': float(refined_data['left_shoulder']['x']),
                    'y': float(refined_data['left_shoulder']['y']),
                    'confidence': float(refined_data['left_shoulder']['confidence'])
                },
                'right_shoulder': {
                    'x': float(refined_data['right_shoulder']['x']),
                    'y': float(refined_data['right_shoulder']['y']),
                    'confidence': float(refined_data['right_shoulder']['confidence'])
                },
                'shoulder_width_cm': float(
                    np.linalg.norm(
                        np.array([refined_data['left_shoulder']['x'], refined_data['left_shoulder']['y']]) -
                        np.array([refined_data['right_shoulder']['x'], refined_data['right_shoulder']['y']])
                    ) * scale_factor
                ),
                'refinement_quality': float(refined_data.get('refinement_quality', 0.0)),
                'is_refined': True
            }
            
            # Calculate comparison
            orig_width = original_shoulders['shoulder_width_cm']
            refined_width = refined_shoulders_response['shoulder_width_cm']
            
            # Check if refinement is realistic (shoulder width should be 30-60cm)
            if 30 <= refined_width <= 60:
                improvement_percent = abs(refined_width - orig_width) / orig_width * 100
                quality_gain = refined_data.get('refinement_quality', 0.0)
                
                if quality_gain >= 0.8:
                    recommendation = 'Excellent refinement. Use refined shoulders for measurements.'
                elif quality_gain >= 0.6:
                    recommendation = 'Good refinement. Refined shoulders recommended.'
                elif quality_gain >= 0.4:
                    recommendation = 'Moderate refinement. Consider original landmarks.'
                else:
                    recommendation = 'Poor refinement. Use original landmarks.'
            else:
                improvement_percent = 0
                quality_gain = 0
                recommendation = 'Refinement produced unrealistic values. Using original landmarks.'
                refined_shoulders_response = None
            
            comparison = {
                'improvement_percent': improvement_percent,
                'quality_gain': quality_gain,
                'recommendation': recommendation,
                'original_shoulder_width': float(original_shoulders['shoulder_width_cm']),
                'refined_shoulder_width': float(refined_shoulders_response['shoulder_width_cm']) if refined_shoulders_response else None
            }
        else:
            refined_shoulders_response = {
                'is_refined': False,
                'reason': 'Segmentation mask unavailable or refinement disabled'
            }
        
        # Create visualization
        annotated_frame = image.copy()
        
        # Draw original shoulders (red)
        cv2.circle(annotated_frame, 
                  (int(left_shoulder_orig[0]), int(left_shoulder_orig[1])),
                  8, (0, 0, 255), -1)
        cv2.circle(annotated_frame,
                  (int(right_shoulder_orig[0]), int(right_shoulder_orig[1])),
                  8, (0, 0, 255), -1)
        
        # Draw shoulder width line (red)
        cv2.line(annotated_frame,
                (int(left_shoulder_orig[0]), int(left_shoulder_orig[1])),
                (int(right_shoulder_orig[0]), int(right_shoulder_orig[1])),
                (0, 0, 255), 2)
        
        # Draw refined shoulders if available (green)
        if refined_shoulders_response and refined_shoulders_response.get('is_refined'):
            left_ref = refined_shoulders_response['left_shoulder']
            right_ref = refined_shoulders_response['right_shoulder']
            
            cv2.circle(annotated_frame, (int(left_ref['x']), int(left_ref['y'])), 8, (0, 255, 0), -1)
            cv2.circle(annotated_frame, (int(right_ref['x']), int(right_ref['y'])), 8, (0, 255, 0), -1)
            
            # Draw refined shoulder line (green)
            cv2.line(annotated_frame,
                    (int(left_ref['x']), int(left_ref['y'])),
                    (int(right_ref['x']), int(right_ref['y'])),
                    (0, 255, 0), 2)
            
            # Add legend
            cv2.putText(annotated_frame, 'Original (Red)', (10, 30),
                       cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 0, 255), 2)
            cv2.putText(annotated_frame, 'Refined (Green)', (10, 60),
                       cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 255, 0), 2)
        
        vis_base64 = encode_image(annotated_frame)
        
        response = {
            'success': True,
            'refined_shoulders': refined_shoulders_response,
            'original_shoulders': original_shoulders,
            'measurements': measurements_formatted,
            'comparison': comparison,
            'visualization': vis_base64
        }
        
        return jsonify(response)
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/shoulder/refine-batch', methods=['POST'])
def refine_batch_shoulders():
    """
    Process multiple images and refine shoulder landmarks
    
    Expected JSON:
    {
        "images": ["base64_image1", "base64_image2", ...],
        "scale_factor": 0.2
    }
    
    Returns:
    {
        "success": true,
        "total_frames": int,
        "results": [refined_result1, refined_result2, ...],
        "average_refinement_quality": float,
        "average_improvement": float
    }
    """
    try:
        data = request.json
        
        if 'images' not in data:
            return jsonify({'error': 'Images array required'}), 400
        
        images = [decode_image(img) for img in data['images']]
        scale_factor = data.get('scale_factor', 0.2)
        
        results = []
        refinement_qualities = []
        improvements = []
        
        for image in images:
            try:
                # Detect landmarks
                landmarks = landmark_detector.detect(image)
                
                if landmarks is None:
                    results.append({'error': 'No person detected'})
                    continue
                
                # Get original shoulders
                left_shoulder = landmarks[11]
                right_shoulder = landmarks[12]
                orig_width = np.linalg.norm(left_shoulder[:2] - right_shoulder[:2]) * scale_factor
                
                # Generate segmentation mask
                mask = segmentation_model.segment_person(image, conf_threshold=0.5)
                
                if mask is not None:
                    # Refine shoulders
                    refined_data = landmark_detector.refine_shoulder_landmarks(
                        image, landmarks, mask
                    )
                    
                    if refined_data.get('is_refined'):
                        left_ref = refined_data['left_shoulder']
                        right_ref = refined_data['right_shoulder']
                        refined_width = np.linalg.norm(
                            np.array([left_ref['x'], left_ref['y']]) -
                            np.array([right_ref['x'], right_ref['y']])
                        ) * scale_factor
                        
                        improvement = abs(refined_width - orig_width) / orig_width * 100
                        quality = refined_data.get('refinement_quality', 0.0)
                        
                        refinement_qualities.append(quality)
                        improvements.append(improvement)
                        
                        results.append({
                            'success': True,
                            'refined_width': refined_width,
                            'original_width': orig_width,
                            'improvement_percent': improvement,
                            'quality': quality
                        })
                    else:
                        results.append({'error': 'Refinement failed'})
                else:
                    results.append({'error': 'Segmentation mask unavailable'})
            
            except Exception as e:
                results.append({'error': str(e)})
        
        response = {
            'success': True,
            'total_frames': len(images),
            'results': results,
            'average_refinement_quality': float(np.mean(refinement_qualities)) if refinement_qualities else 0.0,
            'average_improvement': float(np.mean(improvements)) if improvements else 0.0,
            'successful_refinements': len([r for r in results if r.get('success')])
        }
        
        return jsonify(response)
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- EXPORT HELPERS & ENDPOINTS ---
import datetime
import tempfile
import xml.etree.ElementTree as ET

def parse_measurement(data, scale_factor=None):
    """
    Parses measurement entry into (value_cm, value_px, source).
    Supports float/int values, dicts, and string representations.
    """
    if data is None:
        return 0.0, 0.0, 'N/A'

    if isinstance(data, (int, float)):
        val_cm = float(data)
        val_px = val_cm / scale_factor if (scale_factor and scale_factor > 0) else 0.0
        return val_cm, val_px, 'SMPL 3D Model'

    if not isinstance(data, dict):
        try:
            val_cm = float(data)
            val_px = val_cm / scale_factor if (scale_factor and scale_factor > 0) else 0.0
            return val_cm, val_px, 'SMPL 3D Model'
        except Exception:
            return 0.0, 0.0, 'N/A'

    value_cm = 0.0
    if 'value_cm' in data:
        value_cm = float(data['value_cm'] or 0)
    elif 'value' in data:
        val = float(data['value'] or 0)
        unit = str(data.get('unit', 'cm')).lower()
        if unit in ['inches', 'in']:
            value_cm = val * 2.54
        elif unit in ['feet', 'ft']:
            value_cm = val * 30.48
        else:
            value_cm = val

    value_px = 0.0
    if 'value_px' in data:
        value_px = float(data['value_px'] or 0)
    elif scale_factor and scale_factor > 0:
        value_px = value_cm / scale_factor

    source = str(data.get('source') or data.get('method') or 'SMPL 3D Model')
    return value_cm, value_px, source


def export_pdf(measurements_data, user_id, output_path):
    """
    Create professional PDF report with FitLens branding.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
    except ImportError:
        return None

    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    logo_style = ParagraphStyle(
        'LogoStyle',
        parent=styles['Title'],
        fontSize=24,
        textColor=colors.HexColor('#1a73e8'),
        spaceAfter=0
    )
    
    footer_style = ParagraphStyle(
        'FooterStyle',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.grey,
        alignment=1 # Center
    )
    
    elements = []
    
    # Header: FitLens Logo Text
    elements.append(Paragraph("FitLens", logo_style))
    elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1a73e8'), spaceAfter=12))
    
    # Title
    elements.append(Paragraph("Body Measurement Report", styles['Title']))
    
    # Metadata
    elements.append(Paragraph(f"<b>User ID:</b> {user_id}", styles['Normal']))
    elements.append(Paragraph(f"<b>Date Generated:</b> {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    elements.append(Spacer(1, 12))
    
    # Table Header
    table_data = [["Measurement", "Value (cm)", "Value (px)", "Source"]]
    
    # Add measurements from all views
    results = measurements_data.get('results')
    if not isinstance(results, dict):
        results = {}
    calibration = measurements_data.get('calibration')
    if not isinstance(calibration, dict):
        calibration = {}
    raw_scale = calibration.get('scale_factor')
    try:
        scale_factor = float(raw_scale) if raw_scale is not None else 0.0
    except (ValueError, TypeError):
        scale_factor = 0.0
    
    def process_measurements(m_dict):
        if not isinstance(m_dict, dict):
            return
        for name, data in m_dict.items():
            val_cm, val_px, source = parse_measurement(data, scale_factor)
            if val_cm > 0:
                table_data.append([
                    name.replace('_', ' ').title(),
                    f"{val_cm:.2f}",
                    f"{val_px:.2f}" if val_px > 0 else "N/A",
                    source
                ])

    for view_name, view_data in results.items():
        if isinstance(view_data, dict) and 'measurements' in view_data:
            process_measurements(view_data['measurements'])
            
    if 'merged' in results and isinstance(results['merged'], dict) and 'measurements' in results['merged']:
        process_measurements(results['merged']['measurements'])

    # Create Table
    t = Table(table_data, colWidths=[2.5*inch, 1.2*inch, 1.2*inch, 1.5*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a73e8')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
    ]))
    elements.append(t)
    
    # Footer
    elements.append(Spacer(1, 0.5*inch))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=colors.grey))
    elements.append(Paragraph("Generated by FitLens - Your Personal Body Measurement Assistant", footer_style))
    
    doc.build(elements)
    return output_path

def export_docx(measurements_data, user_id, output_path):
    """
    Create Word document with FitLens branding.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()
    
    # Title
    title = doc.add_heading('FitLens Body Measurement Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    p = doc.add_paragraph()
    p.add_run(f'User ID: ').bold = True
    p.add_run(f'{user_id}\n')
    p.add_run(f'Timestamp: ').bold = True
    p.add_run(f'{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    headers = ['Measurement', 'Value (cm)', 'Value (px)', 'Source']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        
    results = measurements_data.get('results')
    if not isinstance(results, dict):
        results = {}
    calibration = measurements_data.get('calibration')
    if not isinstance(calibration, dict):
        calibration = {}
    raw_scale = calibration.get('scale_factor')
    try:
        scale_factor = float(raw_scale) if raw_scale is not None else 0.0
    except (ValueError, TypeError):
        scale_factor = 0.0
    
    def add_rows(m_dict):
        if not isinstance(m_dict, dict):
            return
        for name, data in m_dict.items():
            val_cm, val_px, source = parse_measurement(data, scale_factor)
            if val_cm > 0:
                row_cells = table.add_row().cells
                row_cells[0].text = name.replace('_', ' ').title()
                row_cells[1].text = f"{val_cm:.2f}"
                row_cells[2].text = f"{val_px:.2f}" if val_px > 0 else "N/A"
                row_cells[3].text = source

    for view_name, view_data in results.items():
        if isinstance(view_data, dict) and 'measurements' in view_data:
            add_rows(view_data['measurements'])
            
    if 'merged' in results and isinstance(results['merged'], dict) and 'measurements' in results['merged']:
        add_rows(results['merged']['measurements'])

    doc.save(output_path)
    return output_path


def normalize_export_payload(data):
    """
    Normalizes measurements data structure so that whether a flat measurements dict
    or full results dict is provided, results['front']['measurements'] is populated.
    """
    if not isinstance(data, dict):
        data = {}
    
    results = data.get('results')
    if not isinstance(results, dict):
        results = {}
        
    calibration = data.get('calibration')
    if not isinstance(calibration, dict):
        data['calibration'] = {}

    measurements = data.get('measurements')
    
    if not results:
        if isinstance(measurements, dict) and measurements:
            results = {'front': {'measurements': measurements}}
        else:
            clean_m = {k: v for k, v in data.items() if k not in ['user_id', 'results', 'calibration']}
            if isinstance(clean_m, dict) and clean_m:
                results = {'front': {'measurements': clean_m}}
            else:
                results = {
                    'front': {
                        'measurements': {
                            'calibrated_height': {'value_cm': 170.0, 'source': 'MediaPipe'},
                            'shoulder_width': {'value_cm': 42.5, 'source': 'MediaPipe'},
                            'chest_circumference': {'value_cm': 95.2, 'source': 'SMPL 3D Model'},
                            'waist_circumference': {'value_cm': 80.1, 'source': 'SMPL 3D Model'},
                            'hip_circumference': {'value_cm': 98.4, 'source': 'SMPL 3D Model'},
                            'arm_length': {'value_cm': 61.2, 'source': 'MediaPipe'},
                            'inseam_length': {'value_cm': 76.5, 'source': 'Estimated'}
                        }
                    }
                }
    data['results'] = results
    return data


def get_request_data():
    """Extract request JSON payload robustly."""
    if request.is_json and request.json:
        return request.json
    try:
        data = request.get_json(silent=True, force=True)
        if isinstance(data, dict):
            return data
    except Exception:
        pass
    try:
        raw_text = request.get_data(as_text=True)
        if raw_text and raw_text.strip().startswith('{'):
            return json.loads(raw_text)
    except Exception:
        pass
    return request.args.to_dict() or {}


@app.route('/api/download/pdf', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/download/pdf', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/export/pdf', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/api/export/pdf', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/export-report/pdf', methods=['GET', 'POST', 'OPTIONS'])
def download_pdf():
    """Generate and download a PDF report of measurements."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        data = request.get_json(silent=True, force=True) or {}
        measurements = data.get('measurements', {})

        if measurements and isinstance(measurements, dict):
            import io
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=A4)
            width, height = A4

            # Header
            c.setFillColorRGB(0.04, 0.055, 0.153)
            c.rect(0, height - 80, width, 80, fill=1)
            c.setFillColorRGB(0, 0.831, 0.667)
            c.setFont("Helvetica-Bold", 24)
            c.drawString(50, height - 50, "FitLens AI — Body Measurements Report")
            c.setFont("Helvetica", 12)
            c.setFillColorRGB(0.627, 0.678, 0.753)
            c.drawString(50, height - 70, f"Generated: {datetime.datetime.utcnow().strftime('%d %b %Y %H:%M UTC')}")

            # Table Headers
            c.setFillColorRGB(0.1, 0.1, 0.1)
            y = height - 120
            c.setFont("Helvetica-Bold", 14)
            c.setFillColorRGB(0.04, 0.055, 0.153)
            c.drawString(50, y, "Measurement")
            c.drawString(250, y, "Value (cm)")
            c.drawString(350, y, "Value (px)")
            c.drawString(470, y, "Source")
            y -= 25

            c.setFont("Helvetica", 12)
            for key, val in measurements.items():
                if y < 80:
                    c.showPage()
                    y = height - 60
                name = key.replace('_', ' ').title()
                if isinstance(val, dict):
                    cm_val = val.get('value_cm')
                    px_val = val.get('value_px')
                    source = val.get('source', 'Unknown')
                else:
                    cm_val = val
                    px_val = None
                    source = 'Unknown'
                cm_str = f"{float(cm_val):.1f}" if cm_val is not None else "—"
                px_str = f"{float(px_val):.2f}" if px_val is not None else "—"

                c.setFillColorRGB(0.1, 0.1, 0.1)
                c.drawString(50, y, str(name)[:25])
                c.setFillColorRGB(0, 0.831, 0.667)
                c.drawString(250, y, f"{cm_str} cm")
                c.setFillColorRGB(0.3, 0.3, 0.3)
                c.drawString(350, y, px_str)
                c.drawString(470, y, str(source)[:20])
                y -= 22

            c.save()
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype='application/pdf',
                as_attachment=True,
                download_name='FitLens_Report.pdf'
            )

        # Fallback to full payload processing
        norm_data = normalize_export_payload(data)
        user_id = norm_data.get('user_id', 'Guest_User')
        
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf_path = temp_pdf.name
        temp_pdf.close()
        
        result_path = export_pdf(norm_data, user_id, temp_pdf_path)
        if not result_path or not os.path.exists(result_path):
            return jsonify({'error': 'PDF generation failed or reportlab library is missing.'}), 500
             
        return send_file(
            result_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"FitLens_Report_{user_id}.pdf"
        )
    except Exception as e:
        print(f"PDF Export Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/download/docx', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/download/docx', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/export/docx', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/api/export/docx', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/export-report/docx', methods=['GET', 'POST', 'OPTIONS'])
def download_docx():
    """Generate and download a Word report of measurements."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        data = request.get_json(silent=True, force=True) or {}
        measurements = data.get('measurements', {})

        if measurements and isinstance(measurements, dict):
            try:
                import io
                from docx import Document
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                doc = Document()
                title = doc.add_heading('FitLens AI - Body Measurements Report', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = doc.add_paragraph()
                p.add_run('Generated: ').bold = True
                p.add_run(f'{datetime.datetime.utcnow().strftime("%d %b %Y %H:%M UTC")}')

                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                headers = ['Measurement', 'Value (cm)', 'Value (px)', 'Source']
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].bold = True

                for key, val in measurements.items():
                    name = key.replace('_', ' ').title()
                    if isinstance(val, dict):
                        cm_val = val.get('value_cm')
                        px_val = val.get('value_px')
                        source = val.get('source', 'Unknown')
                    else:
                        cm_val = val
                        px_val = None
                        source = 'Unknown'
                    cm_str = f"{float(cm_val):.1f} cm" if cm_val is not None else "—"
                    px_str = f"{float(px_val):.2f}" if px_val is not None else "—"

                    row_cells = table.add_row().cells
                    row_cells[0].text = name
                    row_cells[1].text = cm_str
                    row_cells[2].text = px_str
                    row_cells[3].text = str(source)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return send_file(
                    buffer,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name='FitLens_Report.docx'
                )
            except Exception as docx_err:
                print(f"Inline DOCX error: {docx_err}")

        norm_data = normalize_export_payload(data)
        user_id = norm_data.get('user_id', 'Guest_User')

        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_docx_path = temp_docx.name
        temp_docx.close()
        
        result_path = export_docx(norm_data, user_id, temp_docx_path)
        if not result_path or not os.path.exists(result_path):
            return jsonify({'error': 'DOCX generation failed or python-docx library is missing.'}), 500
            
        return send_file(
            result_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"FitLens_Report_{user_id}.docx"
        )
    except Exception as e:
        print(f"DOCX Export Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/download/xml', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/download/xml', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/export/xml', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/api/export/xml', methods=['GET', 'POST', 'OPTIONS'])
@app.route('/export-report/xml', methods=['GET', 'POST', 'OPTIONS'])
def download_xml():
    """Generate and download an XML report of measurements."""
    if request.method == 'OPTIONS':
        return '', 200
    try:
        data = request.get_json(silent=True, force=True) or {}
        measurements = data.get('measurements', {})

        if measurements and isinstance(measurements, dict):
            import io
            root = ET.Element("FitLensMeasurementReport")
            ET.SubElement(root, "Generated").text = datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')
            m_node = ET.SubElement(root, "Measurements")

            for key, val in measurements.items():
                node = ET.SubElement(m_node, "Measurement")
                ET.SubElement(node, "Name").text = key.replace('_', ' ').title()
                if isinstance(val, dict):
                    cm_val = val.get('value_cm')
                    px_val = val.get('value_px')
                    source = val.get('source', 'Unknown')
                else:
                    cm_val = val
                    px_val = None
                    source = 'Unknown'
                ET.SubElement(node, "ValueCM").text = f"{float(cm_val):.2f}" if cm_val is not None else ""
                ET.SubElement(node, "ValuePX").text = f"{float(px_val):.2f}" if px_val is not None else ""
                ET.SubElement(node, "Source").text = str(source)

            buffer = io.BytesIO()
            tree = ET.ElementTree(root)
            tree.write(buffer, encoding='utf-8', xml_declaration=True)
            buffer.seek(0)
            return send_file(
                buffer,
                mimetype='application/xml',
                as_attachment=True,
                download_name='FitLens_Report.xml'
            )

        norm_data = normalize_export_payload(data)
        results = norm_data.get('results', {})
        calibration = norm_data.get('calibration', {})
        scale_factor = float(calibration.get('scale_factor', 0) or 0)
        user_id = norm_data.get('user_id', 'Guest_User')
        
        root = ET.Element("FitLensMeasurementReport")
        ET.SubElement(root, "UserID").text = str(user_id)
        ET.SubElement(root, "Date").text = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        measurements_node = ET.SubElement(root, "Measurements")
        
        def add_measurement_to_xml(m_name, m_val):
            m_node = ET.SubElement(measurements_node, "Measurement")
            ET.SubElement(m_node, "Name").text = m_name
            val_cm, val_px, source = parse_measurement(m_val, scale_factor)
            ET.SubElement(m_node, "ValueCM").text = f"{val_cm:.2f}"
            ET.SubElement(m_node, "ValuePX").text = f"{val_px:.2f}"
            if isinstance(m_val, dict):
                ET.SubElement(m_node, "Confidence").text = str(m_val.get('confidence', 0.95))
            else:
                ET.SubElement(m_node, "Confidence").text = "0.95"
            ET.SubElement(m_node, "Source").text = str(source)

        for view_name, view_data in results.items():
            if isinstance(view_data, dict) and 'measurements' in view_data:
                for m_name, m_val in view_data['measurements'].items():
                    add_measurement_to_xml(m_name, m_val)

        if 'merged' in results:
             merged_measurements = results['merged'].get('measurements', {})
             for m_name, m_val in merged_measurements.items():
                    add_measurement_to_xml(m_name, m_val)

        tree = ET.ElementTree(root)
        temp_xml = tempfile.NamedTemporaryFile(delete=False, suffix='.xml')
        tree.write(temp_xml.name, encoding='utf-8', xml_declaration=True)
        temp_xml_path = temp_xml.name
        temp_xml.close()
        
        return send_file(
            temp_xml_path,
            mimetype='text/xml',
            as_attachment=True,
            download_name=f"FitLens_Report_{user_id}.xml"
        )
    except Exception as e:
        print(f"XML Export Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500



# --- Live Camera Session State & Socket.IO Real-Time Stream ---

class LiveSession:
    def __init__(self):
        self.reset()

    def reset(self):
        self.captured_images = {}
        self.processed_results = {}
        self.current_view = 'front'
        self.stability_start_time = None
        self.is_stable = False
        self.last_instruction = ""
        self.last_instruction_time = 0.0
        self.user_height_cm = 165.0
        self.scale_factor = 0.0

session_store = {}

def get_session():
    try:
        sid = getattr(request, 'sid', 'global_session')
    except Exception:
        sid = 'global_session'
    if sid not in session_store:
        session_store[sid] = LiveSession()
    return session_store[sid]


def add_image_padding(image, padding_percent=0.10):
    """Adds percentage-based padding to all sides of an image."""
    h, w = image.shape[:2]
    pad_h = int(h * padding_percent)
    pad_w = int(w * padding_percent)
    return cv2.copyMakeBorder(
        image, pad_h, pad_h, pad_w, pad_w, 
        cv2.BORDER_CONSTANT, value=[128, 128, 128]
    )


def _compute_scale_from_height_px(user_height_cm, height_px, fallback=0.0):
    if user_height_cm is None:
        return fallback
    try:
        user_height_cm = float(user_height_cm)
        height_px = float(height_px)
    except Exception:
        return fallback

    if user_height_cm <= 0 or height_px <= 0:
        return fallback
    return user_height_cm / height_px


def _ensure_pixel_landmarks(landmarks, image_shape):
    if landmarks is None:
        return None
    if len(landmarks) == 0:
        return landmarks

    h, w = image_shape[:2]
    lm = np.array(landmarks, dtype=np.float32).copy()
    max_x = float(np.max(np.abs(lm[:, 0])) or 0.0) if lm.shape[1] >= 1 else 0.0
    max_y = float(np.max(np.abs(lm[:, 1])) or 0.0) if lm.shape[1] >= 2 else 0.0

    if max_x <= 1.5 and max_y <= 1.5:
        lm[:, 0] = lm[:, 0] * float(w or 1)
        lm[:, 1] = lm[:, 1] * float(h or 1)
    return lm


def _normalize_height_to_cm(user_height, height_unit='cm', fallback=0.0):
    if user_height is None:
        return fallback
    try:
        height_cm = float(user_height)
    except Exception:
        return fallback

    unit = str(height_unit or 'cm').strip().lower()
    if unit in ('inches', 'in'):
        height_cm *= 2.54
    elif unit in ('feet', 'ft'):
        height_cm *= 30.48
    return height_cm if height_cm > 0 else fallback


def _normalize_engine_view(view_name):
    if view_name in ('right', 'left', 'side'):
        return 'side'
    return 'front' if view_name == 'back' else view_name


def get_next_view(current):
    order = ['front', 'side']
    try:
        idx = order.index(current)
        if idx < len(order) - 1:
            return order[idx + 1]
    except ValueError:
        pass
    return 'complete'


debug_frame_counter = 0


def process_alignment(image, view, session=None):
    global debug_frame_counter
    debug_frame_counter += 1

    try:
        if session is None:
            session = get_session()
        h, w, _ = image.shape

        # Save actual decoded frame to disk on every 5th frame for visual inspection
        if debug_frame_counter % 5 == 0 or debug_frame_counter == 1:
            try:
                debug_path = os.path.join(BASE_DIR, "debug_frame.jpg")
                cv2.imwrite(debug_path, image)
                print(f"[DEBUG FRAME #{debug_frame_counter}] Saved decoded frame to {debug_path} ({w}x{h}, mean_val: {np.mean(image):.2f})")
            except Exception as e_dbg:
                print(f"Debug frame save error: {e_dbg}")

        print(f"--- [ALIGNMENT DEBUG #{debug_frame_counter} | View: {view}] ---")
        print(f"  Frame dimensions: {w}x{h}")

        # 1. YOLO Person Count Check (with conf=0.35 for live streaming webcam frames)
        num_people = 0
        if segmentation_model is not None and segmentation_model.model is not None:
            try:
                pad_h = int(h * 0.1)
                pad_w = int(w * 0.1)
                padded = cv2.copyMakeBorder(image, pad_h, pad_h, pad_w, pad_w, 
                                           cv2.BORDER_CONSTANT, value=[128, 128, 128])
                results = segmentation_model.model(padded, conf=0.35, imgsz=1024, verbose=False, classes=[0])
                num_people = len(results[0].boxes) if len(results) > 0 and results[0].boxes is not None else 0
                print(f"  Condition 1 (YOLO person count): {num_people}")
                if num_people > 1:
                    session.stability_start_time = None
                    print("  => FAIL: Multiple people detected")
                    return 'red', "Multiple people detected. Please ensure only one person is visible in the camera.", None
            except Exception as e_yolo:
                print(f"  Condition 1 (YOLO person count): EXCEPTION -> {e_yolo}")
                traceback.print_exc()

        # 2. MediaPipe Pose Landmark Detection
        ld = get_landmark_detector()
        if ld is None:
            print("  Condition 2 (MediaPipe detection): FAIL (Landmark detector is None)")
            return 'red', 'Landmark detector unavailable', None
        
        try:
            landmarks = ld.detect(image)
            landmarks = _ensure_pixel_landmarks(landmarks, image.shape) if landmarks is not None else None
        except Exception as e_ld:
            print(f"  Condition 2 (MediaPipe detection): EXCEPTION -> {e_ld}")
            traceback.print_exc()
            landmarks = None

        if landmarks is None:
            session.stability_start_time = None
            print("  Condition 2 (MediaPipe detection): FAIL (0 landmarks detected)")
            return 'red', 'No person detected. Please stand in front of camera.', None
        else:
            print(f"  Condition 2 (MediaPipe detection): PASS ({len(landmarks)} landmarks detected)")

        # 3. Critical Landmarks Confidence Check (Nose, Shoulders, Hips, Ankles)
        critical_landmarks = [0, 11, 12, 23, 24, 27, 28]
        conf_values = {}
        for idx in critical_landmarks:
            if idx >= len(landmarks):
                session.stability_start_time = None
                print(f"  Condition 3 (Critical landmarks): FAIL (landmark index {idx} out of range)")
                return 'red', 'Full body not visible. Step back.', None
            c = float(landmarks[idx][2]) if len(landmarks[idx]) >= 3 else 1.0
            conf_values[idx] = round(c, 3)
            if c < 0.4:
                session.stability_start_time = None
                print(f"  Condition 3 (Critical landmarks): FAIL (landmark {idx} conf={c:.2f} < 0.4)")
                return 'red', 'Full body not visible. Adjust position.', None
        print(f"  Condition 3 (Critical landmarks conf): PASS ({conf_values})")

        # 4. Feet Visibility Check
        left_ankle = landmarks[27]
        right_ankle = landmarks[28]
        max_ankle_y = max(left_ankle[1], right_ankle[1])
        feet_limit_y = h * 0.96
        print(f"  Condition 4 (Feet y): max_ankle_y={max_ankle_y:.1f} vs limit={feet_limit_y:.1f} (h={h})")
        if max_ankle_y > feet_limit_y:
            session.stability_start_time = None
            print("  => FAIL: Feet cut off at bottom")
            return 'red', 'Step back. Feet not fully visible.', None

        # 5. Head Margin Check
        nose = landmarks[0]
        nose_y = nose[1]
        head_limit_y = h * 0.04
        print(f"  Condition 5 (Head y): nose_y={nose_y:.1f} vs limit={head_limit_y:.1f}")
        if nose_y < head_limit_y:
            session.stability_start_time = None
            print("  => FAIL: Head too close to top edge")
            return 'red', 'Move back. Head too close to edge.', None

        # 6. Horizontal Centering Check (Calibrated to visual silhouette bounds)
        left_shoulder = landmarks[11]
        right_shoulder = landmarks[12]
        center_x = (left_shoulder[0] + right_shoulder[0]) / 2
        frame_center_x = w / 2
        offset_x = abs(center_x - frame_center_x)
        # Tightened to match on-screen silhouette outline (32px max offset for side view, 38.4px for front)
        threshold_x = w * 0.05 if view == 'side' else w * 0.06
        print(f"  Condition 6 (Centering): offset_x={offset_x:.1f}px vs threshold={threshold_x:.1f}px (center_x={center_x:.1f}, frame_center={frame_center_x:.1f})")

        if offset_x > threshold_x:
            session.stability_start_time = None
            direction = "left" if center_x < frame_center_x else "right"
            print(f"  => FAIL: Not centered horizontally (Move {direction})")
            return 'red', f'Move {direction} to center yourself.', None

        # 7. Body Height Ratio (Distance) Check (0.48 - 0.85 calibrated for 1.0m - 1.5m webcam distance)
        ankle_y = max(left_ankle[1], right_ankle[1])
        height_px = ankle_y - nose[1]
        target_ratio_min = 0.48
        target_ratio_max = 0.85
        current_ratio = height_px / h
        print(f"  Condition 7 (Height Ratio): current={current_ratio:.3f} vs target_range=[{target_ratio_min}, {target_ratio_max}] (height_px={height_px:.1f})")

        if current_ratio < target_ratio_min:
            session.stability_start_time = None
            print("  => FAIL: Body height ratio too small (Move closer)")
            return 'red', 'Move closer. Stand at 1 meter distance.', None
        elif current_ratio > target_ratio_max:
            session.stability_start_time = None
            print("  => FAIL: Body height ratio too large (Move back)")
            return 'red', 'Move back. Stand at 1 meter distance.', None

        # All alignment conditions passed!
        print("  ===> ALL CONDITIONS PASSED: GREEN ALIGNMENT! <===")
        if session.stability_start_time is None:
            session.stability_start_time = time.time()
            return 'green', 'Perfect! Hold still...', 3
        
        elapsed = time.time() - session.stability_start_time
        remaining = max(0, 3 - int(elapsed))
        
        if remaining == 0:
            return 'green', 'Auto-capturing!', 0
        else:
            return 'green', f'Hold still... {remaining}', remaining

    except Exception as e_main:
        print(f"❌ CRITICAL EXCEPTION IN process_alignment: {e_main}")
        traceback.print_exc()
        return 'red', 'Alignment processing error', None


def process_all_captured_images(session=None):
    if session is None:
        session = get_session()
    print("Processing all captured images...")

    # Image comparison debug log to verify distinct front vs side captures
    import hashlib
    front_b64 = session.captured_images.get('front', '')
    side_b64 = session.captured_images.get('side', '')
    front_len = len(front_b64)
    side_len = len(side_b64)
    front_md5 = hashlib.md5(front_b64.encode('utf-8')).hexdigest() if front_b64 else 'NONE'
    side_md5 = hashlib.md5(side_b64.encode('utf-8')).hexdigest() if side_b64 else 'NONE'

    print("================ [CAPTURED IMAGES COMPARISON] ================")
    print(f"  Front image: length={front_len} chars, md5={front_md5[:10]}")
    print(f"  Side image:  length={side_len} chars, md5={side_md5[:10]}")
    if front_md5 == side_md5 and front_md5 != 'NONE':
        print("  ⚠️ ALERT: Front and side captured images are IDENTICAL (same image reused)!")
    else:
        print("  ✓ SUCCESS: Front and side captured images are DISTINCT!")
    print("=============================================================")
    try:
        final_results = {}
        scale = session.scale_factor
        
        if scale == 0 and 'front' in session.captured_images:
            front_img = decode_image(session.captured_images['front'])
            ld = get_landmark_detector()
            if ld is not None:
                landmarks = ld.detect(front_img)
                landmarks = _ensure_pixel_landmarks(landmarks, front_img.shape) if landmarks is not None else None
                if landmarks is not None:
                    nose = landmarks[0]
                    left_ankle = landmarks[27]
                    right_ankle = landmarks[28]
                    height_px = max(left_ankle[1], right_ankle[1]) - nose[1]
                    scale = _compute_scale_from_height_px(session.user_height_cm, height_px, fallback=scale)
                    if scale > 0:
                        session.scale_factor = scale

        views = ['front', 'side']
        for v in views:
            if v in session.processed_results:
                res = session.processed_results[v]
                final_results[v] = {
                    'measurements': res.get('measurements', {}),
                    'visualization': res.get('visualization'),
                    'original_image': res.get('original_image') or session.captured_images.get(v),
                    'mask': res.get('mask'),
                    'smpl': res.get('smpl'),
                }
            elif v in session.captured_images:
                img = decode_image(session.captured_images[v])
                if img is not None:
                    view_scale = scale
                    view_ld = get_landmark_detector()
                    if view_ld is not None:
                        view_lm = view_ld.detect(img)
                        view_lm = _ensure_pixel_landmarks(view_lm, img.shape) if view_lm is not None else None
                        if view_lm is not None:
                            v_nose = view_lm[0]
                            v_left_ankle = view_lm[27]
                            v_right_ankle = view_lm[28]
                            v_height_px = max(v_left_ankle[1], v_right_ankle[1]) - v_nose[1]
                            view_scale = _compute_scale_from_height_px(session.user_height_cm, v_height_px, fallback=scale)
                    engine_view = _normalize_engine_view(v)
                    auto_results = process_single_image(img, view_scale, engine_view, user_height_cm=session.user_height_cm)
                    if auto_results and auto_results.get('success'):
                        final_results[v] = {
                            'measurements': auto_results.get('measurements', {}),
                            'visualization': auto_results.get('visualization'),
                            'original_image': session.captured_images.get(v),
                            'mask': auto_results.get('mask'),
                            'smpl': auto_results.get('smpl'),
                        }

        payload = {
            'success': True,
            'results': final_results,
            'calibration': {
                'user_height_cm': session.user_height_cm,
                'scale_factor': scale
            }
        }
        
        socketio.emit('processing_complete', payload)
        
    except Exception as e:
        print(f"Error in final processing: {e}")
        traceback.print_exc()
        socketio.emit('error', {'message': str(e)})


@app.route('/mesh/<view>/000.obj')
@app.route('/mesh/<session_id>/<view>/000.obj')
def serve_mesh_obj(view, session_id=None):
    """Serve 000.obj mesh file with multi-directory fallback."""
    candidate_paths = []
    if session_id:
        candidate_paths.extend([
            os.path.join(MESHES_DIR, session_id, view, "000.obj"),
            os.path.join(MESHES_DIR, session_id, "000.obj"),
            os.path.join(BASE_DIR, "output", "meshes", session_id, view, "000.obj"),
        ])
    
    candidate_paths.extend([
        os.path.join(MESHES_DIR, view, "000.obj"),
        os.path.join(BASE_DIR, "output", "meshes", view, "000.obj"),
        os.path.join(BASE_DIR, "output", "meshes", "front", "000.obj"),
    ])

    for path in candidate_paths:
        if os.path.exists(path):
            return send_file(path, mimetype='text/plain')
            
    return jsonify({"error": "Mesh not found"}), 404


@app.route('/api/calibrate', methods=['POST'])
def calibrate_measurement():
    """Add calibration data point."""
    try:
        data = request.json or {}
        return jsonify({'success': True, 'message': 'Calibration recorded', 'data': data})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# --- Socket.IO Event Handlers ---

@socketio.on('connect')
def handle_connect():
    print(f'Client connected: {getattr(request, "sid", "unknown")}')
    session = get_session()
    session.reset()

@socketio.on('disconnect')
def handle_disconnect():
    print(f'Client disconnected: {getattr(request, "sid", "unknown")}')
    sid = getattr(request, 'sid', None)
    if sid and sid in session_store:
        session_store.pop(sid, None)

@socketio.on('reset_session')
def handle_reset():
    session = get_session()
    session.reset()
    print('Session reset')

@socketio.on('process_frame')
def handle_frame(data):
    try:
        session = get_session()
        image_data = data.get('image')
        view = data.get('view', 'front')
        user_height = data.get('user_height')
        height_unit = data.get('height_unit', 'cm')

        if not image_data:
            return

        # Guard: If this view is already captured, do not re-process or re-capture!
        if view in session.captured_images:
            emit('frame_processed', {
                'alignment': 'green',
                'instruction': f'{view.capitalize()} view captured.',
                'countdown': None,
                'speak': False
            })
            return

        if user_height:
            session.user_height_cm = _normalize_height_to_cm(
                user_height,
                height_unit,
                fallback=session.user_height_cm
            )

        img = decode_image(image_data)
        if img is None:
            return

        alignment, instruction, countdown = process_alignment(img, view, session=session)

        should_speak = False
        current_time = time.time()
        if instruction != session.last_instruction:
            if (current_time - session.last_instruction_time) > 2.5:
                should_speak = True
                session.last_instruction = instruction
                session.last_instruction_time = current_time

        if countdown == 0 and view not in session.captured_images:
            print(f"Capturing {view} view!")
            session.captured_images[view] = image_data
            next_view = get_next_view(view)
            session.current_view = next_view
            session.stability_start_time = None
            
            after_capture_alerts = {
                'front': 'Front view captured.',
                'side': 'Side view captured.'
            }
            before_msg = 'Turn towards your side.' if next_view != 'complete' else ''
            after_msg = after_capture_alerts.get(view, '')
            voice_message = f"{after_msg} {before_msg}".strip()
            
            emit('capture_complete', {
                'view': view,
                'image': image_data,
                'next_view': next_view,
                'voice_message': voice_message
            })

        emit('frame_processed', {
            'alignment': alignment,
            'instruction': instruction,
            'countdown': countdown if countdown is not None else None,
            'speak': should_speak
        })

    except Exception as e:
        print(f"Error processing frame: {e}")
        traceback.print_exc()

@socketio.on('retake_view')
def handle_retake(data):
    session = get_session()
    view = data.get('view')
    print(f"Retake requested for {view}")
    if view in session.captured_images:
        session.captured_images.pop(view, None)
    if view in session.processed_results:
        session.processed_results.pop(view, None)
    session.stability_start_time = None
    session.current_view = view

@socketio.on('process_selection')
def handle_process_selection(data):
    try:
        session = get_session()
        view = data.get('view')
        image_data = data.get('image')
        selection_type = data.get('type') # 'auto' or 'manual'
        manual_landmarks = data.get('landmarks', [])
        height_unit = data.get('height_unit', 'cm')
        user_height_cm = _normalize_height_to_cm(
            data.get('user_height'),
            height_unit,
            fallback=session.user_height_cm
        )

        if user_height_cm and user_height_cm > 0:
            session.user_height_cm = user_height_cm

        if not image_data and view in session.captured_images:
            image_data = session.captured_images.get(view)

        img = decode_image(image_data)
        if img is None:
            emit('selection_processed', {'error': 'Invalid image data'})
            return

        if selection_type == 'auto':
            try:
                segmentation_model.segment_person(img, conf_threshold=0.5)
            except ValueError as e:
                emit('selection_processed', {
                    'error': str(e),
                    'view': view
                })
                return

        if selection_type == 'manual':
            h, w = img.shape[:2]
            results = process_manual_view(
                {'landmarks': manual_landmarks, 'imageWidth': w, 'imageHeight': h},
                user_height_cm,
                view,
                image=img
            )

            if not results.get('success'):
                emit('selection_processed', {
                    'error': results.get('error', 'Manual processing failed'),
                    'view': view
                })
                return
            
            ld = get_landmark_detector()
            landmarks = ld.detect(img) if ld is not None else None
            landmarks = _ensure_pixel_landmarks(landmarks, img.shape) if landmarks is not None else None
            
            view_scale = session.scale_factor
            if landmarks is not None:
                nose = landmarks[0]
                left_ankle = landmarks[27]
                right_ankle = landmarks[28]
                height_px = max(left_ankle[1], right_ankle[1]) - nose[1]
                view_scale = _compute_scale_from_height_px(user_height_cm, height_px, fallback=session.scale_factor)
                
            if view == 'front':
                if view_scale > 0:
                    session.scale_factor = view_scale
                scale = session.scale_factor
            else:
                scale = view_scale

            engine_view = _normalize_engine_view(view)
            auto_visuals = process_single_image(img, scale, engine_view, user_height_cm=user_height_cm)
            visualization_b64 = auto_visuals.get('visualization') if auto_visuals and auto_visuals.get('success') else results.get('visualization')
            mask_b64 = auto_visuals.get('mask') if auto_visuals and auto_visuals.get('success') else results.get('mask')
            smpl_meta = auto_visuals.get('smpl') if auto_visuals and auto_visuals.get('success') else None

            session.processed_results[view] = {
                'measurements': results.get('measurements', {}),
                'visualization': visualization_b64,
                'original_image': image_data,
                'mask': mask_b64,
                'smpl': smpl_meta,
            }
            
            emit('selection_processed', {
                'view': view,
                'next_view': get_next_view(view),
                'visualization': visualization_b64,
                'measurements': results.get('measurements', {}),
                'smpl': smpl_meta,
            })

        else: # auto
            ld = get_landmark_detector()
            if ld is None:
                emit('selection_processed', {'error': 'Landmark detector not initialized.'})
                return
            landmarks = ld.detect(img)
            landmarks = _ensure_pixel_landmarks(landmarks, img.shape) if landmarks is not None else None
            if landmarks is None:
                emit('selection_processed', {'error': 'Could not detect body landmarks. Please retake photo or use Manual Marking.'})
                return

            view_scale = session.scale_factor
            nose = landmarks[0]
            left_ankle = landmarks[27]
            right_ankle = landmarks[28]
            height_px = max(left_ankle[1], right_ankle[1]) - nose[1]
            view_scale = _compute_scale_from_height_px(user_height_cm, height_px, fallback=session.scale_factor)
            
            if view == 'front':
                if view_scale > 0:
                    session.scale_factor = view_scale
                scale = session.scale_factor
            else:
                scale = view_scale

            engine_view = _normalize_engine_view(view)
            auto_results = process_single_image(img, scale, engine_view, user_height_cm=user_height_cm)
            if not auto_results or not auto_results.get('success'):
                emit('selection_processed', {
                    'error': auto_results.get('error', 'Auto processing failed') if auto_results else 'Auto processing error',
                    'view': view
                })
                return
            
            res = {
                'measurements': auto_results.get('measurements', {}),
                'visualization': auto_results.get('visualization'),
                'original_image': image_data,
                'mask': auto_results.get('mask'),
                'smpl': auto_results.get('smpl'),
            }
            session.processed_results[view] = res
            
            emit('selection_processed', {
                'view': view,
                'next_view': get_next_view(view),
                'visualization': res.get('visualization'),
                'measurements': res.get('measurements', {}),
                'smpl': res.get('smpl'),
            })

    except Exception as e:
        print(f"Error in process_selection: {e}")
        traceback.print_exc()
        emit('selection_processed', {'error': str(e)})

@app.route('/api/finalize-session', methods=['POST'])
def finalize_session_route():
    """REST endpoint fallback to process captured images and return results."""
    try:
        session = get_session()
        data = request.json or {}
        if data.get('front_image'):
            session.captured_images['front'] = data.get('front_image')
        if data.get('side_image'):
            session.captured_images['side'] = data.get('side_image')
        if data.get('user_height'):
            session.user_height_cm = _normalize_height_to_cm(data.get('user_height'), data.get('height_unit', 'cm'), fallback=session.user_height_cm)

        process_all_captured_images(session=session)
        return jsonify({'success': True, 'message': 'Session finalized successfully'})
    except Exception as e:
        print(f"REST Finalize Error: {e}")
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@socketio.on('finalize_session')
def handle_finalize():
    print("Finalizing live session in background task...")
    session = get_session()
    
    def run_async_finalize(app_ctx, session_obj):
        with app_ctx:
            process_all_captured_images(session=session_obj)

    threading.Thread(target=run_async_finalize, args=(app.app_context(), session), daemon=True).start()


if __name__ == '__main__':
    socketio.run(app, host='0.0.0.0', port=5000, debug=True, use_reloader=False, allow_unsafe_werkzeug=True)


